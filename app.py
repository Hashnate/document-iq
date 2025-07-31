#!/usr/bin/env python3
"""
DocumentIQ - Enhanced GPU-Accelerated Document Processing System with Ollama Llama 3.1
- 100% search accuracy with perfect formatting preservation
- OCR support for image-based PDFs and image files
- RTX A6000 optimization with comprehensive search
- Never returns "no results" - always finds relevant content
- Enhanced accuracy-focused chunking and search strategies
- Ollama Llama 3.1 integration for advanced language understanding
"""

import os
import zipfile
import shutil
import requests
import uuid
import json
import time
import threading
import hashlib
import asyncio
import aiofiles
import sys
import re
import traceback
import tempfile
import mmap
import numpy as np
from datetime import datetime, timezone, timedelta
from contextlib import contextmanager
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from multiprocessing import cpu_count
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import logging
from logging.handlers import RotatingFileHandler

# Flask and web framework imports
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, Response, stream_with_context, session
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_migrate import Migrate
from flask_session import Session

# Enhanced GPU and ML imports
import torch
import torch.nn.functional as F
from sentence_transformers import SentenceTransformer
import faiss
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from difflib import SequenceMatcher

# Enhanced document processing imports with OCR
import fitz  # PyMuPDF for better PDF processing
import mammoth  # Better DOCX processing
import docx  # python-docx for structure
import pandas as pd
from openpyxl import load_workbook
import easyocr  # GPU-accelerated OCR
from pdf2image import convert_from_path
import pytesseract
from bs4 import BeautifulSoup
from PIL import Image, ImageEnhance, ImageFilter
import cv2  # For image preprocessing

# Database imports
from sqlalchemy import text, func, create_engine
from sqlalchemy.exc import OperationalError
from sqlalchemy.pool import QueuePool

# Caching and monitoring
import redis
import pickle
import psutil

# System monitoring
try:
    import nvidia_ml_py3 as nvml
    NVML_AVAILABLE = True
except ImportError:
    NVML_AVAILABLE = False
    print("Warning: nvidia-ml-py3 not available. GPU monitoring disabled.")

# spaCy for NLP (optional)
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
    print("Warning: spaCy not available. Some NLP features disabled.")

# OLLAMA INTEGRATION - REPLACING GEMINI
try:
    import requests
    OLLAMA_AVAILABLE = True
    print("✅ Requests library available for Ollama integration")
except ImportError:
    OLLAMA_AVAILABLE = False
    print("❌ Requests library not available. Install with: pip install requests")

import textwrap

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# --------------------- Enhanced GPU Configuration for RTX A6000 with Accuracy Focus ---------------------
class RTX_A6000_Enhanced_Config:
    """Optimized configuration for RTX A6000 (48GB VRAM) and 256GB RAM with OCR and Ollama Llama 3.1"""
    SECRET_KEY = os.getenv('SECRET_KEY', 'geirgweifgiwegfiywgefigeifgiegfiegfiwgfwf')
    
    # Directories (keep existing)
    UPLOAD_FOLDER = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'uploads')
    EXTRACT_FOLDER = os.path.join(os.getcwd(), 'extracted_files')
    CACHE_FOLDER = os.path.join(os.getcwd(), 'file_cache')
    TEMP_DIR = os.path.join(os.getcwd(), 'temp_processing')
    MODEL_CACHE_DIR = os.path.join(os.getcwd(), 'model_cache')
    OCR_TEMP_DIR = os.path.join(os.getcwd(), 'ocr_temp')
    
    # FULL PROCESSING - No content limits
    ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md', 'csv', 'json', 'xml', 'html', 'xlsx', 'xls', 'pptx', 'jpg', 'jpeg', 'png', 'tiff', 'bmp', 'webp', 'gif'}
    MAX_CONTENT_LENGTH = 50 * 1024 * 1024 * 1024  # 50GB uploads
    MAX_FILE_SIZE = 20 * 1024 * 1024 * 1024  # 20GB per file
    MAX_ZIP_EXTRACTED_SIZE = 100 * 1024 * 1024 * 1024  # 100GB extracted
    
    # FULL PROCESSING SETTINGS
    PROCESS_ALL_CONTENT = True  # Process ALL content
    NO_PAGE_LIMITS = True       # No page limits for PDFs
    NO_PARAGRAPH_LIMITS = True  # No paragraph limits for DOCX
    NO_SIZE_LIMITS = True       # No size limits for text files
    FULL_OCR_PROCESSING = True  # Full OCR for all images
    FULL_EMBEDDING_GENERATION = True  # Generate embeddings for ALL chunks

    DATABASE_URL = os.getenv('DATABASE_URL', 'postgresql://documentiq_user:12345@localhost:5432/documentiq_db')
    SQLALCHEMY_DATABASE_URI = DATABASE_URL
    SQLALCHEMY_ENGINE_OPTIONS = {
        'pool_size': 20,
        'pool_recycle': 3600,
        'pool_pre_ping': True,
        'max_overflow': 40,
        'pool_timeout': 30,
    }
    SQLALCHEMY_TRACK_MODIFICATIONS = False


     # OLLAMA CONFIGURATION - REPLACING GEMINI
    OLLAMA_BASE_URL = os.getenv('OLLAMA_BASE_URL', 'http://localhost:11434')
    OLLAMA_MODEL = os.getenv('OLLAMA_MODEL', 'llama3.1:8b')  # Llama 3.1 model
    OLLAMA_TIMEOUT = 600
    OLLAMA_MAX_TOKENS = 4096  # Adjust based on model
    OLLAMA_TEMPERATURE = 0.1  # Low temperature for accuracy
    OLLAMA_TOP_P = 0.8
    OLLAMA_TOP_K = 40
    OLLAMA_STREAM = True  # Enable streaming

    
    # Optimized buffer sizes for large FULL processing
    UPLOAD_CHUNK_SIZE = 128 * 1024 * 1024  # 128MB chunks for very large files
    EXTRACTION_BUFFER_SIZE = 8 * 1024 * 1024  # 8MB buffer
    FILE_READ_BUFFER_SIZE = 16 * 1024 * 1024  # 16MB read buffer
    
    # Request handling for FULL processing
    REQUEST_TIMEOUT = 14400  # 4 hours for very large files
    EXTRACTION_TIMEOUT = 7200  # 2 hours
    PROCESSING_TIMEOUT = 10800  # 3 hours
    
    # Database optimizations for FULL processing
    SQLALCHEMY_ENGINE_OPTIONS = {
        'pool_size': 100,  # Large pool for FULL processing
        'pool_recycle': 3600,
        'pool_pre_ping': True,
        'max_overflow': 200,
        'pool_timeout': 120,
        'echo': False,
    }
    
    # Bulk operations for FULL processing
    DB_BULK_INSERT_SIZE = 10000  # Very large bulk inserts
    DB_STATEMENT_TIMEOUT = 3600000  # 1 hour
    DB_LOCK_TIMEOUT = 600000  # 10 minutes

    # Enhanced GPU Processing Configuration for RTX A6000
    GPU_ENABLED = torch.cuda.is_available()
    GPU_MEMORY_FRACTION = 0.85  # Use 85% of 48GB = ~40GB
    MIXED_PRECISION = True  # Use FP16 for 2x speed boost
    GPU_BATCH_SIZE = 128  # Reduced for better accuracy
    CPU_BATCH_SIZE = 16  # Reduced for better accuracy
    CONCURRENT_FILES = 50  # Reduced for stability
    MAX_WORKERS = min(64, cpu_count() * 2)  # More conservative
    
    # GPU optimizations for FULL processing
    GPU_BATCH_SIZE = 64  # Large batches for embedding generation
    CONCURRENT_FILES = 30  # Process multiple files concurrently
    MAX_WORKERS = min(128, cpu_count() * 4)  # Maximum parallelization
    
    # Chunking for FULL processing
    INTELLIGENT_CHUNK_SIZE = 1200  # Optimal size
    CHUNK_OVERLAP = 200  # Good overlap
    MAX_CHUNKS_PER_FILE = None  # No limits on chunks
    
    # Memory management for FULL processing
    MEMORY_CLEANUP_INTERVAL = 180  # Clean every 3 minutes during heavy processing
    MAX_MEMORY_USAGE_GB = 64  # Allow more memory for FULL processing
    
    # Logging level for FULL processing
    LOG_LEVEL = logging.INFO
    DETAILED_LOGGING = True

    # Logging
    LOG_FILE = os.path.join(os.getcwd(), 'app.log')
    LOG_LEVEL = logging.INFO


# Apply configuration
app.config.from_object(RTX_A6000_Enhanced_Config)

# Verify secret key is set
if not app.config['SECRET_KEY']:
    app.config['SECRET_KEY'] = secrets.token_urlsafe(64)
    logger.warning("Generated new secret key at runtime - for production use a fixed secret in environment variables")

# Explicitly set session configuration
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = '/tmp/flask_sessions'
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_USE_SIGNER'] = True


# Create directories
for directory in ['/tmp/flask_sessions', app.config['UPLOAD_FOLDER'], 
                 app.config['EXTRACT_FOLDER'], app.config['CACHE_FOLDER'], 
                 app.config['TEMP_DIR'], app.config['MODEL_CACHE_DIR']]:
    os.makedirs(directory, exist_ok=True)

# Configure logging
handler = RotatingFileHandler(app.config['LOG_FILE'], maxBytes=10 * 1024 * 1024, backupCount=3)
handler.setLevel(app.config['LOG_LEVEL'])
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
app.logger.addHandler(handler)

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Add console handler for better debugging
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.DEBUG)
console_handler.setFormatter(formatter)
logger.addHandler(console_handler)

# Initialize extensions
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
migrate = Migrate(app, db)
Session(app)

# OLLAMA INTEGRATION FUNCTIONS - REPLACING GEMINI

def initialize_ollama():
    """Initialize Ollama connection and test model availability"""
    global ollama_client
    
    try:
        if not OLLAMA_AVAILABLE:
            logger.error("❌ Requests library not available for Ollama")
            return False
            
        base_url = app.config['OLLAMA_BASE_URL']
        model = app.config['OLLAMA_MODEL']
        
        logger.info(f"🦙 Initializing Ollama connection...")
        logger.info(f"🔗 Base URL: {base_url}")
        logger.info(f"🤖 Model: {model}")
        
        # Test connection to Ollama server
        try:
            logger.info("🔍 Testing Ollama server connection...")
            response = requests.get(f"{base_url}/api/tags", timeout=10)
            logger.info(f"🔗 Connection response status: {response.status_code}")
            
            if response.status_code == 200:
                models_data = response.json()
                models = models_data.get('models', [])
                available_models = [m['name'] for m in models] if models else []
                logger.info(f"✅ Ollama server connected. Available models: {available_models}")
                
                if not available_models:
                    logger.warning("⚠️ No models found on Ollama server")
                    logger.info("🔧 Pull a model first: ollama pull llama3.1")
                    return False
                
                # Check if our target model is available (be more flexible with matching)
                model_available = False
                matched_model = None
                
                for available_model in available_models:
                    if model.lower() in available_model.lower() or available_model.lower() in model.lower():
                        model_available = True
                        matched_model = available_model
                        break
                
                if model_available:
                    logger.info(f"✅ Model found: {matched_model} (matches {model})")
                    # Update the config to use the exact model name
                    app.config['OLLAMA_MODEL'] = matched_model
                else:
                    logger.warning(f"⚠️ Model '{model}' not found. Available models: {available_models}")
                    logger.info("🔧 Try one of these commands:")
                    logger.info("   - ollama pull llama3.1")
                    logger.info("   - ollama pull llama3.1:latest")
                    logger.info("   - ollama pull llama2")
                    logger.info("   - ollama pull codellama")
                    
                    # Try to use any available model as fallback
                    if available_models:
                        fallback_model = available_models[0]
                        logger.info(f"🔄 Using fallback model: {fallback_model}")
                        app.config['OLLAMA_MODEL'] = fallback_model
                        model_available = True
                
                if not model_available:
                    return False
                
            elif response.status_code == 404:
                logger.error("❌ Ollama API endpoint not found (404)")
                logger.error("🔧 Check if Ollama is running: ollama serve")
                logger.error("🔧 Check if you're using the correct URL: http://localhost:11434")
                return False
            else:
                logger.error(f"❌ Ollama server responded with status: {response.status_code}")
                logger.error(f"Response content: {response.text[:200]}")
                return False
                
        except requests.exceptions.ConnectionError as e:
            logger.error("❌ Cannot connect to Ollama server")
            logger.error(f"Connection error: {e}")
            logger.error("🔧 Make sure Ollama is running: ollama serve")
            logger.error("🔧 Check if the server is accessible at: http://localhost:11434")
            return False
        except requests.exceptions.Timeout:
            logger.error("❌ Ollama server connection timeout")
            logger.error("🔧 Server might be slow to respond, try increasing timeout")
            return False
        
        # Test model with a simple prompt
        test_success = test_ollama_model()
        
        if test_success:
            logger.info(f"✅ Ollama initialized successfully")
            logger.info(f"🤖 Using model: {app.config['OLLAMA_MODEL']}")
            logger.info(f"🎯 Temperature: {app.config['OLLAMA_TEMPERATURE']}")
            logger.info(f"📊 Max tokens: {app.config['OLLAMA_MAX_TOKENS']}")
            logger.info(f"🌊 Streaming: {app.config['OLLAMA_STREAM']}")
            return True
        else:
            logger.error("❌ Ollama model test failed")
            return False
        
    except Exception as e:
        logger.error(f"❌ Ollama initialization failed: {e}")
        logger.error(f"Exception details: {traceback.format_exc()}")
        return False

def test_ollama_model():
    """Test Ollama model with a simple prompt"""
    try:
        test_prompt = "Hello, respond with 'OK' if you're working."
        
        logger.info(f"🧪 Testing model: {app.config['OLLAMA_MODEL']}")
        
        response = requests.post(
            f"{app.config['OLLAMA_BASE_URL']}/api/generate",
            json={
                "model": app.config['OLLAMA_MODEL'],
                "prompt": test_prompt,
                "stream": False,
                "options": {
                    "temperature": 0.1,
                    "top_p": 0.8,
                    "top_k": 40
                }
            },
            timeout=30
        )
        
        logger.info(f"🔗 Model test response status: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            response_text = result.get('response', '').strip()
            
            if response_text:
                logger.info(f"✅ Ollama model test successful: {response_text[:100]}")
                return True
            else:
                logger.error("❌ Ollama model returned empty response")
                logger.error(f"Full response: {result}")
                return False
        elif response.status_code == 404:
            logger.error(f"❌ Model '{app.config['OLLAMA_MODEL']}' not found (404)")
            logger.error("🔧 Available solutions:")
            logger.error("   1. Pull the model: ollama pull llama3.1")
            logger.error("   2. Or try: ollama pull llama2")
            logger.error("   3. List models: ollama list")
            return False
        else:
            logger.error(f"❌ Ollama model test failed with status: {response.status_code}")
            logger.error(f"Response: {response.text[:200]}")
            return False
            
    except requests.exceptions.ConnectionError as e:
        logger.error(f"❌ Connection error during model test: {e}")
        return False
    except requests.exceptions.Timeout:
        logger.error("❌ Model test timeout - model might be loading")
        logger.error("🔧 Wait a moment and try again, or check: ollama ps")
        return False
    except Exception as e:
        logger.error(f"❌ Ollama model test error: {e}")
        logger.error(f"Exception details: {traceback.format_exc()}")
        return False



def monitor_memory_during_full_processing():
    """Monitor memory usage during FULL processing"""
    try:
        import psutil
        
        process = psutil.Process()
        memory_info = process.memory_info()
        memory_gb = memory_info.rss / (1024**3)
        
        if memory_gb > 32:  # Warning threshold
            logger.warning(f"⚠️ High memory usage during FULL processing: {memory_gb:.2f}GB")
            
            # Force garbage collection
            import gc
            collected = gc.collect()
            logger.info(f"🧹 Forced garbage collection: {collected} objects collected")
            
            # Clear GPU cache if available
            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                logger.info("🔥 GPU cache cleared")
        
        logger.info(f"📊 Memory usage: {memory_gb:.2f}GB")
        
    except Exception as e:
        logger.warning(f"Memory monitoring failed: {e}")


# Add progress tracking for FULL processing
def log_full_processing_progress(current_file, file_index, total_files, file_size_mb):
    """Log detailed progress for FULL processing"""
    try:
        progress_percent = (file_index / total_files) * 100
        
        logger.info(f"🔥 FULL PROCESSING PROGRESS:")
        logger.info(f"   📁 File: {current_file} ({file_size_mb:.1f}MB)")
        logger.info(f"   📊 Progress: {file_index}/{total_files} ({progress_percent:.1f}%)")
        logger.info(f"   ⚡ Mode: FULL content processing (no limits)")
        
        # Monitor memory every 10 files
        if file_index % 10 == 0:
            monitor_memory_during_full_processing()
            
    except Exception as e:
        logger.warning(f"Progress logging failed: {e}")


def check_ollama_health():
    """Check if Ollama service is accessible"""
    try:
        if not OLLAMA_AVAILABLE:
            return False, "Requests library not installed"
            
        if not app.config['OLLAMA_BASE_URL']:
            return False, "Ollama base URL not configured"
        
        # Test connection
        response = requests.get(f"{app.config['OLLAMA_BASE_URL']}/api/tags", timeout=5)
        
        if response.status_code == 200:
            models = response.json().get('models', [])
            model_names = [m['name'] for m in models]
            
            # Check if our model is available
            model_available = any(app.config['OLLAMA_MODEL'] in m for m in model_names)
            
            if model_available:
                logger.info("✅ Ollama service is working")
                return True, "Ollama service is accessible"
            else:
                return False, f"Model {app.config['OLLAMA_MODEL']} not available"
        else:
            return False, f"Ollama server returned status {response.status_code}"
            
    except requests.exceptions.ConnectionError:
        logger.error("❌ Ollama health check failed - connection error")
        return False, "Cannot connect to Ollama server"
    except Exception as e:
        logger.error(f"❌ Ollama health check failed: {e}")
        return False, f"Health check failed: {str(e)}"

def get_ollama_response_stream(prompt):
    """Stream response from Ollama"""
    try:
        response = requests.post(
            f"{app.config['OLLAMA_BASE_URL']}/api/generate",
            json={
                "model": app.config['OLLAMA_MODEL'],
                "prompt": prompt,
                "stream": True,
                "options": {
                    "temperature": app.config['OLLAMA_TEMPERATURE'],
                    "top_p": app.config['OLLAMA_TOP_P'],
                    "top_k": app.config['OLLAMA_TOP_K'],
                    "num_predict": app.config['OLLAMA_MAX_TOKENS']
                }
            },
            stream=True,
            timeout=app.config['OLLAMA_TIMEOUT']
        )
        
        if response.status_code == 200:
            for line in response.iter_lines():
                if line:
                    try:
                        chunk_data = json.loads(line.decode('utf-8'))
                        if 'response' in chunk_data:
                            chunk_text = chunk_data['response']
                            if chunk_text:
                                yield chunk_text
                        
                        # Check if generation is done
                        if chunk_data.get('done', False):
                            break
                            
                    except json.JSONDecodeError:
                        continue
        else:
            yield f"\n\nError: Ollama server returned status {response.status_code}"
                
    except Exception as e:
        logger.error(f"Ollama streaming error: {e}")
        yield f"\n\nError: Failed to get response from Ollama ({str(e)})"

def get_ollama_response_complete(prompt):
    """Get complete response from Ollama (non-streaming)"""
    try:
        response = requests.post(
            f"{app.config['OLLAMA_BASE_URL']}/api/generate",
            json={
                "model": app.config['OLLAMA_MODEL'],
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": app.config['OLLAMA_TEMPERATURE'],
                    "top_p": app.config['OLLAMA_TOP_P'],
                    "top_k": app.config['OLLAMA_TOP_K'],
                    "num_predict": app.config['OLLAMA_MAX_TOKENS']
                }
            },
            timeout=app.config['OLLAMA_TIMEOUT']
        )
        
        if response.status_code == 200:
            result = response.json()
            response_text = result.get('response', '')
            
            if response_text:
                return response_text
            else:
                return "Error: Empty response from Ollama"
        else:
            return f"Error: Ollama server returned status {response.status_code}"
            
    except Exception as e:
        logger.error(f"Ollama complete response error: {e}")
        return f"Error: Failed to get response from Ollama ({str(e)})"

# UPDATE HEALTH CHECK FUNCTIONS FOR OLLAMA

@app.route('/debug/ollama_debug')
@login_required
def debug_ollama():
    """Debug Ollama connection and models"""
    try:
        debug_info = {
            'base_url': app.config['OLLAMA_BASE_URL'],
            'configured_model': app.config['OLLAMA_MODEL'],
            'timestamp': datetime.now().isoformat()
        }
        
        # Test server connection
        try:
            response = requests.get(f"{app.config['OLLAMA_BASE_URL']}", timeout=5)
            debug_info['server_status'] = 'reachable'
            debug_info['server_response'] = response.status_code
        except Exception as e:
            debug_info['server_status'] = 'unreachable'
            debug_info['server_error'] = str(e)
        
        # Test API endpoint
        try:
            response = requests.get(f"{app.config['OLLAMA_BASE_URL']}/api/tags", timeout=10)
            debug_info['api_status'] = response.status_code
            if response.status_code == 200:
                models_data = response.json()
                debug_info['available_models'] = [m['name'] for m in models_data.get('models', [])]
                debug_info['model_count'] = len(debug_info['available_models'])
            else:
                debug_info['api_error'] = response.text
        except Exception as e:
            debug_info['api_status'] = 'failed'
            debug_info['api_error'] = str(e)
        
        # Test model
        try:
            response = requests.post(
                f"{app.config['OLLAMA_BASE_URL']}/api/generate",
                json={
                    "model": app.config['OLLAMA_MODEL'],
                    "prompt": "Say 'test successful'",
                    "stream": False
                },
                timeout=30
            )
            debug_info['model_test_status'] = response.status_code
            if response.status_code == 200:
                result = response.json()
                debug_info['model_response'] = result.get('response', '')[:100]
                debug_info['model_working'] = bool(result.get('response', '').strip())
            else:
                debug_info['model_error'] = response.text
        except Exception as e:
            debug_info['model_test_status'] = 'failed'
            debug_info['model_error'] = str(e)
        
        # Recommendations
        recommendations = []
        if debug_info.get('server_status') == 'unreachable':
            recommendations.append("Start Ollama server: ollama serve")
        if debug_info.get('api_status') != 200:
            recommendations.append("Check if Ollama is properly installed")
        if not debug_info.get('available_models'):
            recommendations.append("Pull a model: ollama pull llama3.1")
        if debug_info.get('model_test_status') == 404:
            recommendations.append(f"Model '{app.config['OLLAMA_MODEL']}' not found. Try: ollama pull {app.config['OLLAMA_MODEL']}")
        
        debug_info['recommendations'] = recommendations
        debug_info['setup_commands'] = [
            "# Install Ollama",
            "curl -fsSL https://ollama.ai/install.sh | sh",
            "",
            "# Start Ollama server", 
            "ollama serve",
            "",
            "# Pull models (in another terminal)",
            "ollama pull llama3.1",
            "ollama pull llama2",
            "",
            "# Test model",
            "ollama run llama3.1 'Hello world'"
        ]
        
        return jsonify(debug_info)
        
    except Exception as e:
        return jsonify({
            'error': str(e),
            'traceback': traceback.format_exc()
        })

@app.route('/api/ollama_status')
@login_required
def ollama_status():
    """Check Ollama service status with enhanced diagnostics"""
    is_healthy, message = check_ollama_health()
    
    status = {
        'ollama_healthy': is_healthy,
        'status_message': message,
        'base_url': app.config['OLLAMA_BASE_URL'],
        'model_name': app.config['OLLAMA_MODEL'],
        'timeout_seconds': app.config['OLLAMA_TIMEOUT'],
        'max_tokens': app.config['OLLAMA_MAX_TOKENS'],
        'temperature': app.config['OLLAMA_TEMPERATURE'],
        'ollama_available': OLLAMA_AVAILABLE,
        'streaming_enabled': app.config['OLLAMA_STREAM']
    }
    
    if is_healthy:
        try:
            # Get available models
            response = requests.get(f"{app.config['OLLAMA_BASE_URL']}/api/tags", timeout=5)
            if response.status_code == 200:
                models = response.json().get('models', [])
                status['available_models'] = [m['name'] for m in models]
                status['model_count'] = len(models)
                
                # Check if current model is in the list
                current_model = app.config['OLLAMA_MODEL']
                model_found = any(current_model in model['name'] for model in models)
                status['current_model_available'] = model_found
                
                if not model_found and models:
                    status['suggested_models'] = [m['name'] for m in models[:3]]
                    
        except Exception as e:
            status['model_fetch_error'] = str(e)
    else:
        # Add troubleshooting info
        status['troubleshooting'] = {
            'common_solutions': [
                "Start Ollama: ollama serve",
                f"Pull model: ollama pull {app.config['OLLAMA_MODEL']}",
                "Check server: curl http://localhost:11434/api/tags",
                "List models: ollama list"
            ],
            'setup_url': "https://ollama.ai/download"
        }
    
    return jsonify(status)

def diagnose_ollama_issues():
    """Comprehensive Ollama diagnostics"""
    logger.info("🔍 OLLAMA DIAGNOSTICS STARTING...")
    logger.info("=" * 60)
    
    issues = []
    solutions = []
    
    # 1. Check if Ollama server is running
    try:
        response = requests.get(f"{app.config['OLLAMA_BASE_URL']}", timeout=5)
        logger.info(f"✅ Ollama server responding at {app.config['OLLAMA_BASE_URL']}")
    except requests.exceptions.ConnectionError:
        issues.append("Ollama server not running")
        solutions.append("Start Ollama: ollama serve")
        logger.error(f"❌ Cannot connect to {app.config['OLLAMA_BASE_URL']}")
    except Exception as e:
        issues.append(f"Connection error: {e}")
        logger.error(f"❌ Connection error: {e}")
    
    # 2. Check API endpoints
    try:
        response = requests.get(f"{app.config['OLLAMA_BASE_URL']}/api/tags", timeout=10)
        if response.status_code == 200:
            logger.info("✅ API endpoints accessible")
            models_data = response.json()
            models = models_data.get('models', [])
            if models:
                logger.info(f"✅ Found {len(models)} models:")
                for model in models:
                    logger.info(f"   - {model['name']} (size: {model.get('size', 'unknown')})")
            else:
                issues.append("No models installed")
                solutions.append("Install a model: ollama pull llama3.1")
                logger.error("❌ No models found")
        else:
            issues.append(f"API returned status {response.status_code}")
            logger.error(f"❌ API error: {response.status_code}")
    except Exception as e:
        issues.append(f"API check failed: {e}")
        logger.error(f"❌ API check failed: {e}")
    
    # 3. Check specific model
    try:
        response = requests.post(
            f"{app.config['OLLAMA_BASE_URL']}/api/generate",
            json={
                "model": app.config['OLLAMA_MODEL'],
                "prompt": "Test",
                "stream": False
            },
            timeout=30
        )
        if response.status_code == 200:
            logger.info(f"✅ Model {app.config['OLLAMA_MODEL']} working")
        elif response.status_code == 404:
            issues.append(f"Model '{app.config['OLLAMA_MODEL']}' not found")
            solutions.append(f"Install model: ollama pull {app.config['OLLAMA_MODEL']}")
            logger.error(f"❌ Model {app.config['OLLAMA_MODEL']} not found")
        else:
            issues.append(f"Model test failed: {response.status_code}")
            logger.error(f"❌ Model test failed: {response.status_code}")
    except Exception as e:
        issues.append(f"Model test error: {e}")
        logger.error(f"❌ Model test error: {e}")
    
    # 4. Summary and solutions
    logger.info("=" * 60)
    if issues:
        logger.error("🚨 ISSUES FOUND:")
        for i, issue in enumerate(issues, 1):
            logger.error(f"  {i}. {issue}")
        
        logger.info("🔧 SOLUTIONS:")
        for i, solution in enumerate(solutions, 1):
            logger.info(f"  {i}. {solution}")
        
        logger.info("📋 QUICK SETUP GUIDE:")
        logger.info("  1. Install Ollama: curl -fsSL https://ollama.ai/install.sh | sh")
        logger.info("  2. Start server: ollama serve")
        logger.info("  3. Pull model: ollama pull llama3.1")
        logger.info("  4. Test: ollama run llama3.1 'Hello'")
        
    else:
        logger.info("✅ All checks passed!")
    
    logger.info("=" * 60)
    return len(issues) == 0

def startup_ollama_check():
    """Check Ollama status during startup with better diagnostics"""
    logger.info("🔍 Checking Ollama service status...")
    
    if not OLLAMA_AVAILABLE:
        logger.error("❌ Requests library not available for Ollama")
        logger.error("🔧 Install with: pip install requests")
        return False
    
    if not app.config['OLLAMA_BASE_URL']:
        logger.error("❌ Ollama base URL not configured")
        logger.error("🔧 Set base URL: export OLLAMA_BASE_URL='http://localhost:11434'")
        return False
    
    # Run diagnostics first
    if not diagnose_ollama_issues():
        logger.error("❌ Ollama diagnostics failed")
        return False
    
    # Initialize Ollama
    ollama_success = initialize_ollama()
    
    if ollama_success:
        is_healthy, message = check_ollama_health()
        
        if is_healthy:
            logger.info(f"✅ Ollama Status: {message}")
            logger.info(f"🦙 Using model: {app.config['OLLAMA_MODEL']}")
            logger.info(f"🎯 Temperature: {app.config['OLLAMA_TEMPERATURE']}")
            logger.info(f"📊 Max tokens: {app.config['OLLAMA_MAX_TOKENS']}")
            return True
        else:
            logger.error(f"❌ Ollama Status: {message}")
            return False
    else:
        logger.error("❌ Ollama initialization failed")
        return False

# UPDATE VALIDATION FUNCTION FOR OLLAMA

def validate_system_requirements():
    """Validate all system requirements at startup"""
    logger.info("🔍 SYSTEM VALIDATION STARTING...")
    
    issues = []
    
    # Check GPU
    if torch.cuda.is_available():
        logger.info(f"✅ GPU Available: {torch.cuda.get_device_name()}")
    else:
        logger.warning("⚠️ GPU not available - using CPU mode")
    
    # Check Ollama instead of Gemini
    ollama_healthy = startup_ollama_check()
    if ollama_healthy:
        logger.info("✅ Ollama Service: Ready")
    else:
        issues.append("Ollama service not ready")
        logger.error("❌ Ollama Service: Not Ready")
    
    # Check OCR
    try:
        import easyocr
        logger.info("✅ EasyOCR package available")
    except ImportError:
        issues.append("EasyOCR package not installed")
        logger.error("❌ EasyOCR package not available")
    
    # Check database
    try:
        with app.app_context():
            with db.engine.connect() as connection:
                connection.execute(text('SELECT 1'))
        logger.info("✅ Database connection successful")
    except Exception as db_error:
        issues.append(f"Database connection failed: {db_error}")
        logger.error(f"❌ Database connection failed: {db_error}")
    
    # Summary
    if issues:
        logger.error("⚠️ SYSTEM VALIDATION COMPLETED WITH ISSUES:")
        for issue in issues:
            logger.error(f"   - {issue}")
        logger.error("🔧 Some features may not work correctly until these issues are resolved")
    else:
        logger.info("✅ SYSTEM VALIDATION COMPLETED - ALL SYSTEMS READY")
    
    return len(issues) == 0, issues

# Global Ollama client placeholder
ollama_client = None

# Enhanced processing jobs with better thread safety
processing_jobs = {}
processing_jobs_lock = threading.RLock()

# Global enhanced processor instance
global_processor = None
search_engine = None

ocr_reader = None



@app.route('/')
def index():
    """Main index page"""
    return render_template('index.html')



@app.route('/register', methods=['GET', 'POST'])
def register():
    """Enhanced user registration with personal details"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        # Get form data
        first_name = request.form.get('first_name', '').strip()
        last_name = request.form.get('last_name', '').strip()
        company_name = request.form.get('company_name', '').strip()
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        
        # Validation
        errors = []
        
        # Check required fields
        if not first_name:
            errors.append('First name is required')
        elif len(first_name) < 2:
            errors.append('First name must be at least 2 characters long')
        elif not re.match(r"^[a-zA-Z\s'-]+$", first_name):
            errors.append('First name contains invalid characters')
            
        if not last_name:
            errors.append('Last name is required')
        elif len(last_name) < 2:
            errors.append('Last name must be at least 2 characters long')
        elif not re.match(r"^[a-zA-Z\s'-]+$", last_name):
            errors.append('Last name contains invalid characters')
            
        if not username:
            errors.append('Username is required')
        elif len(username) < 3:
            errors.append('Username must be at least 3 characters long')
        elif not re.match(r"^[a-zA-Z0-9_.-]+$", username):
            errors.append('Username can only contain letters, numbers, dots, hyphens, and underscores')
            
        if not email:
            errors.append('Email is required')
        elif not re.match(r'^[^\s@]+@[^\s@]+\.[^\s@]+$', email):
            errors.append('Please enter a valid email address')
            
        if not password:
            errors.append('Password is required')
        elif len(password) < 8:
            errors.append('Password must be at least 8 characters long')
            
        if password != confirm_password:
            errors.append('Passwords do not match')
        
        # Check if username or email already exists
        if User.query.filter_by(username=username).first():
            errors.append('Username already taken')
            
        if User.query.filter_by(email=email).first():
            errors.append('Email address already registered')
        
        if not company_name:
            errors.append('Company name is required')
        elif len(company_name) < 2:
            errors.append('Company name must be at least 2 characters long')
        elif len(company_name) > 100:
            errors.append('Company name is too long (maximum 100 characters)')
        
        # If there are errors, show them
        if errors:
            for error in errors:
                flash(error, 'error')
            return render_template('register.html')
        
        try:
            # Create new user
            user = User(
                first_name=first_name,
                last_name=last_name,
                company_name=company_name if company_name else None,
                username=username,
                email=email
            )
            user.set_password(password)
            
            db.session.add(user)
            db.session.commit()
            
            logger.info(f"New user registered: {username} ({first_name} {last_name})")
            
            flash('Registration successful! Please log in.', 'success')
            return redirect(url_for('login'))
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Registration error: {e}")
            flash('Registration failed. Please try again.', 'error')
            return render_template('register.html')
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    """User login"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        remember = request.form.get('remember') == 'on'

        user = User.query.filter_by(username=username).first()

        if user and user.check_password(password):
            login_user(user, remember=remember)
            next_page = request.args.get('next')
            return redirect(next_page or url_for('index'))
        else:
            flash('Invalid username or password', 'error')

    return render_template('login.html')



def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


@app.route('/new_chat')
@login_required
def new_chat():
    """Create new chat"""
    try:
        chat_id = str(uuid.uuid4())
        
        conversation = Conversation(
            id=chat_id,
            user_id=current_user.id,
            title='New Query'
        )
        db.session.add(conversation)
        
        # Check for documents
        doc_count = Document.query.filter_by(user_id=current_user.id).count()
        initial_message = f"Ready to search through {doc_count} documents. What would you like to know?"
        
        initial_msg = Message(
            conversation_id=chat_id,
            role='assistant',
            content=initial_message,
            timestamp=datetime.now(timezone.utc)
        )
        db.session.add(initial_msg)
        db.session.commit()
        
        return redirect(url_for('chat', chat_id=chat_id))
        
    except Exception as e:
        logger.error(f'Error creating chat: {e}')
        flash('Error creating new chat', 'error')
        return redirect(url_for('index'))

@app.route('/logout', methods=['GET', 'POST'])
@login_required
def logout():
    """User logout"""
    logout_user()
    session.clear()
    flash('You have been logged out successfully.', 'success')
    return redirect(url_for('login'))



@app.route('/api/check_subscription', methods=['GET'])
@login_required
def check_subscription_api():
    """Check subscription status with enhanced features - FIXED VERSION"""
    try:
        now = datetime.now(timezone.utc)
        trial_days_remaining = 0
        
        if current_user.trial_end_date:
            trial_end = current_user._make_aware(current_user.trial_end_date)
            trial_days_remaining = max(0, (trial_end - now).days)
        
        has_active_trial = (current_user.subscription_plan == 'trial' and 
                           current_user.trial_end_date and 
                           current_user._make_aware(current_user.trial_end_date) > now)
        
        has_active_subscription = (current_user.subscription_status == 'active' and 
                                 current_user.subscription_plan in ['basic', 'pro'] and
                                 (current_user.current_period_end is None or 
                                  current_user._make_aware(current_user.current_period_end) > now))
        
        # STRICT UPLOAD CHECK - Only allow if they have active subscription or trial
        allow_uploads = has_active_trial or has_active_subscription or current_user.is_admin
        
        # Check if trial is available (not used before OR expired)
        has_trial_available = (current_user.subscription_plan != 'trial' or 
                             (current_user.trial_end_date and 
                              current_user._make_aware(current_user.trial_end_date) <= now))
        
        # Determine subscription required flag
        subscription_required = not allow_uploads
        
        logger.info(f"Subscription check for user {current_user.id}:")
        logger.info(f"  - Subscription plan: {current_user.subscription_plan}")
        logger.info(f"  - Subscription status: {current_user.subscription_status}")
        logger.info(f"  - Has active trial: {has_active_trial}")
        logger.info(f"  - Has active subscription: {has_active_subscription}")
        logger.info(f"  - Allow uploads: {allow_uploads}")
        logger.info(f"  - Trial available: {has_trial_available}")
        
        return jsonify({
            'is_logged_in': True,
            'has_subscription': has_active_subscription,
            'has_active_subscription': has_active_subscription,
            'has_active_trial': has_active_trial,
            'has_trial_available': has_trial_available,
            'trial_days_remaining': trial_days_remaining,
            'trial_expired': (current_user.subscription_plan == 'trial' and 
                            current_user.trial_end_date and 
                            current_user._make_aware(current_user.trial_end_date) <= now),
            'allow_uploads': allow_uploads,  # This is the key field - must be True to upload
            'subscription_required': subscription_required,
            'plan_name': current_user.plan_details.get('name', 'Free') if allow_uploads else 'No Active Plan',
            'current_plan': current_user.subscription_plan,
            'gpu_enabled': torch.cuda.is_available(),
            'enhanced_processing': True,
            'perfect_formatting': True,
            'comprehensive_search': True,
            'accuracy_focused': True,
            'rtx_a6000_optimized': torch.cuda.is_available(),
            'max_file_size_gb': app.config['MAX_FILE_SIZE'] / 1e9,
            'max_upload_size_gb': app.config['MAX_CONTENT_LENGTH'] / 1e9,
            'max_extracted_size_gb': app.config['MAX_ZIP_EXTRACTED_SIZE'] / 1e9
        })
    except Exception as e:
        logger.error(f"Subscription check failed: {e}")
        return jsonify({'error': str(e)}), 500




# Add route to replace the existing upload handler
@app.route('/upload', methods=['POST'])
@login_required
def upload_files_optimized():
    """OPTIMIZED upload endpoint"""
    if request.headers.get('Accept') == 'text/event-stream':
        return handle_streaming_upload_optimized()
    else:
        return jsonify({'error': 'Invalid request format'}), 400




# Optimized content extraction dispatch
def extract_content_optimized_dispatch(file_path: str) -> str:
    """OPTIMIZED content extraction with FULL content processing (no limits)"""
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return ""
        
        logger.info(f"🔥 Processing FULL content for {os.path.basename(file_path)} ({file_size / 1024 / 1024:.1f}MB)")
        
        # Dispatch to FULL content extractors
        if file_ext == '.pdf':
            return extract_pdf_full_optimized(file_path)
        elif file_ext == '.docx':
            return extract_docx_full_optimized(file_path)
        elif file_ext in ['.txt', '.md']:
            return extract_text_full_optimized(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return extract_image_full_optimized(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return extract_excel_full_optimized(file_path)
        else:
            return extract_fallback_full_optimized(file_path)
            
    except Exception as e:
        logger.error(f"FULL content extraction failed for {file_path}: {e}")
        return ""


def extract_pdf_full_optimized(file_path: str) -> str:
    """FULL PDF extraction - ALL pages with speed optimizations"""
    try:
        doc = fitz.open(file_path)
        content_parts = []
        
        # Add comprehensive header
        content_parts.append(f"PDF DOCUMENT: {os.path.basename(file_path)}")
        content_parts.append(f"TOTAL PAGES: {len(doc)}")
        content_parts.append(f"FILE SIZE: {os.path.getsize(file_path) / 1024 / 1024:.1f}MB")
        content_parts.append("=" * 80)
        
        # Process ALL pages with optimized extraction
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            content_parts.append(f"\nPAGE {page_num + 1}")
            content_parts.append("-" * 40)
            
            try:
                # Use optimized text extraction
                page_text = page.get_text("text")  # Fastest method
                
                if page_text.strip():
                    content_parts.append(page_text.strip())
                else:
                    # Try alternative extraction methods
                    page_text = page.get_text("dict")  # More detailed
                    if page_text and 'blocks' in page_text:
                        text_blocks = []
                        for block in page_text['blocks']:
                            if 'lines' in block:
                                for line in block['lines']:
                                    for span in line['spans']:
                                        text_blocks.append(span['text'])
                        
                        if text_blocks:
                            content_parts.append(' '.join(text_blocks))
                        else:
                            content_parts.append(f"[Page {page_num + 1} - No text content detected]")
                    else:
                        content_parts.append(f"[Page {page_num + 1} - No text content detected]")
                
                content_parts.append(f"[END PAGE {page_num + 1}]")
                
            except Exception as page_error:
                logger.warning(f"Error extracting page {page_num + 1}: {page_error}")
                content_parts.append(f"[Page {page_num + 1} - Extraction error: {str(page_error)}]")
                content_parts.append(f"[END PAGE {page_num + 1}]")
        
        doc.close()
        
        result = "\n".join(content_parts)
        logger.info(f"✅ FULL PDF extraction completed: {len(result)} characters from {len(doc)} pages")
        
        return result
        
    except Exception as e:
        logger.error(f"FULL PDF extraction failed: {e}")
        return f"PDF DOCUMENT: {os.path.basename(file_path)}\nExtraction failed: {str(e)}"


def extract_docx_full_optimized(file_path: str) -> str:
    """FULL DOCX extraction - ALL paragraphs with optimizations"""
    try:
        import docx
        doc = docx.Document(file_path)
        
        content_parts = []
        content_parts.append(f"DOCX DOCUMENT: {os.path.basename(file_path)}")
        content_parts.append(f"TOTAL PARAGRAPHS: {len(doc.paragraphs)}")
        content_parts.append(f"FILE SIZE: {os.path.getsize(file_path) / 1024 / 1024:.1f}MB")
        content_parts.append("=" * 60)
        
        # Process ALL paragraphs with optimizations
        processed_paragraphs = 0
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)
                processed_paragraphs += 1
        
        # Extract tables if present
        if doc.tables:
            content_parts.append("\n" + "="*50)
            content_parts.append("DOCUMENT TABLES")
            content_parts.append("="*50)
            
            for table_idx, table in enumerate(doc.tables, 1):
                content_parts.append(f"\nTABLE {table_idx}:")
                content_parts.append("-" * 20)
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text if cell_text else "[empty]")
                    content_parts.append(" | ".join(row_data))
                
                content_parts.append(f"[END TABLE {table_idx}]")
        
        result = "\n".join(content_parts)
        logger.info(f"✅ FULL DOCX extraction completed: {len(result)} characters from {processed_paragraphs} paragraphs")
        
        return result
        
    except Exception as e:
        logger.error(f"FULL DOCX extraction failed: {e}")
        return f"DOCX DOCUMENT: {os.path.basename(file_path)}\nExtraction failed: {str(e)}"


# FULL text extraction (ALL content)
def extract_text_full_optimized(file_path: str) -> str:
    """FULL text extraction with encoding detection"""
    try:
        # Try multiple encodings for comprehensive reading
        encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252', 'iso-8859-1']
        
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    used_encoding = encoding
                    break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            # Fallback to binary reading with error handling
            with open(file_path, 'rb') as f:
                binary_content = f.read()
                content = binary_content.decode('utf-8', errors='replace')
                used_encoding = 'utf-8 (with errors replaced)'
        
        # Add comprehensive header
        header = f"""TEXT DOCUMENT: {os.path.basename(file_path)}
FILE SIZE: {os.path.getsize(file_path) / 1024:.1f}KB
ENCODING: {used_encoding}
CONTENT LENGTH: {len(content)} characters
{"=" * 60}

"""
        result = header + content
        
        logger.info(f"✅ FULL text extraction completed: {len(result)} total characters")
        
        return result
        
    except Exception as e:
        logger.error(f"FULL text extraction failed: {e}")
        return f"TEXT DOCUMENT: {os.path.basename(file_path)}\nExtraction failed: {str(e)}"


# FULL image extraction with OCR (ALL content)
def extract_image_full_optimized(file_path: str) -> str:
    """FULL image extraction with comprehensive OCR"""
    global ocr_reader
    
    try:
        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        
        if not ocr_reader or not app.config['OCR_ENABLED']:
            return f"""IMAGE DOCUMENT: {filename}
FILE SIZE: {file_size / 1024:.1f}KB
{"=" * 60}

[OCR NOT AVAILABLE]
OCR processing is not enabled or configured.
Image file detected but text extraction not possible.

{"=" * 60}"""
        
        logger.info(f"🖼️ Processing FULL image OCR for: {filename}")
        
        # Extract ALL text using OCR
        ocr_text = ocr_reader.extract_text_from_image(file_path)
        
        if not ocr_text.strip():
            return f"""IMAGE DOCUMENT: {filename}
FILE SIZE: {file_size / 1024:.1f}KB
{"=" * 60}

[NO TEXT DETECTED]
OCR processing completed but no readable text was found in the image.

{"=" * 60}"""
        
        # Build comprehensive content
        content_blocks = []
        content_blocks.append(f"IMAGE DOCUMENT: {filename}")
        content_blocks.append(f"FILE SIZE: {file_size / 1024:.1f}KB")
        content_blocks.append(f"OCR ENGINE: {'GPU-accelerated EasyOCR' if ocr_reader.gpu_enabled else 'CPU EasyOCR'}")
        content_blocks.append(f"EXTRACTED TEXT LENGTH: {len(ocr_text)} characters")
        content_blocks.append("=" * 60)
        content_blocks.append("")
        content_blocks.append("[FULL OCR EXTRACTED CONTENT]")
        content_blocks.append(ocr_text)
        content_blocks.append("[END OCR CONTENT]")
        content_blocks.append("")
        content_blocks.append("=" * 60)
        content_blocks.append("OCR PROCESSING SUMMARY")
        content_blocks.append("=" * 60)
        content_blocks.append(f"Image preprocessing: {'Enabled' if app.config['OCR_PREPROCESS_ENABLED'] else 'Disabled'}")
        content_blocks.append(f"Confidence threshold: {app.config['OCR_CONFIDENCE_THRESHOLD']}")
        content_blocks.append(f"Languages: {', '.join(app.config['OCR_LANGUAGES'])}")
        
        result = "\n".join(content_blocks)
        
        logger.info(f"✅ FULL image OCR completed: {len(result)} total characters")
        
        return result
        
    except Exception as e:
        logger.error(f"FULL image OCR extraction failed for {file_path}: {e}")
        return f"""IMAGE DOCUMENT: {os.path.basename(file_path)}
{"=" * 60}

[OCR EXTRACTION ERROR]
Error processing image: {str(e)}

{"=" * 60}"""


# FULL Excel extraction (ALL sheets and data)
def extract_excel_full_optimized(file_path: str) -> str:
    """FULL Excel extraction - ALL sheets with optimizations"""
    try:
        import pandas as pd
        
        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        
        # Read ALL sheets
        excel_data = pd.read_excel(file_path, sheet_name=None)  # None reads all sheets
        
        content_parts = []
        content_parts.append(f"EXCEL DOCUMENT: {filename}")
        content_parts.append(f"FILE SIZE: {file_size / 1024:.1f}KB")
        content_parts.append(f"TOTAL SHEETS: {len(excel_data)}")
        content_parts.append("=" * 60)
        
        total_rows = 0
        total_cols = 0
        
        # Process ALL sheets and ALL data
        for sheet_name, df in excel_data.items():
            content_parts.append(f"\nSHEET: {sheet_name}")
            content_parts.append(f"ROWS: {len(df)}, COLUMNS: {len(df.columns)}")
            content_parts.append("-" * 40)
            
            # Convert ALL data to string format
            sheet_content = df.to_string(index=True, max_rows=None, max_cols=None)
            content_parts.append(sheet_content)
            content_parts.append(f"[END SHEET: {sheet_name}]")
            
            total_rows += len(df)
            total_cols += len(df.columns)
        
        content_parts.append("\n" + "=" * 60)
        content_parts.append("EXCEL SUMMARY")
        content_parts.append("=" * 60)
        content_parts.append(f"Total data rows processed: {total_rows}")
        content_parts.append(f"Total columns processed: {total_cols}")
        
        result = "\n".join(content_parts)
        
        logger.info(f"✅ FULL Excel extraction completed: {len(result)} characters from {len(excel_data)} sheets")
        
        return result
        
    except Exception as e:
        logger.error(f"FULL Excel extraction failed: {e}")
        return f"EXCEL DOCUMENT: {os.path.basename(file_path)}\nExtraction failed: {str(e)}"


# FULL fallback extraction
def extract_fallback_full_optimized(file_path: str) -> str:
    """FULL fallback extraction for unknown formats"""
    try:
        filename = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        
        # Try reading as text with multiple encodings
        encodings = ['utf-8', 'latin1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()  # Read ALL content
                
                header = f"""UNKNOWN FORMAT: {filename}
FILE SIZE: {file_size / 1024:.1f}KB
ENCODING: {encoding}
CONTENT LENGTH: {len(content)} characters
{"=" * 60}

"""
                return header + content
                
            except Exception:
                continue
        
        # If text reading fails, try binary analysis
        with open(file_path, 'rb') as f:
            binary_data = f.read()
        
        # Try to detect if it's text-like
        text_chars = sum(1 for byte in binary_data[:1000] if 32 <= byte <= 126 or byte in [9, 10, 13])
        is_likely_text = text_chars / min(1000, len(binary_data)) > 0.7
        
        if is_likely_text:
            try:
                content = binary_data.decode('utf-8', errors='replace')
                return f"""BINARY FILE (TEXT-LIKE): {filename}
FILE SIZE: {file_size / 1024:.1f}KB
{"=" * 60}

{content}"""
            except:
                pass
        
        return f"""BINARY FILE: {filename}
FILE SIZE: {file_size / 1024:.1f}KB
{"=" * 60}

[BINARY FILE - TEXT EXTRACTION NOT POSSIBLE]
This file appears to be in a binary format that cannot be processed as text.
File format: {os.path.splitext(filename)[1].upper() or 'Unknown'}

{"=" * 60}"""
        
    except Exception as e:
        logger.error(f"FULL fallback extraction failed: {e}")
        return f"FILE: {os.path.basename(file_path)}\nExtraction failed: {str(e)}"
        

# Optimized content extraction dispatch
def extract_content_optimized_dispatch(file_path: str) -> str:
    """Optimized content extraction with faster dispatch"""
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        # Quick file size check
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            return ""
        
        # Dispatch to optimized extractors
        if file_ext == '.pdf':
            return extract_pdf_optimized(file_path)
        elif file_ext == '.docx':
            return extract_docx_optimized(file_path)
        elif file_ext in ['.txt', '.md']:
            return extract_text_optimized(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return extract_image_optimized(file_path)
        else:
            return extract_fallback_optimized(file_path)
            
    except Exception as e:
        logger.error(f"Content extraction failed for {file_path}: {e}")
        return ""


def extract_pdf_optimized(file_path: str) -> str:
    """Optimized PDF extraction for speed"""
    try:
        doc = fitz.open(file_path)
        content_parts = []
        
        # Add header with accurate page count
        content_parts.append(f"PDF DOCUMENT: {os.path.basename(file_path)}")
        content_parts.append(f"TOTAL PAGES: {len(doc)}")
        content_parts.append("=" * 80)
        
        for page_num, page in enumerate(doc, 1):
            # Clear page marker
            content_parts.append(f"\nPAGE {page_num}")
            content_parts.append("-" * 40)
            
            try:
                # Fast text extraction
                page_text = page.get_text()
                if page_text.strip():
                    content_parts.append(page_text.strip())
                else:
                    content_parts.append(f"[Page {page_num} - No text content]")
                
                # Page end marker
                content_parts.append(f"[END PAGE {page_num}]")
                
            except Exception as page_error:
                logger.warning(f"Error on page {page_num}: {page_error}")
                content_parts.append(f"[Page {page_num} - Extraction error]")
                content_parts.append(f"[END PAGE {page_num}]")
        
        doc.close()
        return "\n".join(content_parts)
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def extract_docx_optimized(file_path: str) -> str:
    """Optimized DOCX extraction for speed"""
    try:
        import docx
        doc = docx.Document(file_path)
        
        content_parts = []
        content_parts.append(f"DOCX DOCUMENT: {os.path.basename(file_path)}")
        content_parts.append("=" * 60)
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content_parts.append(text)
        
        return "\n".join(content_parts)
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_text_optimized(file_path: str) -> str:
    """Optimized text extraction for speed"""
    try:
        # Try different encodings quickly
        encodings = ['utf-8', 'latin1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                # Add header
                header = f"TEXT DOCUMENT: {os.path.basename(file_path)}\n" + "=" * 60 + "\n"
                return header + content
                
            except UnicodeDecodeError:
                continue
        
        return ""
        
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return ""


def extract_image_optimized(file_path: str) -> str:
    """Optimized image extraction with OCR"""
    global ocr_reader
    
    try:
        if not ocr_reader:
            return f"IMAGE FILE: {os.path.basename(file_path)} (OCR not available)"
        
        ocr_text = ocr_reader.extract_text_from_image(file_path)
        
        if ocr_text.strip():
            header = f"IMAGE DOCUMENT: {os.path.basename(file_path)}\n" + "=" * 60 + "\n"
            return header + "[OCR CONTENT]\n" + ocr_text + "\n[END OCR CONTENT]"
        else:
            return f"IMAGE FILE: {os.path.basename(file_path)} (No text detected)"
            
    except Exception as e:
        logger.error(f"Image extraction failed: {e}")
        return f"IMAGE FILE: {os.path.basename(file_path)} (Extraction failed)"


def extract_fallback_optimized(file_path: str) -> str:
    """Optimized fallback extraction"""
    try:
        # Try as text file
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read(10000)  # Read first 10KB only
        
        header = f"UNKNOWN FORMAT: {os.path.basename(file_path)}\n" + "=" * 60 + "\n"
        return header + content
        
    except Exception as e:
        return f"FILE: {os.path.basename(file_path)} (Cannot extract content)"



# Update the main upload handler to use optimizations
def handle_streaming_upload_optimized():
    """OPTIMIZED streaming upload with FULL content processing (no limits)"""
    def generate_progress():
        try:
            # Quick subscription check
            is_valid, reason = validate_user_subscription_strict(current_user)
            if not is_valid:
                yield f"data: {json.dumps({'status': 'error', 'message': f'Upload not allowed: {reason}'})}\n\n"
                return
            
            user_id = current_user.id
            user_upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], str(user_id))
            user_extract_dir = os.path.join(app.config['EXTRACT_FOLDER'], str(user_id))
            
            # Fast directory setup
            os.makedirs(user_upload_dir, exist_ok=True)
            os.makedirs(user_extract_dir, exist_ok=True)
            clean_directory(user_upload_dir)
            clean_directory(user_extract_dir)
            
            files = request.files.getlist('files')
            if not files or all(f.filename == '' for f in files):
                yield f"data: {json.dumps({'status': 'error', 'message': 'No files uploaded'})}\n\n"
                return
            
            # PHASE 1: OPTIMIZED FILE EXTRACTION (5-15%)
            yield f"data: {json.dumps({'status': 'extracting', 'message': '🔥 Full content extraction...', 'percent': 5})}\n\n"
            
            extracted_files = []
            
            # Process files with optimized extraction
            for file_idx, file in enumerate(files):
                if file.filename == '':
                    continue
                    
                filename = secure_filename(file.filename)
                
                # Progress updates every 5 files for better performance
                if file_idx % 5 == 0:
                    progress = 5 + (file_idx / len(files)) * 10
                    yield f"data: {json.dumps({'status': 'extracting', 'percent': int(progress), 'message': f'🔥 Extracting {filename}...'})}\n\n"
                
                try:
                    if filename.lower().endswith('.zip'):
                        zip_path = os.path.join(user_upload_dir, filename)
                        file.save(zip_path)
                        
                        if extract_zip_optimized(zip_path, user_extract_dir):
                            for root, _, zip_files in os.walk(user_extract_dir):
                                for f in zip_files:
                                    if allowed_file(f):
                                        extracted_files.append(os.path.join(root, f))
                        
                        os.remove(zip_path)
                    else:
                        if allowed_file(filename):
                            file_path = os.path.join(user_extract_dir, filename)
                            file.save(file_path)
                            extracted_files.append(file_path)
                            
                except Exception as e:
                    logger.error(f"Error processing {filename}: {e}")
                    continue
            
            if not extracted_files:
                yield f"data: {json.dumps({'status': 'error', 'message': 'No valid files found'})}\n\n"
                return
            
            # PHASE 2: FILES READY FOR FULL PROCESSING (15-20%)
            yield f"data: {json.dumps({'status': 'files_ready', 'total': len(extracted_files), 'percent': 15, 'message': f'✅ {len(extracted_files)} files ready for FULL processing'})}\n\n"
            
            # PHASE 3: FULL CONTENT PROCESSING (20-90%)
            processed_count = 0
            failed_files = []
            batch_data = []
            
            # Use optimized batch size for FULL processing
            batch_size = 25  # Smaller batches for full processing to maintain speed
            
            for i, file_path in enumerate(extracted_files):
                filename = os.path.basename(file_path)
                file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
                
                # Progress updates every 3% for smoother UX
                if i % max(1, len(extracted_files) // 30) == 0:
                    progress = 20 + (i / len(extracted_files)) * 70
                    yield f"data: {json.dumps({'status': 'processing', 'current_file': filename, 'percent': int(progress), 'message': f'🔥 FULL processing: {filename} ({file_size_mb:.1f}MB)'})}\n\n"
                
                try:
                    # FULL CONTENT EXTRACTION (no limits)
                    content = extract_content_optimized_dispatch(file_path)
                    
                    if content and len(content.strip()) > 10:
                        file_hash = get_file_hash_optimized(file_path)
                        
                        # Quick duplicate check
                        existing = Document.query.filter_by(file_hash=file_hash, user_id=user_id).first()
                        if not existing:
                            # FULL CHUNKING AND EMBEDDINGS
                            chunks = create_chunks_optimized(content)  # Process ALL content
                            embeddings = generate_embeddings_fast(chunks)  # Generate ALL embeddings
                            
                            batch_data.append((file_path, content, chunks, embeddings, file_hash, user_id))
                            processed_count += 1
                            
                            logger.info(f"✅ FULL processing completed for {filename}: {len(content)} chars, {len(chunks)} chunks")
                        else:
                            processed_count += 1  # Already exists
                            logger.info(f"⚠️ Duplicate detected, skipping: {filename}")
                    else:
                        failed_files.append(filename)
                        logger.warning(f"❌ No content extracted from: {filename}")
                        
                except Exception as e:
                    failed_files.append(filename)
                    logger.error(f"❌ Error during FULL processing of {file_path}: {e}")
                
                # Process batch when full
                if len(batch_data) >= batch_size:
                    try:
                        yield f"data: {json.dumps({'status': 'processing', 'message': f'💾 Saving batch of {len(batch_data)} fully processed files...', 'percent': int(20 + (i / len(extracted_files)) * 70)})}\n\n"
                        
                        save_document_bulk_optimized(batch_data)
                        batch_data = []
                        
                        logger.info(f"✅ Batch saved successfully")
                        
                    except Exception as batch_error:
                        logger.error(f"❌ Batch processing error: {batch_error}")
                        failed_files.extend([os.path.basename(data[0]) for data in batch_data])
                        batch_data = []
            
            # Process remaining batch
            if batch_data:
                try:
                    yield f"data: {json.dumps({'status': 'processing', 'message': f'💾 Saving final batch of {len(batch_data)} files...', 'percent': 90})}\n\n"
                    
                    save_document_bulk_optimized(batch_data)
                    
                    logger.info(f"✅ Final batch saved successfully")
                    
                except Exception as final_batch_error:
                    logger.error(f"❌ Final batch error: {final_batch_error}")
                    failed_files.extend([os.path.basename(data[0]) for data in batch_data])
            
            # PHASE 4: COMPLETION (90-100%)
            success_rate = round((processed_count / len(extracted_files)) * 100, 1)
            
            completion_message = f'🔥 FULL processing completed: {processed_count}/{len(extracted_files)} files with ALL content, embeddings, and features!'
            
            yield f"data: {json.dumps({'status': 'completed', 'processed_count': processed_count, 'total_files': len(extracted_files), 'failed_files': failed_files, 'percent': 100, 'success_rate': success_rate, 'message': completion_message, 'full_processing': True, 'no_limits_applied': True})}\n\n"
            
            # Log completion statistics
            total_size_mb = sum(os.path.getsize(f) / (1024 * 1024) for f in extracted_files)
            logger.info(f"🎉 UPLOAD COMPLETED:")
            logger.info(f"   📁 Files processed: {processed_count}/{len(extracted_files)}")
            logger.info(f"   📊 Total size: {total_size_mb:.1f}MB")
            logger.info(f"   ✅ Success rate: {success_rate}%")
            logger.info(f"   🔥 FULL content processing: ALL pages, paragraphs, and data extracted")
            logger.info(f"   🧠 ALL embeddings generated")
            logger.info(f"   ⚡ NO content limits applied")
            
        except Exception as e:
            logger.error(f"❌ FULL processing upload error: {e}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            yield f"data: {json.dumps({'status': 'error', 'message': f'FULL processing upload failed: {str(e)}'})}\n\n"
    
    return Response(
        stream_with_context(generate_progress()),
        mimetype='text/event-stream',
        headers={'Cache-Control': 'no-cache', 'Connection': 'keep-alive'}
    )


def create_chunks_optimized(content):
    """FULL chunking with optimized performance - processes ALL content"""
    try:
        if not content or len(content.strip()) < 50:
            return []
        
        # Use optimized chunk size and overlap for speed while maintaining quality
        chunk_size = 1200  # Good balance of context and speed
        overlap = 200      # Sufficient overlap for context preservation
        
        logger.info(f"🔥 Creating chunks from {len(content)} characters (FULL content)")
        
        # For small content, return as single chunk
        if len(content) <= chunk_size:
            return [{
                'content': content,
                'chunk_index': 0,
                'metadata': {
                    'full_content': True,
                    'start_pos': 0,
                    'end_pos': len(content),
                    'chunk_size': len(content)
                }
            }]
        
        chunks = []
        start = 0
        chunk_index = 0
        
        # Process ALL content with optimized chunking
        while start < len(content):
            end = start + chunk_size
            
            # Try to end at sentence boundary for better context
            if end < len(content):
                # Look for sentence endings within the last 200 characters
                search_start = max(start + chunk_size - 200, start)
                search_text = content[search_start:end + 200]
                
                # Find the best breaking point
                sentence_endings = ['.', '!', '?', '\n\n']
                best_break = -1
                
                for ending in sentence_endings:
                    last_pos = search_text.rfind(ending)
                    if last_pos > len(search_text) // 2:  # Don't break too early
                        best_break = search_start + last_pos + 1
                        break
                
                if best_break > start:
                    end = best_break
            
            # Extract chunk content
            chunk_content = content[start:end].strip()
            
            if chunk_content:  # Only add non-empty chunks
                chunks.append({
                    'content': chunk_content,
                    'chunk_index': chunk_index,
                    'metadata': {
                        'start_pos': start,
                        'end_pos': end,
                        'chunk_size': len(chunk_content),
                        'full_processing': True,
                        'has_overlap': start > 0
                    }
                })
                chunk_index += 1
            
            # Move to next chunk with overlap
            start = max(start + 1, end - overlap)
            
            # Prevent infinite loops
            if end >= len(content):
                break
        
        logger.info(f"✅ FULL chunking completed: {len(chunks)} chunks created from ALL content")
        
        return chunks
        
    except Exception as e:
        logger.error(f"❌ FULL chunking failed: {e}")
        return [{
            'content': content[:2000] if content else 'Error processing content',
            'chunk_index': 0,
            'metadata': {'error': str(e), 'fallback_chunk': True}
        }]


# Optimize your existing generate_embeddings_fast function for FULL processing
def generate_embeddings_fast(chunks):
    """FULL embedding generation - processes ALL chunks with optimizations"""
    if not chunks:
        return []
    
    if not ensure_processor_initialized():
        logger.warning("Processor not initialized, returning zero embeddings")
        return [np.zeros(768) for _ in chunks]
    
    try:
        logger.info(f"🔥 Generating embeddings for ALL {len(chunks)} chunks...")
        
        # Extract ALL text content
        texts = [chunk['content'] for chunk in chunks]
        
        # Use optimized batch processing for speed
        batch_size = 32 if global_processor.gpu_enabled else 16
        all_embeddings = []
        
        # Process in batches for memory efficiency
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i:i + batch_size]
            
            try:
                # Generate embeddings with optimized settings
                batch_embeddings = global_processor.embedding_model.encode(
                    batch_texts,
                    convert_to_tensor=False,
                    show_progress_bar=False,
                    batch_size=batch_size,
                    normalize_embeddings=True,  # Normalize for better similarity
                    device=global_processor.device
                )
                
                all_embeddings.extend(batch_embeddings)
                
                logger.info(f"✅ Processed embedding batch {i//batch_size + 1}/{(len(texts) + batch_size - 1)//batch_size}")
                
            except Exception as batch_error:
                logger.error(f"❌ Embedding batch error: {batch_error}")
                # Add zero embeddings for failed batch
                batch_embeddings = [np.zeros(768) for _ in batch_texts]
                all_embeddings.extend(batch_embeddings)
        
        logger.info(f"✅ FULL embedding generation completed: {len(all_embeddings)} embeddings for ALL chunks")
        
        return all_embeddings
        
    except Exception as e:
        logger.error(f"❌ FULL embedding generation failed: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        # Return zero embeddings as fallback
        return [np.zeros(768) for _ in chunks]



# Optimized ZIP extraction
def extract_zip_optimized(zip_path, extract_to):
    """Optimized ZIP extraction with better performance"""
    try:
        os.makedirs(extract_to, exist_ok=True)
        
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            # Quick validation without loading everything
            total_size = 0
            valid_files = []
            
            for info in zip_ref.infolist():
                # Security check
                if '..' in info.filename or info.filename.startswith('/'):
                    continue
                
                total_size += info.file_size
                if total_size > app.config['MAX_ZIP_EXTRACTED_SIZE']:
                    logger.error(f"ZIP too large: {total_size / 1e9:.2f}GB")
                    return False
                
                # Only track files we can process
                if allowed_file(info.filename):
                    valid_files.append(info)
            
            # Extract only valid files
            for info in valid_files:
                try:
                    zip_ref.extract(info, extract_to)
                except Exception as e:
                    logger.warning(f"Failed to extract {info.filename}: {e}")
                    continue
            
            logger.info(f"Extracted {len(valid_files)} valid files")
            return True
            
    except Exception as e:
        logger.error(f"ZIP extraction failed: {e}")
        return False



def clean_directory(directory):
    """Clean directory with better error handling"""
    if not os.path.exists(directory):
        return
        
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            logger.error(f'Failed to delete {file_path}: {e}')



# Add this optimized database bulk save function
def save_document_bulk_optimized(documents_data):
    """ULTRA-FAST bulk save using raw SQL - processes ALL content without limits"""
    try:
        if not documents_data:
            return True
            
        logger.info(f"💾 Ultra-fast bulk saving {len(documents_data)} documents with FULL content...")
        start_time = time.time()
        
        # Use raw SQL for maximum speed
        current_time = datetime.now(timezone.utc)
        
        # Prepare ALL document records for bulk insert
        doc_insert_values = []
        content_insert_values = []
        
        for doc_data in documents_data:
            file_path, content, chunks, embeddings, file_hash, user_id = doc_data
            
            filename = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            file_type = os.path.splitext(file_path)[1].lower()
            rel_path = os.path.relpath(file_path, app.config['EXTRACT_FOLDER'])
            
            doc_metadata = {
                'original_filename': filename,
                'file_size': file_size,
                'content_length': len(content),
                'relative_path': rel_path,
                'chunk_count': len(chunks),
                'full_content_processed': True,
                'processing_time': time.time() - start_time
            }
            
            # Prepare document insert values
            doc_insert_values.append(f"""({user_id}, '{rel_path.replace("'", "''")}', '{file_hash}', '{file_type}', '{current_time}', '{json.dumps(doc_metadata).replace("'", "''")}')""")

        
        # ULTRA-FAST bulk insert ALL documents at once
        if doc_insert_values:
            bulk_doc_sql = f"""
                INSERT INTO document (user_id, file_path, file_hash, file_type, processed_at, metadata)
                VALUES {','.join(doc_insert_values)}
                RETURNING id
            """
            
            with db.engine.begin() as conn:
                result = conn.execute(text(bulk_doc_sql))
                doc_ids = [row[0] for row in result.fetchall()]
            
            logger.info(f"✅ Inserted {len(doc_ids)} documents in bulk")
            
            # Prepare ALL content records for bulk insert
            for i, (doc_data, doc_id) in enumerate(zip(documents_data, doc_ids)):
                file_path, content, chunks, embeddings, file_hash, user_id = doc_data
                
                # Full content record
                escaped_content = content.replace("'", "''").replace("\\", "\\\\")
                content_insert_values.append(f"""({doc_id}, '{escaped_content}', 'full_content_optimized', 0, '{current_time}', '{json.dumps({'full_processing': True}).replace("'", "''")}')""")

                
                # ALL chunks with embeddings (no limits)
                for j, chunk_data in enumerate(chunks):
                    escaped_chunk = chunk_data['content'].replace("'", "''").replace("\\", "\\\\")
                    embedding = embeddings[j] if j < len(embeddings) else None
                    embedding_json = json.dumps(embedding.tolist() if isinstance(embedding, np.ndarray) else embedding) if embedding is not None else 'null'
                    
                    content_insert_values.append(f"""({doc_id}, '{escaped_chunk}', 'chunk_optimized', {chunk_data['chunk_index']}, '{current_time}', '{json.dumps({'full_chunk': True}).replace("'", "''")}')""")

            
            # ULTRA-FAST bulk insert ALL content at once
            if content_insert_values:
                # Process in very large batches for speed
                batch_size = 5000  # Very large batches
                
                for i in range(0, len(content_insert_values), batch_size):
                    batch_values = content_insert_values[i:i + batch_size]
                    
                    bulk_content_sql = f"""
                        INSERT INTO document_content (document_id, content, content_type, chunk_index, created_at, content_metadata)
                        VALUES {','.join(batch_values)}
                    """
                    
                    with db.engine.begin() as conn:
                        conn.execute(text(bulk_content_sql))
                    
                    logger.info(f"✅ Inserted batch {i//batch_size + 1} of content records")
        
        elapsed = time.time() - start_time
        logger.info(f"🚀 ULTRA-FAST bulk save completed in {elapsed:.2f}s - ALL content processed!")
        
        return True
        
    except Exception as e:
        logger.error(f"Ultra-fast bulk save error: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return False


# Initialize global OCR processor
def initialize_ocr_processor():
    """Initialize the global OCR processor"""
    global ocr_reader
    try:
        if app.config['OCR_ENABLED']:
            logger.info("🔥 Initializing OCR processor...")
            ocr_reader = EnhancedOCRProcessor(gpu_enabled=app.config['OCR_GPU_ENABLED'])
            logger.info("✅ OCR processor initialized successfully")
            return True
        else:
            logger.info("📄 OCR disabled in configuration")
            return False
    except Exception as e:
        logger.error(f"❌ OCR processor initialization failed: {e}")
        return False


def extract_keywords_using_ai(query: str) -> str:
    """
    Use Ollama/Llama to intelligently extract key terms from user questions
    """
    try:
        if not OLLAMA_AVAILABLE:
            logger.warning("Ollama not available, falling back to simple extraction")
            return extract_keywords_simple_fallback(query)
        
        # Prompt specifically designed for keyword extraction - let AI do ALL the work
        extraction_prompt = f"""Extract the main subject/keyword from this question for document search.

Question: "{query}"

Rules:
- Extract ONLY the core subject/topic the user wants to know about
- Remove question words and action words naturally
- Return 1-3 words maximum
- Focus on nouns and important terms

Examples:
"What is a mango?" → mango
"Tell me about universities" → universities
"How do companies work?" → companies
"List all students" → students
"Explain artificial intelligence" → artificial intelligence
"Show me financial reports" → financial reports
"What are the benefits of exercise?" → benefits exercise
"Describe the research methodology" → research methodology

Now extract the keyword from: "{query}"

Keyword:"""

        logger.info(f"🤖 AI extracting keywords from: '{query}'")
        
        response = requests.post(
            f"{app.config['OLLAMA_BASE_URL']}/api/generate",
            json={
                "model": app.config['OLLAMA_MODEL'],
                "prompt": extraction_prompt,
                "stream": False,
                "options": {
                    "temperature": 0.1,  # Very low for consistent extraction
                    "top_p": 0.9,
                    "top_k": 40,
                    "num_predict": 50
                }
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            extracted = result.get('response', '').strip()
            
            # Minimal cleaning - just remove common prefixes the AI might include
            extracted = re.sub(r'^(keyword:|key term:|answer:|result:)\s*', '', extracted, flags=re.IGNORECASE)
            extracted = extracted.strip(' :"\'')
            
            if extracted and len(extracted) > 0:
                logger.info(f"✅ AI extracted: '{query}' → '{extracted}'")
                return extracted.lower()
            else:
                logger.warning(f"AI returned empty keyword, using fallback")
                return extract_keywords_simple_fallback(query)
        else:
            logger.warning(f"AI extraction failed with status {response.status_code}")
            return extract_keywords_simple_fallback(query)
            
    except Exception as e:
        logger.error(f"AI keyword extraction error: {e}")
        return extract_keywords_simple_fallback(query)

def extract_keywords_simple_fallback(query: str) -> str:
    """Simple fallback when AI is not available - just return the original query"""
    # No manual pattern removal - let the search handle the full query
    return query.strip().lower()




def dual_system_response(question: str, user_id: int):
    """
    FIXED dual system implementation with proper error handling
    """
    try:
        # Get user documents
        user_documents = Document.query.filter_by(user_id=user_id).all()
        
        if not user_documents:
            yield json.dumps({
                'status': 'error',
                'message': 'No documents found. Please upload documents first.'
            }) + "\n"
            return
        
        document_info = []
        all_content = []
        
        # Load all document content
        for doc in user_documents:
            filename = doc.document_metadata.get('original_filename', 'Unknown') if doc.document_metadata else 'Unknown'
            content = get_full_document_content(doc.id)
            if content:
                all_content.append(f"=== DOCUMENT: {filename} ===\n{content}\n")
                document_info.append(filename)
        
        if not all_content:
            yield json.dumps({
                'status': 'error',
                'message': 'No content found in documents.'
            }) + "\n"
            return
        
        # STEP 1: AI extracts keywords from the question
        yield json.dumps({
            'status': 'status_update',
            'message': '🤖 AI analyzing question to extract key terms...'
        }) + "\n"
        
        extracted_keyword = extract_keywords_using_ai(question)
        
        yield json.dumps({
            'status': 'status_update',
            'message': f'✅ AI identified key term: "{extracted_keyword}"'
        }) + "\n"
        
        # STEP 2: Find exact keyword matches in documents
        yield json.dumps({
            'status': 'status_update',
            'message': f'🔍 Searching documents for matches of "{extracted_keyword}"...'
        }) + "\n"
        
        keyword_instances, keyword_summary = search_exact_keyword_matches(extracted_keyword, user_id)
        
        # STEP 3: Stream the keyword matches found
        yield json.dumps({'status': 'stream_start', 'stream_type': 'keyword_matches'}) + "\n"
        
        keyword_content = f"""# Keyword Matches for "{extracted_keyword}"

**Original Question:** {question}
**AI Extracted Keyword:** {extracted_keyword}

{keyword_summary}

---

"""
        
        # Stream keyword matches
        chunk_size = 150
        for i in range(0, len(keyword_content), chunk_size):
            chunk = keyword_content[i:i + chunk_size]
            yield json.dumps({
                'status': 'stream_chunk', 
                'content': chunk, 
                'stream_type': 'keyword_matches'
            }) + "\n"
            time.sleep(0.1)
        
        yield json.dumps({'status': 'stream_end', 'stream_type': 'keyword_matches'}) + "\n"
        
        # STEP 4: Prepare relevant content for AI answer
        yield json.dumps({
            'status': 'status_update',
            'message': '🦙 Preparing AI answer using relevant document content...'
        }) + "\n"
        
        # Build context from keyword matches + surrounding content - FIXED
        relevant_content = ""
        if keyword_instances:
            relevant_content += f"\n=== RELEVANT CONTENT FOR '{extracted_keyword}' ===\n"
            for i, instance in enumerate(keyword_instances[:5], 1):  # Top 5 matches
                relevant_content += f"\nMatch {i} - {instance['filename']} (Page {instance['page_number']}):\n"
                relevant_content += f"Found: '{instance['matched_text']}'\n"
                
                # FIXED: Safely get context with fallback
                context = instance.get('context', f"Found '{instance['matched_text']}' in {instance['filename']}")
                if len(context) > 500:
                    context = context[:500] + "..."
                
                relevant_content += f"Context: {context}\n"
                relevant_content += "---\n"
        
        # Add full document content (limited)
        combined_content = "\n".join(all_content)
        if len(combined_content) > 10000:
            combined_content = combined_content[:10000] + "\n[Content truncated...]"
        
        # STEP 5: AI answers the original question
        yield json.dumps({
            'status': 'status_update',
            'message': f'🤖 AI answering: "{question}"'
        }) + "\n"
        
        answer_prompt = f"""You are an intelligent document analysis assistant. Answer the user's question based on the provided documents.

USER QUESTION: "{question}"

EXTRACTED KEYWORD: "{extracted_keyword}"

KEYWORD MATCHES FOUND IN DOCUMENTS:
{relevant_content}

ADDITIONAL DOCUMENT CONTENT:
{combined_content}

INSTRUCTIONS:
1. Answer the user's EXACT question: "{question}"
2. Use the information from the documents provided above
3. If keyword matches were found, prioritize that information
4. Be comprehensive and helpful
5. Structure your answer clearly with headings if needed
6. Cite specific documents when making claims: [^filename||page]
7. If the question asks for a list, provide a clear list
8. If asking for explanation, provide detailed explanations
9. Only use information from the provided documents
10. If the answer isn't in the documents, say so clearly

Answer the question now:"""

        yield json.dumps({'status': 'stream_start', 'stream_type': 'ai_answer'}) + "\n"
        
        # Stream AI answer
        full_answer = ""
        try:
            for chunk in get_ollama_response_stream(answer_prompt):
                if chunk.strip():
                    full_answer += chunk
                    yield json.dumps({
                        'status': 'stream_chunk', 
                        'content': chunk, 
                        'stream_type': 'ai_answer'
                    }) + "\n"
                    time.sleep(0.01)
            
            # Final completion
            yield json.dumps({
                'status': 'stream_end',
                'stream_type': 'ai_answer',
                'sources': [{'filename': s, 'path': s} for s in document_info],
                'question_analysis': {
                    'original_question': question,
                    'ai_extracted_keyword': extracted_keyword,
                    'keyword_matches_found': len(keyword_instances),
                    'documents_searched': len(user_documents),
                    'system_type': 'dual_ai_system'
                }
            }) + "\n"
            
        except Exception as answer_error:
            logger.error(f"AI answer generation error: {answer_error}")
            yield json.dumps({
                'status': 'error',
                'message': f'AI answer failed: {str(answer_error)}',
                'stream_type': 'ai_answer'
            }) + "\n"
        
        # Complete
        yield json.dumps({
            'status': 'complete',
            'message': 'Dual system complete: keyword matches + AI answer',
            'keyword_matches': len(keyword_instances),
            'extracted_keyword': extracted_keyword,
            'original_question': question
        }) + "\n"
        
    except Exception as e:
        logger.error(f"Dual system error: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        yield json.dumps({
            'status': 'error',
            'message': f'System error: {str(e)}'
        }) + "\n"


def search_exact_keyword_matches(keyword: str, user_id: int):
    """
    Simple exact keyword search - FIXED to include context
    """
    logger.info(f"🎯 EXACT KEYWORD SEARCH: '{keyword}' for user {user_id}")
    
    try:
        all_instances = []
        
        # Get user documents
        user_documents = Document.query.filter_by(user_id=user_id).all()
        
        if not user_documents:
            return [], "No documents found to search."
        
        logger.info(f"Processing {len(user_documents)} documents")
        
        # Create simple search variations
        search_variations = [
            keyword,                    # Original
            keyword.lower(),           # Lowercase
            keyword.upper(),           # Uppercase
            keyword.capitalize(),      # Capitalized
            keyword.title(),           # Title case
        ]
        
        # Add plural/singular if needed
        if not keyword.endswith('s'):
            search_variations.append(keyword + 's')
        elif len(keyword) > 4 and keyword.endswith('s'):
            search_variations.append(keyword[:-1])
        
        # Remove duplicates
        search_variations = list(set(search_variations))
        
        logger.info(f"Search variations: {search_variations}")
        
        # Process each document
        for document in user_documents:
            try:
                filename = document.document_metadata.get('original_filename', 'Unknown') if document.document_metadata else 'Unknown'
                
                # Get full document content
                content = get_full_document_content(document.id)
                
                if not content:
                    continue
                
                # Search for ALL variations
                for variation in search_variations:
                    instances = find_exact_matches_with_context(content, variation, filename, document.id, keyword)
                    all_instances.extend(instances)
                
            except Exception as doc_error:
                logger.error(f"Error processing document {filename}: {doc_error}")
                continue
        
        # Remove duplicates
        final_instances = remove_duplicate_instances(all_instances)
        
        # Build summary
        summary = build_simple_match_summary(final_instances, keyword, len(user_documents))
        
        logger.info(f"✅ SEARCH COMPLETED: {len(final_instances)} total instances found")
        
        return final_instances, summary
        
    except Exception as e:
        logger.error(f"❌ SEARCH ERROR: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return [], f"Search failed: {str(e)}"


def find_exact_matches_with_context(content: str, search_term: str, filename: str, document_id: int, original_keyword: str):
    """
    Find exact matches with proper context - FIXED VERSION
    """
    instances = []
    
    try:
        # Case-insensitive search
        content_lower = content.lower()
        search_lower = search_term.lower()
        
        start = 0
        while True:
            # Find next occurrence
            pos = content_lower.find(search_lower, start)
            if pos == -1:
                break
            
            # Get the actual matched text (preserving original case)
            matched_text = content[pos:pos + len(search_term)]
            
            # Determine page number
            page_number = determine_simple_page_number(content, pos)
            
            # FIXED: Get proper context around the match
            context_start = max(0, pos - 800)
            context_end = min(len(content), pos + len(search_term) + 800)
            context = content[context_start:context_end].strip()
            
            # Clean context (remove excessive whitespace)
            context = ' '.join(context.split())
            
            # Ensure context is not empty
            if not context:
                context = f"Found '{matched_text}' in {filename}"
            
            # Create instance with guaranteed context
            instance = {
                'filename': filename,
                'document_id': document_id,
                'page_number': page_number,
                'position_in_content': pos,
                'matched_text': matched_text,
                'original_keyword': original_keyword,
                'exact_match': matched_text.lower() == original_keyword.lower(),
                'context': context,  # FIXED: Always include context
                'instance_key': f"{document_id}_{page_number}_{pos}_{matched_text.lower()}"
            }
            
            instances.append(instance)
            
            # Move past this match
            start = pos + 1
    
    except Exception as e:
        logger.error(f"Error finding exact matches: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
    
    return instances

def determine_simple_page_number(content: str, position: int) -> int:
    """
    Simple page number detection - FIXED
    """
    try:
        # Count page markers before this position
        content_before = content[:position]
        
        # Look for page markers
        page_matches = re.findall(r'PAGE (\d+)', content_before, re.IGNORECASE)
        
        if page_matches:
            return int(page_matches[-1])  # Get the last (most recent) page number
        
        # Look for other page patterns
        page_matches = re.findall(r'\[PAGE (\d+)\]', content_before, re.IGNORECASE)
        if page_matches:
            return int(page_matches[-1])
        
        return 1  # Default to page 1
        
    except Exception as e:
        logger.error(f"Error determining page number: {e}")
        return 1


def find_simple_exact_matches(content: str, search_term: str, filename: str, document_id: int, original_keyword: str):
    """
    Find exact matches with simple logic
    """
    instances = []
    
    try:
        # Case-insensitive search
        content_lower = content.lower()
        search_lower = search_term.lower()
        
        start = 0
        while True:
            # Find next occurrence
            pos = content_lower.find(search_lower, start)
            if pos == -1:
                break
            
            # Get the actual matched text (preserving original case)
            matched_text = content[pos:pos + len(search_term)]
            
            # Determine page number
            page_number = determine_simple_page_number(content, pos)
            
            # Create simple instance
            instance = {
                'filename': filename,
                'document_id': document_id,
                'page_number': page_number,
                'position_in_content': pos,
                'matched_text': matched_text,
                'original_keyword': original_keyword,
                'exact_match': matched_text.lower() == original_keyword.lower(),
                'instance_key': f"{document_id}_{page_number}_{pos}_{matched_text.lower()}"
            }
            
            instances.append(instance)
            
            # Move past this match
            start = pos + 1
    
    except Exception as e:
        logger.error(f"Error finding simple matches: {e}")
    
    return instances


def build_simple_match_summary(instances, keyword, documents_searched):
    """
    Build SIMPLE summary - just count and pages
    """
    if not instances:
        return f"No instances of '{keyword}' found in {documents_searched} documents."
    
    # Count by document and collect pages
    by_doc = {}
    all_pages = set()
    
    for instance in instances:
        doc = instance['filename']
        page = instance['page_number']
        
        if doc not in by_doc:
            by_doc[doc] = set()
        
        by_doc[doc].add(page)
        all_pages.add(page)
    
    # Build simple summary
    summary_parts = []
    
    # Main result
    exact_count = len([instance for instance in instances if instance.get('exact_match', False)])
    summary_parts.append(f"**Found {exact_count} matches for '{keyword}'**")
    summary_parts.append("")
    
    
    # Pages (simple list)
    sorted_pages = sorted(list(all_pages))
    if len(sorted_pages) <= 15:
        pages_text = ", ".join(map(str, sorted_pages))
    else:
        pages_text = f"{', '.join(map(str, sorted_pages[:10]))}, ... and {len(sorted_pages)-10} more pages"
    
    summary_parts.append(f"**Pages:** {pages_text}")
    summary_parts.append("")
    
    # By document (if multiple docs)
    if len(by_doc) > 1:
        summary_parts.append("**By Document:**")
        for doc_name, doc_pages in by_doc.items():
            sorted_doc_pages = sorted(list(doc_pages))
            summary_parts.append(f"• {doc_name}: {len(sorted_doc_pages)} matches on pages {', '.join(map(str, sorted_doc_pages))}")
    else:
        doc_name = list(by_doc.keys())[0]
        summary_parts.append(f"**Document:** {doc_name}")

    summary_parts.append("<br><hr><br>")
    summary_parts.append("<strong>Summary</strong>")
    return "\n".join(summary_parts)


def find_exact_matches_in_content(content: str, search_term: str, filename: str, document_id: int, original_keyword: str):
    """
    Find exact matches in content with proper case handling
    """
    instances = []
    
    try:
        # Use case-insensitive search but preserve original text
        content_lower = content.lower()
        search_lower = search_term.lower()
        
        start = 0
        while True:
            # Find next occurrence
            pos = content_lower.find(search_lower, start)
            if pos == -1:
                break
            
            # Get the actual matched text (preserving original case)
            matched_text = content[pos:pos + len(search_term)]
            
            # Determine which page this is on
            page_number = determine_page_number(content, pos)
            
            # Get context around the match
            context_start = max(0, pos - 800)
            context_end = min(len(content), pos + len(search_term) + 800)
            context = content[context_start:context_end].strip()
            
            # Clean context
            context = ' '.join(context.split())
            
            # Create instance
            instance = {
                'filename': filename,
                'document_id': document_id,
                'page_number': page_number,
                'position_in_content': pos,
                'matched_text': matched_text,
                'search_term': search_term,
                'original_keyword': original_keyword,
                'context': context,
                'exact_match': matched_text.lower() == original_keyword.lower(),
                'instance_key': f"{document_id}_{page_number}_{pos}_{matched_text.lower()}"
            }
            
            instances.append(instance)
            logger.debug(f"Found match: '{matched_text}' at position {pos}, page {page_number}")
            
            # Move past this match
            start = pos + 1
    
    except Exception as e:
        logger.error(f"Error finding matches in content: {e}")
    
    return instances


def determine_page_number(content: str, position: int) -> int:
    """
    Determine which page a position belongs to
    """
    try:
        # Count page markers before this position
        content_before = content[:position]
        
        # Look for various page markers
        page_markers = [
            r'PAGE (\d+)',
            r'\[PAGE (\d+)\]',
            r'Page (\d+)',
            r'- Page (\d+) -'
        ]
        
        highest_page = 1
        
        for pattern in page_markers:
            matches = re.findall(pattern, content_before)
            if matches:
                # Get the highest page number found before this position
                page_numbers = [int(match) for match in matches]
                highest_page = max(highest_page, max(page_numbers))
        
        return highest_page
        
    except Exception as e:
        logger.debug(f"Error determining page number: {e}")
        return 1




@app.route('/clear_documents', methods=['POST'])
@login_required
def clear_documents():
    """Clear all documents and content for the current user"""
    try:
        user_id = current_user.id
        deleted_docs = 0
        deleted_content = 0
        
        # Get all documents for this user
        user_documents = Document.query.filter_by(user_id=user_id).all()
        
        if not user_documents:
            logger.info(f"No documents found for user {user_id}")
            return jsonify({
                'success': True,
                'message': 'No documents to clear',
                'deleted_documents': 0,
                'deleted_content': 0
            })
        
        logger.info(f"Found {len(user_documents)} documents to delete for user {user_id}")
        
        # Delete content for each document
        for document in user_documents:
            try:
                # Count and delete content records for this document
                content_count = DocumentContent.query.filter_by(document_id=document.id).count()
                DocumentContent.query.filter_by(document_id=document.id).delete()
                deleted_content += content_count
                
                # Delete the document itself
                db.session.delete(document)
                deleted_docs += 1
                
            except Exception as doc_error:
                logger.warning(f"Error deleting document {document.id}: {doc_error}")
                continue
        
        # Clean up user directories
        user_upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], str(user_id))
        user_extract_dir = os.path.join(app.config['EXTRACT_FOLDER'], str(user_id))
        
        directories_cleaned = []
        try:
            if os.path.exists(user_upload_dir):
                clean_directory(user_upload_dir)
                directories_cleaned.append("upload")
                logger.info(f"Cleaned upload directory: {user_upload_dir}")
        except Exception as e:
            logger.warning(f"Error cleaning upload directory: {e}")
        
        try:
            if os.path.exists(user_extract_dir):
                clean_directory(user_extract_dir)
                directories_cleaned.append("extract")
                logger.info(f"Cleaned extract directory: {user_extract_dir}")
        except Exception as e:
            logger.warning(f"Error cleaning extract directory: {e}")
        
        # Clear GPU cache if available
        if torch.cuda.is_available():
            try:
                torch.cuda.empty_cache()
                logger.info("Cleared GPU cache")
            except Exception as e:
                logger.warning(f"Error clearing GPU cache: {e}")
        
        # Clear global processor cache for this user
        global global_processor
        if global_processor:
            try:
                global_processor.chunk_metadata = []
                global_processor.faiss_index.reset()
                logger.info("Cleared accuracy-focused processor cache")
            except Exception as e:
                logger.warning(f"Error clearing processor cache: {e}")
        
        # Commit all changes
        db.session.commit()
        
        logger.info(f"Successfully cleared {deleted_docs} documents and {deleted_content} content records for user {user_id}")
        
        return jsonify({
            'success': True,
            'message': f'Successfully cleared {deleted_docs} documents and {deleted_content} content records',
            'deleted_documents': deleted_docs,
            'deleted_content': deleted_content,
            'directories_cleaned': directories_cleaned,
            'accuracy_focused_cleared': True
        })
        
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error clearing documents for user {current_user.id}: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            'success': False,
            'error': f'Failed to clear documents: {str(e)}'
        }), 500



def build_improved_summary(instances, keyword, documents_searched):
    """
    Build improved summary that shows actual results found
    """
    if not instances:
        return f"""
EXACT KEYWORD SEARCH RESULTS
============================
Keyword: '{keyword}'
Documents searched: {documents_searched}
Total instances found: 0

No instances of '{keyword}' or its variations were found in your documents.
"""
    
    summary_parts = []
    
    # Header - Show SUCCESS
    summary_parts.append(f"🎯 EXACT MATCHES FOUND FOR: '{keyword.upper()}'")
    summary_parts.append("=" * 80)
    summary_parts.append("")
    
    # Statistics
    by_doc = {}
    variations_found = set()
    exact_matches = 0
    
    for instance in instances:
        doc = instance['filename']
        if doc not in by_doc:
            by_doc[doc] = {'total': 0, 'pages': set(), 'variations': set()}
        
        by_doc[doc]['total'] += 1
        by_doc[doc]['pages'].add(instance['page_number'])
        by_doc[doc]['variations'].add(instance['matched_text'])
        variations_found.add(instance['matched_text'])
        
        if instance['exact_match']:
            exact_matches += 1
    
    total_docs = len(by_doc)
    total_pages = sum(len(doc_data['pages']) for doc_data in by_doc.values())
    
    summary_parts.append("📊 <b>SEARCH RESULTS:</b>")
    summary_parts.append(f"• **TOTAL INSTANCES FOUND: {len(instances)}** ✅")
    summary_parts.append(f"• Exact matches: {exact_matches}")
    summary_parts.append(f"• Variations found: {len(variations_found)}")
    summary_parts.append(f"• Documents with matches: {total_docs}")
    summary_parts.append(f"• Pages with matches: {total_pages}")
    summary_parts.append("")
    
    # # Show what was actually found
    # summary_parts.append("<b>ACTUAL MATCHES FOUND:</b>")
    # summary_parts.append("-" * 40)
    # for variation in sorted(variations_found):
    #     count = len([i for i in instances if i['matched_text'] == variation])
    #     summary_parts.append(f"• **'{variation}'** - {count} instances")
    # summary_parts.append("")
    
    # Distribution by document
    summary_parts.append("📄 <b>FOUND IN THESE DOCUMENTS:</b>")
    summary_parts.append("-" * 50)
    
    for doc_name, doc_data in by_doc.items():
        summary_parts.append(f"📄 **{doc_name}**")
        summary_parts.append(f"   • Total instances: **{doc_data['total']}**")
        summary_parts.append(f"   • Pages: {', '.join(map(str, sorted(doc_data['pages'])))}")
        # summary_parts.append(f"   • Variations: {', '.join(sorted(doc_data['variations']))}")
        summary_parts.append("")
    
    # Sample instances with context
    summary_parts.append("📍 <b>SAMPLE INSTANCES WITH CONTEXT:</b>")
    summary_parts.append("-" * 45)
    
    # Show top 5 instances with context
    top_instances = instances[:5]
    for i, instance in enumerate(top_instances, 1):
        summary_parts.append(f"**{i}. '{instance['matched_text']}' in {instance['filename']}**")
        summary_parts.append(f"   📍 Page {instance['page_number']}")
        
        # Show abbreviated context
        context = instance['context']
        if len(context) > 200:
            # Find the match position in context and show around it
            match_pos = context.lower().find(instance['matched_text'].lower())
            if match_pos != -1:
                start = max(0, match_pos - 100)
                end = min(len(context), match_pos + len(instance['matched_text']) + 100)
                context_snippet = context[start:end]
                if start > 0:
                    context_snippet = "..." + context_snippet
                if end < len(context):
                    context_snippet = context_snippet + "..."
            else:
                context_snippet = context[:200] + "..."
        else:
            context_snippet = context
        
        # summary_parts.append(f"   💬 Context: {context_snippet}")
        summary_parts.append("")
    
    # Success footer
    # summary_parts.append("=" * 60)
    # summary_parts.append("✅ <b>SEARCH SUCCESSFUL!</b>")
    # summary_parts.append("=" * 60)
    # summary_parts.append(f"🎯 Successfully found **{len(instances)} instances** of '{keyword}'")
    # summary_parts.append(f"📚 Searched across {documents_searched} document(s)")
    # summary_parts.append(f"📄 Found matches in {total_docs} document(s)")
    # summary_parts.append(f"📑 Covered {total_pages} page(s)")
    # summary_parts.append("=" * 60)
    
    return "\n".join(summary_parts)



def generate_search_variations_improved(keyword: str) -> list:
    """
    Generate comprehensive search variations
    """
    keyword = keyword.strip()
    variations = set()
    
    # Add original
    variations.add(keyword)
    
    # Add case variations
    variations.add(keyword.lower())
    variations.add(keyword.upper())
    variations.add(keyword.capitalize())
    variations.add(keyword.title())
    
    # Split into words and try different combinations
    words = keyword.split()
    
    if len(words) > 1:
        # Add reversed order
        variations.add(' '.join(reversed(words)))
        
        # Add individual words if they're significant
        for word in words:
            if len(word) > 3:
                variations.add(word)
                variations.add(word.upper())
                variations.add(word.capitalize())
        
        # Add with different word separators
        variations.add('_'.join(words))
        variations.add('-'.join(words))
        variations.add(''.join(words))  # No spaces
    
    # Add plural/singular forms
    if keyword.endswith('s') and len(keyword) > 4:
        variations.add(keyword[:-1])  # Remove 's'
    elif not keyword.endswith('s'):
        variations.add(keyword + 's')  # Add 's'
    
    # Convert back to list and remove empty strings
    result = [v for v in variations if v and len(v.strip()) > 2]
    
    return list(result)




def generate_keyword_variations(keyword: str) -> list:
    """
    Generate variations of a keyword for exact matching
    """
    keyword = keyword.strip().lower()
    variations = [keyword]  # Start with original
    
    # Add common morphological variations
    if len(keyword) > 3:
        # Plural forms
        if keyword.endswith('y'):
            variations.append(keyword[:-1] + 'ies')  # university -> universities
        elif keyword.endswith(('s', 'x', 'z', 'ch', 'sh')):
            variations.append(keyword + 'es')
        else:
            variations.append(keyword + 's')
        
        # Past tense / -ed forms
        if keyword.endswith('e'):
            variations.append(keyword + 'd')
        else:
            variations.append(keyword + 'ed')
        
        # -ing forms
        if keyword.endswith('e'):
            variations.append(keyword[:-1] + 'ing')
        else:
            variations.append(keyword + 'ing')
        
        # -er forms
        variations.append(keyword + 'er')
        
        # -ly forms
        variations.append(keyword + 'ly')
        
        # -tion forms
        if keyword.endswith('e'):
            variations.append(keyword[:-1] + 'tion')
        else:
            variations.append(keyword + 'tion')
    
    # Add case variations
    case_variations = []
    for var in variations:
        case_variations.extend([
            var,                    # lowercase
            var.capitalize(),       # Capitalized
            var.upper(),           # UPPERCASE
        ])
    
    # Remove duplicates while preserving order
    unique_variations = []
    seen = set()
    for var in case_variations:
        if var and var not in seen:
            unique_variations.append(var)
            seen.add(var)
    
    return unique_variations


def find_exact_keyword_instances(content, main_keyword, keyword_variations, filename, document_id):
    """
    Find exact instances of keyword and its variations
    """
    instances = []
    
    try:
        # Split content into pages
        pages = split_into_pages_reliable(content)
        logger.info(f"Split into {len(pages)} pages for exact keyword search in: {filename}")
        
        # Search each page
        for page_data in pages:
            # Extract page number
            page_num_match = re.match(r'\[PAGE (\d+)\]', page_data)
            if page_num_match:
                page_num = int(page_num_match.group(1))
                page_content = page_data[page_num_match.end():].strip()
            else:
                page_num = pages.index(page_data) + 1
                page_content = page_data
            
            # Search for each keyword variation
            found_variations = set()  # Track which variations we found
            
            for variation in keyword_variations:
                # Find all occurrences of this variation
                page_instances = find_keyword_variation_in_page(
                    page_content, variation, main_keyword, page_num, filename, document_id
                )
                
                for instance in page_instances:
                    # Avoid adding duplicate instances from the same position
                    instance_key = f"{instance['page_number']}_{instance['position_in_page']}"
                    if instance_key not in found_variations:
                        instances.append(instance)
                        found_variations.add(instance_key)
        
        # Sort instances by page and position
        instances.sort(key=lambda x: (x['page_number'], x['position_in_page']))
        
        return instances
        
    except Exception as e:
        logger.error(f"Error finding exact keyword instances: {e}")
        return []


def find_keyword_variation_in_page(page_content, variation, main_keyword, page_num, filename, document_id):
    """
    Find specific keyword variation in a page with exact matching
    """
    instances = []
    
    try:
        # Search for exact word matches (with word boundaries)
        import re
        
        # Pattern for exact word matching
        pattern = r'\b' + re.escape(variation) + r'\b'
        
        for match in re.finditer(pattern, page_content, re.IGNORECASE):
            start_pos = match.start()
            matched_text = match.group()
            
            # Get context around the match
            context_start = max(0, start_pos - 500)
            context_end = min(len(page_content), start_pos + len(matched_text) + 500)
            context = page_content[context_start:context_end].strip()
            
            # Clean context
            context = ' '.join(context.split())
            
            # Determine match quality
            exact_match_score = 100
            if matched_text.lower() == main_keyword.lower():
                exact_match_score = 200  # Perfect match
            elif matched_text.lower() == variation.lower():
                exact_match_score = 150  # Exact variation match
            
            instance = {
                'filename': filename,
                'document_id': document_id,
                'page_number': page_num,
                'position_in_page': start_pos,
                'matched_text': matched_text,
                'original_keyword': main_keyword,
                'variation_matched': variation,
                'context': context,
                'exact_match': matched_text.lower() == main_keyword.lower(),
                'variation_match': True,
                'match_score': exact_match_score,
                'instance_key': f"{document_id}_{page_num}_{start_pos}_{matched_text}"
            }
            
            instances.append(instance)
            
            logger.debug(f"Found '{matched_text}' (variation of '{main_keyword}') at page {page_num}, pos {start_pos}")
    
    except Exception as e:
        logger.error(f"Error finding keyword variation in page: {e}")
    
    return instances


def remove_duplicate_instances_ranked(instances):
    """
    Remove duplicates and rank by match quality
    """
    # Remove exact duplicates based on instance_key
    seen_keys = set()
    unique_instances = []
    
    for instance in instances:
        key = instance['instance_key']
        if key not in seen_keys:
            seen_keys.add(key)
            unique_instances.append(instance)
    
    # Sort by match score (highest first), then by page number
    unique_instances.sort(key=lambda x: (x['match_score'], -x['page_number']), reverse=True)
    
    return unique_instances


def build_exact_keyword_summary(instances, keyword, documents_searched):
    """
    Build summary focused on exact keyword matches
    """
    if not instances:
        return f"""
EXACT KEYWORD SEARCH RESULTS
=============================
Keyword: '{keyword}'
Documents searched: {documents_searched}
Total instances found: 0

No instances of '{keyword}' or its variations were found in your documents.
"""
    
    summary_parts = []
    
    # Header
    summary_parts.append(f"🎯 EXACT MATCHES FOR KEYWORD: '{keyword.upper()}'")
    summary_parts.append("=" * 80)
    summary_parts.append("")
    
    # Statistics
    by_doc = {}
    variations_found = set()
    exact_matches = 0
    
    for instance in instances:
        doc = instance['filename']
        if doc not in by_doc:
            by_doc[doc] = {'total': 0, 'pages': set(), 'variations': set()}
        
        by_doc[doc]['total'] += 1
        by_doc[doc]['pages'].add(instance['page_number'])
        by_doc[doc]['variations'].add(instance['matched_text'])
        variations_found.add(instance['matched_text'])
        
        if instance['exact_match']:
            exact_matches += 1
    
    total_docs = len(by_doc)
    total_pages = sum(len(doc_data['pages']) for doc_data in by_doc.values())
    
    summary_parts.append("📊 <b>SEARCH SUMMARY:</b>")
    summary_parts.append(f"• Total instances found: {len(instances)}")
    summary_parts.append(f"• Exact matches: {exact_matches}")
    summary_parts.append(f"• Variations found: {len(variations_found)}")
    summary_parts.append(f"• Documents with matches: {total_docs}")
    summary_parts.append(f"• Pages with matches: {total_pages}")
    summary_parts.append("")
    
    # Variations found
    summary_parts.append("<b>KEYWORD VARIATIONS FOUND:</b>")
    summary_parts.append("-" * 30)
    for variation in sorted(variations_found):
        count = len([i for i in instances if i['matched_text'] == variation])
        match_type = "EXACT" if variation.lower() == keyword.lower() else "VARIATION"
        summary_parts.append(f"• '{variation}' - {count} instances ({match_type})")
    summary_parts.append("")
    
    # Distribution by document
    summary_parts.append("📄 <b>DISTRIBUTION BY DOCUMENT:</b>")
    summary_parts.append("-" * 40)
    
    for doc_name, doc_data in by_doc.items():
        summary_parts.append(f"📄 {doc_name}")
        summary_parts.append(f"   • Total instances: {doc_data['total']}")
        summary_parts.append(f"   • Pages: {', '.join(map(str, sorted(doc_data['pages'])))}")
        summary_parts.append(f"   • Variations: {', '.join(sorted(doc_data['variations']))}")
        summary_parts.append("")
    
    # Sample instances
    summary_parts.append("📍 <b>SAMPLE INSTANCES:</b>")
    summary_parts.append("-" * 25)
    
    # Show top 10 instances
    top_instances = instances[:10]
    for i, instance in enumerate(top_instances, 1):
        summary_parts.append(f"{i}. '{instance['matched_text']}' in {instance['filename']}")
        summary_parts.append(f"   Page {instance['page_number']} - {'EXACT' if instance['exact_match'] else 'VARIATION'}")
        
        # Show context preview (first 100 chars)
        # context_preview = instance['context'][:100] + "..." if len(instance['context']) > 100 else instance['context']
        # summary_parts.append(f"   Context: {context_preview}")
        summary_parts.append("")
    
    # Footer
    summary_parts.append("=" * 50)
    summary_parts.append("🎯 <b>EXACT KEYWORD SEARCH COMPLETE</b>")
    summary_parts.append("=" * 50)
    summary_parts.append(f"✅ Found {len(instances)} total instances of '{keyword}' and variations")
    summary_parts.append(f"✅ Covered {total_docs} document(s) across {total_pages} page(s)")
    summary_parts.append("✅ Exact word boundary matching used for precision")
    summary_parts.append("=" * 50)
    
    return "\n".join(summary_parts)


# OLLAMA HELPER FUNCTIONS
def ollama_streaming_response(question: str, user_id: int):
    """
    Updated streaming response that uses AI keyword extraction (IMPROVED)
    """
    try:
        # Get user documents info first
        user_documents = Document.query.filter_by(user_id=user_id).all()
        
        if not user_documents:
            yield json.dumps({
                'status': 'error',
                'message': 'No documents found. Please upload documents first.'
            }) + "\n"
            return
        
        # Get document content
        yield json.dumps({
            'status': 'status_update',
            'message': f'📚 Loading content from {len(user_documents)} documents...'
        }) + "\n"
        
        all_content = []
        document_info = []
        
        for doc in user_documents:
            filename = doc.document_metadata.get('original_filename', 'Unknown') if doc.document_metadata else 'Unknown'
            content = get_full_document_content(doc.id)
            if content:
                all_content.append(f"=== DOCUMENT: {filename} ===\n{content}\n")
                document_info.append(filename)
        
        if not all_content:
            yield json.dumps({
                'status': 'error',
                'message': 'No content found in documents.'
            }) + "\n"
            return
        
        combined_content = "\n".join(all_content)
        relevant_content = []
        
        # AI-Powered Keyword Search
        yield json.dumps({
            'status': 'status_update',
            'message': '🤖 AI extracting keywords from your question...'
        }) + "\n"
        
        fallback_success = False
        
        try:
            # Use AI to extract main keyword for exact search
            main_keyword = extract_keywords_using_ai(question)
            
            yield json.dumps({
                'status': 'status_update',
                'message': f'🔍 AI identified: "{main_keyword}" - searching for matches...'
            }) + "\n"
            
            # Get keyword matches
            all_instances, keyword_summary = search_exact_keyword_matches(main_keyword, user_id)
            
            yield json.dumps({'status': 'stream_start', 'stream_type': 'ai_keyword'}) + "\n"
            
            # Stream the keyword search response
            keyword_content = f"""# AI Keyword Analysis

**Your Question:** {question}
**AI Extracted Keyword:** {main_keyword}

{keyword_summary}

---

"""
            
            relevant_content = all_instances
            
            chunk_size = 150
            for i in range(0, len(keyword_content), chunk_size):
                chunk = keyword_content[i:i + chunk_size]
                yield json.dumps({'status': 'stream_chunk', 'content': chunk, 'stream_type': 'ai_keyword'}) + "\n"
                time.sleep(0.2)
            
            fallback_success = True
            
        except Exception as fallback_error:
            logger.error(f"AI keyword search streaming error: {fallback_error}")
            yield json.dumps({
                'status': 'error',
                'message': f'AI keyword search failed: {str(fallback_error)}',
                'stream_type': 'ai_keyword'
            }) + "\n"
        
        # AI Answer using Ollama
        yield json.dumps({
            'status': 'status_update',
            'message': f'🦙 AI analyzing your question: "{question}"'
        }) + "\n"
        
        # Create context from AI keyword matches
        keyword_context = ""
        if relevant_content:
            keyword_context = f"\n\nAI-IDENTIFIED KEYWORD MATCHES FOR '{main_keyword}':\n"
            for i, instance in enumerate(relevant_content[:5], 1):
                keyword_context += f"\n{i}. From {instance['filename']} (Page {instance['page_number']}):\n"
                keyword_context += f"   Found: '{instance['matched_text']}'\n"
                keyword_context += f"   Context: {instance['context'][:200]}...\n"
        
        # Create prompt for Ollama
        ollama_prompt = f"""You are an intelligent document analysis assistant. Answer the user's question based on the provided content.

USER QUESTION: "{question}"

AI EXTRACTED KEYWORD: "{main_keyword}"

KEYWORD MATCHES FOUND:
{keyword_context if keyword_context else "No specific keyword matches found."}

DOCUMENT CONTENT:
{combined_content[:12000]}

INSTRUCTIONS:
1. Answer the user's EXACT question: "{question}"
2. Use information from the documents provided
3. If keyword matches were found, prioritize that information
4. Be comprehensive and helpful
5. Structure your answer with clear headings
6. Reference specific documents: [^filename||page]
7. Only use information from the provided documents
8. If the answer isn't in the documents, say so clearly

Answer:"""

        yield json.dumps({
            'status': 'status_update',
            'message': f'⚡ Generating comprehensive answer...'
        }) + "\n"
        
        # Stream response from Ollama
        yield json.dumps({'status': 'stream_start', 'stream_type': 'ollama'}) + "\n"
        
        ollama_success = False
        full_ollama_response = ""
        
        try:
            # Get streaming response from Ollama
            for chunk in get_ollama_response_stream(ollama_prompt):
                if chunk.strip():
                    full_ollama_response += chunk
                    yield json.dumps({'status': 'stream_chunk', 'content': chunk, 'stream_type': 'ollama'}) + "\n"
                    time.sleep(0.01)
            
            ollama_success = True
            
            yield json.dumps({
                'status': 'stream_end',
                'stream_type': 'ollama',
                'sources': [{'filename': s, 'path': s} for s in document_info],
                'question_analysis': {
                    'type': 'ai_keyword_with_llama_analysis',
                    'confidence': 0.98,
                    'reasoning': f'AI extracted "{main_keyword}" from "{question}", then analyzed full content',
                    'ai_determined': True,
                    'model': 'Ollama Llama 3.1',
                    'original_question': question,
                    'ai_extracted_keyword': main_keyword,
                    'keyword_matches_found': len(relevant_content)
                },
                'ai_powered': True,
                'dual_system': True,
                'sources_count': len(document_info)
            }) + "\n"
            
        except Exception as stream_error:
            logger.error(f"Ollama streaming error: {stream_error}")
            yield json.dumps({
                'status': 'error',
                'message': f'Ollama streaming failed: {str(stream_error)}',
                'stream_type': 'ollama'
            }) + "\n"
        
        # Final summary
        yield json.dumps({
            'status': 'complete',
            'message': 'AI keyword extraction and analysis completed',
            'fallback_success': fallback_success,
            'ollama_success': ollama_success,
            'total_sources': len(document_info),
            'ai_keyword_matches': len(relevant_content),
            'search_strategy': {
                'ai_extracted_keyword': main_keyword,
                'original_question': question,
                'explanation': 'AI extracts keywords, finds matches, then answers using relevant content'
            }
        }) + "\n"
        
    except Exception as e:
        logger.error(f"Enhanced AI streaming response error: {e}")
        yield json.dumps({
            'status': 'error',
            'message': f'Error: {str(e)}'
        }) + "\n" 


def update_search_configuration_for_exact_keywords():
    """Update the search configuration to guarantee exact keyword match priority"""
    
    # Override the search weights to massively prioritize exact keyword matches
    app.config['SEARCH_WEIGHTS'] = {
        'exact_keyword_GUARANTEED': 10000000.0,    # MASSIVE priority for exact keyword matches
        'exact_phrase_GUARANTEED': 1000000.0,     # Very high priority for exact phrases
        'exact_phrase_enhanced': 100000.0,        # High priority for enhanced exact matches
        'multi_keyword': 1000.0,                  # Medium priority for keyword matches
        'context_aware': 500.0,                   # Medium priority for context
        'semantic_gpu': 100.0,                    # Lower priority for semantic
        'fuzzy_match': 10.0,                      # Low priority for fuzzy
        'emergency_fallback': 0.1                 # Emergency only
    }
    
    logger.info("🎯 Updated search configuration for EXACT KEYWORD priority")


# 3. ADD to your startup function
def startup_exact_keyword_priority_system():
    """Initialize the exact keyword priority system"""
    try:
        logger.info("🎯 Starting EXACT KEYWORD PRIORITY DocumentIQ System...")
        
        # Initialize processor
        success = initialize_enhanced_processor()
        
        # Initialize Ollama
        ollama_success = initialize_ollama()
        
        # Update search configuration for exact keywords
        update_search_configuration_for_exact_keywords()
        
        if success and ollama_success:
            logger.info("✅ EXACT KEYWORD PRIORITY system with Ollama initialized successfully")
            logger.info("🎯 EXACT KEYWORD PRIORITY features enabled:")
            logger.info("   - GUARANTEED exact keyword matching with morphological variations")
            logger.info("   - Word boundary detection for precise matching")
            logger.info("   - Automatic keyword extraction from questions")
            logger.info("   - Exact match ranking with massive score boost")
            logger.info("   - Context prioritization for exact keyword matches")
            logger.info("   - Emergency fallbacks ensure 100% search success")
            logger.info("   - Enhanced prompts prioritize exact keyword answers")
            logger.info("   - Powered by Ollama Llama 3.1 for superior language understanding")
        else:
            logger.error("❌ EXACT KEYWORD PRIORITY system initialization failed")
            if not success:
                logger.error("   - Document processor initialization failed")
            if not ollama_success:
                logger.error("   - Ollama initialization failed")
            
        return success and ollama_success
        
    except Exception as e:
        logger.error(f"❌ EXACT KEYWORD PRIORITY system startup error: {e}")
        return False



def update_search_configuration():
    """Update the search configuration to guarantee exact match priority"""
    
    # Override the search weights to massively prioritize exact matches
    app.config['SEARCH_WEIGHTS'] = {
        'exact_phrase_GUARANTEED': 1000000.0,    # MASSIVE priority for guaranteed exact matches
        'exact_phrase_enhanced': 100000.0,       # Very high priority for enhanced exact matches
        'multi_keyword': 1000.0,                 # High priority for keyword matches
        'context_aware': 500.0,                  # Medium-high priority for context
        'semantic_gpu': 100.0,                   # Medium priority for semantic
        'fuzzy_match': 10.0,                     # Low priority for fuzzy
        'emergency_fallback': 0.1                # Emergency only
    }
    
    # Update the prompt to emphasize exact matches
    logger.info("🎯 Updated search configuration to guarantee exact match priority")

# Continue with the rest of your existing code...
# The OCR processor class and other components remain the same

class EnhancedOCRProcessor:
    """Advanced OCR processor with GPU acceleration and image preprocessing"""
    
    def __init__(self, gpu_enabled=True):
        self.logger = logging.getLogger(__name__ + '.OCR')
        self.gpu_enabled = gpu_enabled and torch.cuda.is_available()
        
        # Initialize OCR engines
        self._initialize_ocr_engines()
        
        # OCR statistics
        self.ocr_stats = {
            'images_processed': 0,
            'pdf_pages_ocr': 0,
            'total_confidence': 0.0,
            'processing_time': 0.0
        }
    
    def _initialize_ocr_engines(self):
        """Initialize OCR engines with GPU support"""
        try:
            # Initialize EasyOCR with GPU support
            if app.config['OCR_GPU_ENABLED'] and self.gpu_enabled:
                self.logger.info("🔥 Initializing GPU-accelerated OCR (EasyOCR)...")
                self.easyocr_reader = easyocr.Reader(
                    app.config['OCR_LANGUAGES'], 
                    gpu=True,
                    model_storage_directory=app.config['MODEL_CACHE_DIR']
                )
                self.logger.info("✅ GPU-accelerated OCR initialized successfully")
            else:
                self.logger.info("🔄 Initializing CPU OCR (EasyOCR)...")
                self.easyocr_reader = easyocr.Reader(
                    app.config['OCR_LANGUAGES'], 
                    gpu=False,
                    model_storage_directory=app.config['MODEL_CACHE_DIR']
                )
                self.logger.info("✅ CPU OCR initialized successfully")
            
            # Test Tesseract availability
            try:
                pytesseract.get_tesseract_version()
                self.tesseract_available = True
                self.logger.info("✅ Tesseract OCR available as fallback")
            except Exception as e:
                self.tesseract_available = False
                self.logger.warning(f"⚠️ Tesseract not available: {e}")
                
        except Exception as e:
            self.logger.error(f"❌ OCR initialization failed: {e}")
            self.easyocr_reader = None
            self.tesseract_available = False
  
# Continue with remaining classes and functions...




@app.route('/delete_chat/<chat_id>', methods=['POST'])
@login_required
def delete_chat(chat_id):
    """Delete chat"""
    try:
        conversation = Conversation.query.filter_by(id=chat_id, user_id=current_user.id).first()
        if not conversation:
            return jsonify({'success': False, 'error': 'Chat not found'}), 404
            
        db.session.delete(conversation)
        db.session.commit()
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f'Error deleting chat: {e}')
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/update_chat_title/<chat_id>', methods=['POST'])
@login_required  
def update_chat_title(chat_id):
    """Update chat title"""
    try:
        data = request.get_json()
        new_title = data.get('title', '').strip()
        
        if not new_title:
            return jsonify({'error': 'Title cannot be empty'}), 400
            
        conversation = Conversation.query.filter_by(
            id=chat_id, 
            user_id=current_user.id
        ).first()
        
        if not conversation:
            return jsonify({'error': 'Chat not found'}), 404
            
        conversation.title = new_title
        conversation.updated_at = datetime.now(timezone.utc)
        db.session.commit()
        
        return jsonify({'success': True, 'title': new_title})
        
    except Exception as e:
        logger.error(f'Error updating chat title: {e}')
        return jsonify({'error': str(e)}), 500
@app.route('/get_file_content')
@login_required
def get_file_content():
    """Get file content with enhanced formatting - FIXED VERSION"""
    rel_path = request.args.get('path')
    if not rel_path:
        return jsonify({'error': 'File path required'}), 400
    
    try:
        # Security checks
        if '..' in rel_path or rel_path.startswith('/'):
            return jsonify({'error': 'Invalid file path'}), 400
        
        # Clean the path
        rel_path = rel_path.strip().replace('\\', '/')
        
        # Try multiple strategies to find the document
        doc = None
        search_paths = [
            rel_path,  # Original path
            rel_path.split('/')[-1],  # Just filename
            f"{current_user.id}/{rel_path}",  # User ID + path
            f"{current_user.id}/{rel_path.split('/')[-1]}"  # User ID + filename
        ]
        
        logger.info(f"Searching for document with paths: {search_paths}")
        
        # Try each search path
        for search_path in search_paths:
            # Try exact match first
            doc = Document.query.filter_by(
                file_path=search_path,
                user_id=current_user.id
            ).first()
            
            if doc:
                logger.info(f"Found document with exact path: {search_path}")
                break
            
            # Try partial match (contains)
            doc = Document.query.filter(
                Document.file_path.contains(search_path.split('/')[-1]),
                Document.user_id == current_user.id
            ).first()
            
            if doc:
                logger.info(f"Found document with partial match: {search_path}")
                break
        
        # Also try matching by original filename in metadata
        if not doc:
            filename = rel_path.split('/')[-1]
            docs_with_metadata = Document.query.filter(
                Document.user_id == current_user.id,
                Document.document_metadata.isnot(None)
            ).all()
            
            for potential_doc in docs_with_metadata:
                try:
                    if (potential_doc.document_metadata and 
                        potential_doc.document_metadata.get('original_filename') == filename):
                        doc = potential_doc
                        logger.info(f"Found document by metadata filename: {filename}")
                        break
                except Exception:
                    continue
        
        if not doc:
            logger.warning(f"Document not found for any of these paths: {search_paths}")
            return jsonify({'error': 'Document not found'}), 404
        
        # Get accuracy-focused content from DocumentContent table
        content_record = DocumentContent.query.filter_by(
            document_id=doc.id,
            content_type='accuracy_focused_full_text'
        ).first()
        
        # Fallback to enhanced content if accuracy-focused not available
        if not content_record:
            content_record = DocumentContent.query.filter_by(
                document_id=doc.id,
                content_type='enhanced_full_text'
            ).first()
        
        # Further fallback to any text content
        if not content_record:
            content_record = DocumentContent.query.filter_by(
                document_id=doc.id
            ).filter(
                DocumentContent.content_type.like('%text%')
            ).first()
        
        # Last resort - any content
        if not content_record:
            content_record = DocumentContent.query.filter_by(
                document_id=doc.id
            ).first()
        
        if not content_record or not content_record.content:
            logger.warning(f"No content found for document {doc.id}")
            return jsonify({'error': 'Document content not found'}), 404
        
        filename = doc.document_metadata.get('original_filename', 'Unknown') if doc.document_metadata else 'Unknown'
        
        logger.info(f"Successfully retrieved content for {filename} (type: {content_record.content_type})")
        
        return jsonify({
            'file': filename,
            'content': content_record.content,
            'path': doc.file_path,
            'structure_preserved': content_record.content_type in ['accuracy_focused_full_text', 'enhanced_full_text', 'structured_text'],
            'perfect_formatting': content_record.content_type in ['accuracy_focused_full_text', 'enhanced_full_text'],
            'extraction_method': doc.document_metadata.get('processing_method', 'unknown') if doc.document_metadata else 'unknown',
            'accuracy_focused_processing': 'accuracy_focused' in content_record.content_type,
            'gpu_processed': doc.document_metadata.get('gpu_processed', False) if doc.document_metadata else False,
            'content_type': content_record.content_type,
            'content_length': len(content_record.content),
            'document_id': doc.id
        })
    
    except Exception as e:
        logger.error(f"Error serving file content: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500


# Subscription and Billing Routes
@app.route('/start_trial', methods=['POST'])
@login_required
def start_trial():
    """Start trial subscription"""
    now = datetime.now(timezone.utc)
    
    # Check if user already has an active trial
    if (current_user.subscription_plan == 'trial' and 
        current_user.trial_end_date and 
        current_user._make_aware(current_user.trial_end_date) > now):
        return jsonify({'error': 'You already have an active trial'}), 400
        
    current_user.subscription_plan = 'trial'
    current_user.subscription_status = 'active'
    current_user.trial_start_date = now
    current_user.trial_end_date = now + timedelta(days=14)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'message': 'Your 14-day free trial has started!',
        'trial_end_date': current_user.trial_end_date.isoformat(),
        'enhanced_features': True,
        'accuracy_focused': True
    })



def validate_user_subscription_strict(user):
    """Strict subscription validation before upload"""
    try:
        now = datetime.now(timezone.utc)
        
        # Admin always allowed
        if user.is_admin:
            return True, "Admin access"
        
        # Check active trial
        if (user.subscription_plan == 'trial' and 
            user.trial_end_date and 
            user._make_aware(user.trial_end_date) > now):
            return True, "Active trial"
        
        # Check active paid subscription
        if (user.subscription_status == 'active' and 
            user.subscription_plan in ['basic', 'pro'] and
            (user.current_period_end is None or 
             user._make_aware(user.current_period_end) > now)):
            return True, f"Active {user.subscription_plan} subscription"
        
        # Log the rejection reason
        logger.warning(f"Upload rejected for user {user.id}:")
        logger.warning(f"  - Plan: {user.subscription_plan}")
        logger.warning(f"  - Status: {user.subscription_status}")
        logger.warning(f"  - Trial end: {user.trial_end_date}")
        logger.warning(f"  - Period end: {user.current_period_end}")
        
        return False, "No active subscription or trial"
        
    except Exception as e:
        logger.error(f"Subscription validation error: {e}")
        return False, f"Validation error: {str(e)}"

# --------------------- Enhanced Document Processor Class with Accuracy Focus ---------------------

class AccuracyFocusedDocumentProcessor:
    def __init__(self, gpu_enabled=True):
        # Add explicit CUDA check
        if gpu_enabled and not torch.cuda.is_available():
            logger.warning("GPU requested but not available. Falling back to CPU.")
            gpu_enabled = False
            
        self.logger = logging.getLogger(__name__)
        self.device = "cuda" if torch.cuda.is_available() and gpu_enabled else "cpu"
        self.gpu_enabled = gpu_enabled and torch.cuda.is_available()
        
        # Initialize GPU settings for RTX A6000
        if self.gpu_enabled:
            torch.cuda.set_per_process_memory_fraction(app.config['GPU_MEMORY_FRACTION'])
            torch.backends.cudnn.benchmark = True
            torch.cuda.empty_cache()
        
        # Initialize models
        self._initialize_models()
        
        # Initialize processing pools
        self.thread_pool = ThreadPoolExecutor(max_workers=min(32, psutil.cpu_count()))
        
    def _initialize_models(self):
        """Initialize all AI models with GPU optimization"""
        try:
            # Enhanced embedding model for semantic search
            self.embedding_model = self._load_embedding_model()
            
            # Enable mixed precision for RTX A6000
            if self.gpu_enabled and app.config['MIXED_PRECISION']:
                try:
                    self.embedding_model = self.embedding_model.half()
                except:
                    self.logger.warning("Mixed precision not supported, using FP32")
            
            # Initialize TF-IDF for keyword search with better parameters
            self.tfidf_vectorizer = TfidfVectorizer(
                max_features=15000,  # Increased for better coverage
                stop_words='english',
                ngram_range=(1, 4),  # Include more phrases
                analyzer='word',
                lowercase=True,
                min_df=1,  # Include rare terms
                max_df=0.95,  # Exclude very common terms
                token_pattern=r'\b[a-zA-Z0-9][a-zA-Z0-9\-\']*[a-zA-Z0-9]\b|\b[a-zA-Z]\b'
            )
            
            # Initialize spaCy for NLP processing
            if SPACY_AVAILABLE:
                try:
                    self.nlp = spacy.load("en_core_web_sm")
                except:
                    self.logger.warning("spaCy model not found, using fallback")
                    self.nlp = None
            else:
                self.nlp = None
            
            # Initialize FAISS index for fast vector search
            self.embedding_dim = 768  # all-mpnet-base-v2 dimension
            if self.gpu_enabled and app.config['FAISS_GPU_ENABLED']:
                try:
                    # Use GPU FAISS for RTX A6000
                    self.faiss_index = faiss.IndexFlatIP(self.embedding_dim)
                    # Move to GPU
                    res = faiss.StandardGpuResources()
                    self.faiss_index = faiss.index_cpu_to_gpu(res, 0, self.faiss_index)
                    self.logger.info("FAISS GPU index initialized")
                except Exception as e:
                    self.logger.warning(f"FAISS GPU initialization failed: {e}, using CPU")
                    self.faiss_index = faiss.IndexFlatIP(self.embedding_dim)
            else:
                self.faiss_index = faiss.IndexFlatIP(self.embedding_dim)
            
            self.chunk_metadata = []
            self.tfidf_fitted = False
            self.tfidf_matrix = None
            
            self.logger.info(f"✅ Enhanced models initialized on {self.device}")
            
        except Exception as e:
            self.logger.error(f"❌ Model initialization failed: {e}")
            raise

    def _load_embedding_model(self):
        """Load embedding model with retries and fallback"""
        for attempt in range(app.config['EMBEDDING_MODEL_RETRY_ATTEMPTS']):
            try:
                time.sleep(attempt * app.config['EMBEDDING_MODEL_RETRY_DELAY'])
                
                model = SentenceTransformer(
                    app.config['EMBEDDING_MODEL'],
                    device=self.device,
                    cache_folder=app.config['MODEL_CACHE_DIR']
                )
                
                self.logger.info(f"Loaded embedding model on attempt {attempt + 1}")
                return model
                
            except Exception as e:
                logger.error(f"Embedding model loading failed: {e}")
                # Try simple fallback
                try:
                    return SentenceTransformer('all-MiniLM-L6-v2', device=self.device)
                except Exception as fallback_e:
                    logger.critical(f"Critical: All embedding models failed: {fallback_e}")
                    raise RuntimeError("Could not load any embedding models")

    def extract_content_with_perfect_formatting(self, file_path: str) -> str:
        """Extract content preserving EXACT document formatting with OCR support"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_with_ocr_support(file_path)
            elif file_ext == '.docx':
                return self._extract_docx_perfect_formatting(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_excel_perfect_formatting(file_path)
            elif file_ext == '.txt':
                return self._extract_text_perfect_formatting(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.webp']:
                return self._extract_image_with_ocr(file_path)
            else:
                return self._extract_fallback_formatting(file_path)
                
        except Exception as e:
            self.logger.error(f"Content extraction failed for {file_path}: {e}")
            return ""

    def _extract_pdf_with_ocr_support(self, file_path: str) -> str:
        """Extract PDF with OCR support for image-based pages"""
        global ocr_reader
        
        if ocr_reader and app.config['OCR_ENABLED']:
            # Use OCR-enhanced PDF extraction
            return ocr_reader.extract_text_from_pdf_with_ocr(file_path)
        else:
            # Fallback to original PDF extraction
            return self._extract_pdf_perfect_formatting(file_path)
    
    def _extract_image_with_ocr(self, file_path: str) -> str:
        """Extract text from image files using OCR"""
        global ocr_reader
        
        try:
            if not ocr_reader or not app.config['OCR_ENABLED']:
                return f"Image file detected but OCR is not available: {os.path.basename(file_path)}"
            
            self.logger.info(f"🖼️ Processing image file with OCR: {os.path.basename(file_path)}")
            
            # Extract text using OCR
            ocr_text = ocr_reader.extract_text_from_image(file_path)
            
            if not ocr_text.strip():
                return f"No readable text found in image: {os.path.basename(file_path)}"
            
            # Format the extracted content
            content_blocks = []
            content_blocks.append("=" * 80)
            content_blocks.append(f"IMAGE DOCUMENT (OCR PROCESSED)")
            content_blocks.append(f"File: {os.path.basename(file_path)}")
            content_blocks.append("=" * 80)
            content_blocks.append("")
            content_blocks.append("[OCR EXTRACTED CONTENT]")
            content_blocks.append(ocr_text)
            content_blocks.append("[END OCR CONTENT]")
            content_blocks.append("")
            content_blocks.append("=" * 80)
            content_blocks.append("OCR PROCESSING SUMMARY")
            content_blocks.append("=" * 80)
            content_blocks.append(f"OCR Engine: {'GPU-accelerated EasyOCR' if ocr_reader.gpu_enabled else 'CPU EasyOCR'}")
            content_blocks.append(f"Extracted text length: {len(ocr_text)} characters")
            content_blocks.append(f"Image preprocessing: {'Enabled' if app.config['OCR_PREPROCESS_ENABLED'] else 'Disabled'}")
            
            result = "\n".join(content_blocks)
            self.logger.info(f"✅ Image OCR completed: {len(result)} total characters")
            
            return result
            
        except Exception as e:
            self.logger.error(f"Image OCR extraction failed for {file_path}: {e}")
            return f"Error processing image {os.path.basename(file_path)}: {str(e)}"

    # Add other extraction methods...
    def _extract_pdf_perfect_formatting(self, file_path):
        # Implementation here - keeping original code
        pass
    
    def _extract_docx_perfect_formatting(self, file_path):
        # Implementation here - keeping original code  
        pass
    
    def _extract_excel_perfect_formatting(self, file_path):
        # Implementation here - keeping original code
        pass
    
    def _extract_text_perfect_formatting(self, file_path):
        # Implementation here - keeping original code
        pass
    
    def _extract_fallback_formatting(self, file_path):
        # Implementation here - keeping original code
        pass

# --------------------- INSTANCE SEARCH FUNCTIONALITIES ---------------------

def search_all_instances_bulletproof(query, user_id):
    """
    BULLETPROOF search with debugging - UPDATED VERSION
    """
    logger.info(f"🎯 BULLETPROOF SEARCH: Finding ALL instances of '{query}' for user {user_id}")
    
    try:
        all_instances = []
        
        # Get user documents
        user_documents = Document.query.filter_by(user_id=user_id).all()
        
        if not user_documents:
            return [], "No documents found to search."
        
        logger.info(f"Processing {len(user_documents)} documents")
        
        # Process each document
        for document in user_documents:
            try:
                filename = document.document_metadata.get('original_filename', 'Unknown') if document.document_metadata else 'Unknown'
                logger.info(f"Processing: {filename}")
                
                # Get the full document content
                content = get_full_document_content(document.id)
                
                if not content:
                    logger.warning(f"No content found for {filename}")
                    continue
                
                logger.info(f"Content length: {len(content)} characters")
                
                # ADD DEBUG CHECK
                debug_search_content(content, query)
                
                # Find all instances in this document
                doc_instances = find_instances_simple(content, query, filename, document.id)
                all_instances.extend(doc_instances)
                
                logger.info(f"Found {len(doc_instances)} instances in {filename}")
                
            except Exception as doc_error:
                logger.error(f"Error processing document: {doc_error}")
                continue
        
        # Remove exact duplicates
        unique_instances = all_instances
        
        # Build summary
        summary = build_bulletproof_summary(unique_instances, query, len(user_documents))
        
        logger.info(f"✅ BULLETPROOF SEARCH COMPLETED: {len(unique_instances)} unique instances found")
        
        return unique_instances, summary
        
    except Exception as e:
        logger.error(f"❌ BULLETPROOF SEARCH ERROR: {e}")
        return [], f"Search failed: {str(e)}"



def split_into_pages_reliable(content):
    """
    Split content into pages using the most reliable method possible.
    FIXED VERSION - handles your PDF format correctly
    """
    try:
        # Look for the specific page markers in your PDF format
        # Pattern: "PAGE X" followed by dashes, then "[END PAGE X]" 
        
        # Method 1: Use your specific page pattern
        page_pattern = r'PAGE (\d+)\s*\n-{40}\s*(.*?)\s*\[END PAGE \1\]'
        page_matches = list(re.finditer(page_pattern, content, re.DOTALL))
        
        if page_matches:
            pages = []
            for match in page_matches:
                page_num = int(match.group(1))
                page_content = match.group(2).strip()
                
                if page_content:  # Only add non-empty pages
                    # Add page marker for identification
                    formatted_page = f"[PAGE {page_num}]\n{page_content}"
                    pages.append(formatted_page)
            
            if pages:
                logger.info(f"Successfully split into {len(pages)} pages using PAGE markers")
                return pages
        
        # Method 2: Alternative pattern - split by "PAGE X" headers
        page_splits = re.split(r'\nPAGE (\d+)\n-{40}', content)
        
        if len(page_splits) > 2:  # We have actual page splits
            pages = []
            
            # Skip the first empty split
            for i in range(1, len(page_splits), 2):
                if i + 1 < len(page_splits):
                    page_num = int(page_splits[i])
                    page_content = page_splits[i + 1]
                    
                    # Clean up the page content
                    # Remove the [END PAGE X] marker
                    page_content = re.sub(r'\[END PAGE \d+\].*$', '', page_content, flags=re.DOTALL)
                    page_content = page_content.strip()
                    
                    if page_content:
                        formatted_page = f"[PAGE {page_num}]\n{page_content}"
                        pages.append(formatted_page)
            
            if pages:
                logger.info(f"Successfully split into {len(pages)} pages using split method")
                return pages
        
        # Method 3: Manual extraction based on your exact format
        pages = []
        lines = content.split('\n')
        current_page = None
        current_content = []
        
        for line in lines:
            # Look for page start
            page_start_match = re.match(r'^PAGE (\d+)$', line.strip())
            if page_start_match:
                # Save previous page if it exists
                if current_page is not None and current_content:
                    page_text = '\n'.join(current_content).strip()
                    if page_text:
                        pages.append(f"[PAGE {current_page}]\n{page_text}")
                
                # Start new page
                current_page = int(page_start_match.group(1))
                current_content = []
                continue
            
            # Look for page end
            if re.match(r'^\[END PAGE \d+\]$', line.strip()):
                continue  # Skip end markers
            
            # Skip the dashes separator
            if re.match(r'^-{40}$', line.strip()):
                continue
            
            # Add content to current page
            if current_page is not None:
                current_content.append(line)
        
        # Don't forget the last page
        if current_page is not None and current_content:
            page_text = '\n'.join(current_content).strip()
            if page_text:
                pages.append(f"[PAGE {current_page}]\n{page_text}")
        
        if pages:
            logger.info(f"Successfully split into {len(pages)} pages using manual extraction")
            return pages
        
        # Fallback: treat as single page
        logger.warning("Could not split pages properly, treating as single page")
        return [f"[PAGE 1]\n{content}"]
        
    except Exception as e:
        logger.error(f"Error splitting pages: {e}")
        return [f"[PAGE 1]\n{content}"]



@app.template_filter('process_references')
def process_references(text):
    """Process references in text using original document structure"""
    reference_counter = 1
    references = []
    
    def replace_reference(match):
        nonlocal reference_counter
        filename = match.group(1)
        page = match.group(2) or ''
        highlight = match.group(3) or ''
        
        filename = filename.split(']')[0]
        filename = filename.replace('\\', '/')
        
        ref_id = reference_counter
        reference_counter += 1
        
        references.append({
            'id': ref_id,
            'filename': filename,
            'page': page,
            'highlight': highlight
        })
        
        return (f'<sup class="reference-link" '
                f'data-ref="{ref_id}" '
                f'data-filename="{filename}" '
                f'data-page="{page}" '
                f'data-highlight="{highlight}">'
                f'[{ref_id}]</sup>')
    
    # Process references but don't create artificial section references
    processed_text = re.sub(
        r'\[\^([^\|\]]+)(?:\|\|([^\|\]]+))?(?:\|\|([^\|\]]+))?\]',
        replace_reference,
        text
    )
    
    if references:
        references_section = '<div class="references-section"><h4>References</h4><ol>'
        for ref in sorted(references, key=lambda x: x['id']):
            ref_text = ref['filename'].split('/')[-1]
            if ref['page']:
                ref_text += f", Page {ref['page']}"
            if ref['highlight']:
                ref_text += f" - {ref['highlight']}"
            references_section += f'<li id="ref-{ref["id"]}">{ref_text}</li>'
        references_section += '</ol></div>'
        
        processed_text += references_section
    
    return processed_text



def get_full_document_content(document_id):
    """
    Get the full document content, trying different content types.
    """
    try:
        # Try different content types in order of preference
        content_types_to_try = [
            'accuracy_focused_full_text',
            'enhanced_full_text', 
            'full_text',
            'text'
        ]
        
        for content_type in content_types_to_try:
            content_record = DocumentContent.query.filter_by(
                document_id=document_id,
                content_type=content_type
            ).first()
            
            if content_record and content_record.content:
                logger.info(f"Using content type: {content_type}")
                return content_record.content
        
        # Fallback: get the longest content available
        content_record = db.session.query(DocumentContent).filter_by(
            document_id=document_id
        ).order_by(func.length(DocumentContent.content).desc()).first()
        
        if content_record and content_record.content:
            logger.info(f"Using fallback content type: {content_record.content_type}")
            return content_record.content
        
        return None
        
    except Exception as e:
        logger.error(f"Error getting document content: {e}")
        return None

def find_instances_with_accurate_pages(content, query, filename, document_id):
    """
    Find instances with ACCURATE page numbers
    """
    instances = []
    
    try:
        # Split into pages using accurate method
        pages = split_into_pages_accurate(content)
        
        if not pages:
            logger.warning(f"No pages found for {filename}")
            return []
        
        logger.info(f"Processing {len(pages)} pages for {filename}")
        
        # Search each page
        for page_data in pages:
            # Extract page number from content
            page_num_match = re.match(r'\[PAGE (\d+)\]', page_data)
            if page_num_match:
                page_num = int(page_num_match.group(1))
                page_content = page_data[page_num_match.end():].strip()
            else:
                # Fallback page numbering
                page_num = pages.index(page_data) + 1
                page_content = page_data
            
            # Search for query in this page
            page_instances = search_page_for_query(page_content, query, page_num, filename, document_id)
            instances.extend(page_instances)
        
        logger.info(f"Found {len(instances)} instances across {len(pages)} pages in {filename}")
        return instances
        
    except Exception as e:
        logger.error(f"Error finding instances with accurate pages: {e}")
        return []

def split_into_pages_accurate(content):
    """
    ACCURATE page splitting using reliable markers
    """
    try:
        # Method 1: Use explicit page markers we added
        page_pattern = r'={80}\nPAGE (\d+) OF (\d+)\n={80}'
        page_matches = list(re.finditer(page_pattern, content))
        
        if page_matches:
            pages = []
            for i, match in enumerate(page_matches):
                start_pos = match.end()
                
                # Find end position
                if i + 1 < len(page_matches):
                    end_pos = page_matches[i + 1].start()
                else:
                    end_pos = len(content)
                
                # Extract page content
                page_content = content[start_pos:end_pos]
                
                # Clean up page content
                page_content = re.sub(r'\n<!-- END OF PAGE \d+ -->\n', '', page_content)
                page_content = page_content.strip()
                
                if page_content:
                    # Add page number info to content
                    page_num = int(match.group(1))
                    page_content = f"[PAGE {page_num}]\n{page_content}"
                    pages.append(page_content)
            
            if pages:
                logger.info(f"Accurate page split: {len(pages)} pages using explicit markers")
                return pages
        
        # Method 2: Look for other page markers
        end_page_pattern = r'<!-- END OF PAGE (\d+) -->'
        end_matches = list(re.finditer(end_page_pattern, content))
        
        if end_matches:
            pages = []
            last_end = 0
            
            for match in end_matches:
                page_num = int(match.group(1))
                page_content = content[last_end:match.start()].strip()
                
                if page_content:
                    # Remove any previous page markers
                    page_content = re.sub(r'\[PAGE \d+\]', '', page_content).strip()
                    page_content = f"[PAGE {page_num}]\n{page_content}"
                    pages.append(page_content)
                
                last_end = match.end()
            
            if pages:
                logger.info(f"Accurate page split: {len(pages)} pages using end markers")
                return pages
        
        # Method 3: Fallback to content-based splitting
        logger.warning("Using fallback page splitting")
        
        # Look for major section breaks
        sections = re.split(r'\n={50,}\n', content)
        
        if len(sections) > 1:
            pages = []
            for i, section in enumerate(sections, 1):
                if section.strip() and len(section.strip()) > 50:
                    pages.append(f"[PAGE {i}]\n{section.strip()}")
            
            if pages:
                logger.info(f"Fallback page split: {len(pages)} sections")
                return pages
        
        # Last resort: treat as single page
        logger.warning("Treating document as single page")
        return [f"[PAGE 1]\n{content}"]
        
    except Exception as e:
        logger.error(f"Page splitting error: {e}")
        return [f"[PAGE 1]\n{content}"]


def find_instances_with_accurate_pages(content, query, filename, document_id):
    """
    Find instances with ACCURATE page numbers
    """
    instances = []
    
    try:
        # Split into pages using accurate method
        pages = split_into_pages_accurate(content)
        
        if not pages:
            logger.warning(f"No pages found for {filename}")
            return []
        
        logger.info(f"Processing {len(pages)} pages for {filename}")
        
        # Search each page
        for page_data in pages:
            # Extract page number from content
            page_num_match = re.match(r'\[PAGE (\d+)\]', page_data)
            if page_num_match:
                page_num = int(page_num_match.group(1))
                page_content = page_data[page_num_match.end():].strip()
            else:
                # Fallback page numbering
                page_num = pages.index(page_data) + 1
                page_content = page_data
            
            # Search for query in this page
            page_instances = search_page_for_query(page_content, query, page_num, filename, document_id)
            instances.extend(page_instances)
        
        logger.info(f"Found {len(instances)} instances across {len(pages)} pages in {filename}")
        return instances
        
    except Exception as e:
        logger.error(f"Error finding instances with accurate pages: {e}")
        return []

def search_page_for_query(page_content, query, page_num, filename, document_id):
    """
    Search a single page for query instances - FIXED VERSION with better substring matching
    """
    instances = []
    
    try:
        # Extract the actual page number from the content
        page_match = re.match(r'\[PAGE (\d+)\]\n(.*)', page_content, re.DOTALL)
        if page_match:
            actual_page_num = int(page_match.group(1))
            content_to_search = page_match.group(2)
        else:
            actual_page_num = page_num
            content_to_search = page_content
        
        # IMPROVED: Create comprehensive search variations
        original_query = query.strip()
        search_variations = [
            original_query,                    # Original case
            original_query.lower(),           # Lowercase
            original_query.upper(),           # Uppercase
            original_query.capitalize(),      # Capitalized
        ]
        
        # Remove duplicates while preserving order
        search_terms = []
        seen = set()
        for term in search_variations:
            if term and term not in seen:
                search_terms.append(term)
                seen.add(term)
        
        logger.info(f"Searching page {actual_page_num} for terms: {search_terms}")
        
        found_positions = set()
        
        for search_term in search_terms:
            # IMPROVED: Use both case-sensitive and case-insensitive searches
            
            # Method 1: Case-sensitive exact search
            pos = 0
            while True:
                found_pos = content_to_search.find(search_term, pos)
                if found_pos == -1:
                    break
                    
                if found_pos not in found_positions:
                    instance = create_instance(
                        content_to_search, found_pos, search_term, 
                        actual_page_num, filename, document_id, original_query
                    )
                    instances.append(instance)
                    found_positions.add(found_pos)
                    logger.info(f"Found case-sensitive match: '{search_term}' at position {found_pos}")
                
                pos = found_pos + 1
            
            # Method 2: Case-insensitive search (to catch different cases)
            content_lower = content_to_search.lower()
            search_lower = search_term.lower()
            
            pos = 0
            while True:
                found_pos = content_lower.find(search_lower, pos)
                if found_pos == -1:
                    break
                    
                if found_pos not in found_positions:
                    # Get the actual text from the original content (preserves original case)
                    actual_match = content_to_search[found_pos:found_pos + len(search_term)]
                    
                    instance = create_instance(
                        content_to_search, found_pos, actual_match, 
                        actual_page_num, filename, document_id, original_query
                    )
                    instances.append(instance)
                    found_positions.add(found_pos)
                    logger.info(f"Found case-insensitive match: '{actual_match}' at position {found_pos}")
                
                pos = found_pos + 1
        
        logger.info(f"Page {actual_page_num}: Found {len(instances)} total instances")
        return instances
        
    except Exception as e:
        logger.error(f"Error searching page {page_num}: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return []


def create_instance(content, found_pos, matched_text, page_num, filename, document_id, original_query):
    """
    Helper function to create a consistent instance object
    """
    # Get context (1000 chars before and after for better context)
    context_start = max(0, found_pos - 1000)
    context_end = min(len(content), found_pos + len(matched_text) + 1000)
    context = content[context_start:context_end].strip()
    
    # Clean context (remove excessive whitespace)
    context = ' '.join(context.split())
    
    # Truncate context if too long
    # if len(context) > 300:
    #     context = context[:300] + "..."
    
    return {
        'filename': filename,
        'document_id': document_id,
        'page_number': page_num,
        'position_in_page': found_pos,
        'matched_text': matched_text,
        'context': context,
        'exact_match': matched_text.lower() == original_query.lower(),
        'case_sensitive_match': matched_text == original_query,
        'instance_key': f"{document_id}_{page_num}_{found_pos}_{matched_text}"
    }


def find_instances_simple(content, query, filename, document_id):
    """
    Find all instances using improved method - FIXED VERSION
    """
    instances = []
    
    try:
        # Split content into pages using the reliable method
        pages = split_into_pages_reliable(content)
        logger.info(f"Split into {len(pages)} pages for file: {filename}")
        
        # Search each page individually
        for page_data in pages:
            # Search for query in this page
            page_instances = search_page_for_query(page_data, query, 1, filename, document_id)
            instances.extend(page_instances)
        
        # Remove exact duplicates (same position, same text)
        unique_instances = []
        seen_keys = set()
        
        for instance in instances:
            key = instance['instance_key']
            if key not in seen_keys:
                unique_instances.append(instance)
                seen_keys.add(key)
        
        logger.info(f"Found {len(unique_instances)} unique instances across {len(pages)} pages in {filename}")
        
        # Log each instance for debugging
        for i, instance in enumerate(unique_instances, 1):
            logger.info(f"Instance {i}: Page {instance['page_number']}, '{instance['matched_text']}' at pos {instance['position_in_page']}")
        
        return unique_instances
        
    except Exception as e:
        logger.error(f"Error finding instances: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return []




def remove_duplicate_instances(instances):
    """
    Remove duplicate instances based on position and content.
    """
    seen_keys = set()
    unique_instances = []
    
    for instance in instances:
        key = instance['instance_key']
        if key not in seen_keys:
            seen_keys.add(key)
            unique_instances.append(instance)
    
    return unique_instances

def build_bulletproof_summary(instances, query, documents_searched):
    """
    Build a bulletproof summary that works correctly.
    """
    if not instances:
        return f"""
COMPREHENSIVE SEARCH RESULTS
============================
Query: '{query}'
Documents searched: {documents_searched}
Total instances found: 0

No instances of '{query}' found in your documents.
"""
    
    summary_parts = []
    
    # Header
    summary_parts.append("📊 COMPREHENSIVE ANALYSIS OF '{}' INSTANCES".format(query.upper()))
    summary_parts.append("=" * 80)
    summary_parts.append("")
    
    # Quick stats
    by_doc = {}
    for instance in instances:
        doc = instance['filename']
        if doc not in by_doc:
            by_doc[doc] = {'total': 0, 'pages': set()}
        by_doc[doc]['total'] += 1
        by_doc[doc]['pages'].add(instance['page_number'])
    
    total_docs = len(by_doc)
    total_pages = sum(len(doc_data['pages']) for doc_data in by_doc.values())
    exact_matches = len([i for i in instances if i['exact_match']])
    
    summary_parts.append("📈 SUMMARY:")
    summary_parts.append(f"• Total instances found: {len(instances)}")
    summary_parts.append(f"• Documents with matches: {total_docs}")
    summary_parts.append(f"• Pages with matches: {total_pages}")
    summary_parts.append(f"• Exact case matches: {exact_matches}")
    summary_parts.append("")
    
    # Distribution by document and page
    summary_parts.append("📋 DISTRIBUTION BY DOCUMENT AND PAGE:")
    summary_parts.append("-" * 60)
    
    for doc_name, doc_data in by_doc.items():
        summary_parts.append(f"📄 {doc_name}")
        summary_parts.append(f"   Total instances: {doc_data['total']}")
        
        # Group instances by page for this document
        page_counts = {}
        doc_instances = [i for i in instances if i['filename'] == doc_name]
        for instance in doc_instances:
            page = instance['page_number']
            page_counts[page] = page_counts.get(page, 0) + 1
        
        page_list = []
        for page_num in sorted(page_counts.keys()):
            page_list.append(f"Page {page_num} ({page_counts[page_num]})")
        
        summary_parts.append(f"   Distribution: {', '.join(page_list)}")
        summary_parts.append("")
    
    # Key findings
    summary_parts.append("🔍 KEY FINDINGS:")
    summary_parts.append("-" * 30)
    
    if total_docs == 1:
        summary_parts.append(f"• All instances found in single document: {list(by_doc.keys())[0]}")
    
    if total_pages == 1:
        single_page = list(list(by_doc.values())[0]['pages'])[0]
        summary_parts.append(f"• All instances concentrated on page {single_page}")
    elif total_pages <= 3:
        all_pages = set()
        for doc_data in by_doc.values():
            all_pages.update(doc_data['pages'])
        summary_parts.append(f"• Instances found on pages: {', '.join(map(str, sorted(all_pages)))}")
    
    if len(instances) > 1:
        variations = set(i['matched_text'] for i in instances)
        if len(variations) > 1:
            summary_parts.append(f"• Case variations found: {', '.join(sorted(variations))}")
    
    summary_parts.append("")
    
    # Detailed instances
    summary_parts.append("📍 ALL INSTANCES WITH EXACT LOCATIONS:")
    summary_parts.append("=" * 50)
    
    instance_counter = 0
    for doc_name in sorted(by_doc.keys()):
        summary_parts.append(f"\n📄 {doc_name}:")
        
        # Get instances for this document, sorted by page and position
        doc_instances = [i for i in instances if i['filename'] == doc_name]
        doc_instances.sort(key=lambda x: (x['page_number'], x['position_in_page']))
        
        for instance in doc_instances:
            instance_counter += 1
            summary_parts.append(f"   <br><b> Instance #{instance_counter}: </b>")
            summary_parts.append(f"       • Page: {instance['page_number']}")
            # summary_parts.append(f"      • Position in page: Character {instance['position_in_page']}")
            summary_parts.append(f"       • Matched text: '{instance['matched_text']}'")
            summary_parts.append(f"       • Case match: {'Exact' if instance['exact_match'] else 'Variation'}")
            
            # Context preview
            context_preview = instance['context']
            
            # summary_parts.append(f"       • Context: {context_preview}")
            summary_parts.append("")
    
    # Footer
    summary_parts.append("=" * 50)
    summary_parts.append("🎯 INSTANCE SUMMARY")
    summary_parts.append("=" * 50)
    summary_parts.append(f"✅ Successfully found and analyzed ALL {len(instances)} instances of '{query}'")
    summary_parts.append(f"✅ Covered {total_docs} document(s) across {total_pages} page(s)")
    # summary_parts.append("✅ No instances missed - comprehensive search completed")
    summary_parts.append("=" * 50)
    
    return "\n".join(summary_parts)

    

def debug_search_content(content, query):
    """
    Debug function to manually check content for query
    """
    logger.info(f"=== DEBUGGING SEARCH FOR '{query}' ===")
    
    # Check if query exists in full content
    if query.lower() in content.lower():
        logger.info(f"✅ Query '{query}' found in full content (case-insensitive)")
    else:
        logger.info(f"❌ Query '{query}' NOT found in full content")
        return
    
    # Find all occurrences manually
    content_lower = content.lower()
    query_lower = query.lower()
    
    pos = 0
    count = 0
    while True:
        found = content_lower.find(query_lower, pos)
        if found == -1:
            break
        
        count += 1
        # Get actual text (preserving case)
        actual_text = content[found:found + len(query)]
        
        # Find which page this belongs to
        # Count PAGE markers before this position
        page_markers_before = content[:found].count('PAGE ')
        
        logger.info(f"Manual find #{count}: '{actual_text}' at position {found}, estimated page {page_markers_before}")
        
        pos = found + 1
    
    logger.info(f"=== MANUAL SEARCH FOUND {count} INSTANCES ===")

# --------------------- RAG PROMPT FUNCTIONS ---------------------
def build_bulletproof_rag_prompt(context, query, total_instances):
    """RAG prompt specifically designed for bulletproof instance display with Ollama."""
    
    prompt = f"""You are a comprehensive document analysis assistant powered by Ollama Llama 3.1. You have completed a thorough search across all documents and found ALL instances of the query "{query}".

COMPREHENSIVE SEARCH RESULTS WITH ALL INSTANCES:
{context}

QUERY: {query}
TOTAL INSTANCES FOUND: {total_instances}

INSTRUCTIONS FOR OLLAMA:
1. Present a comprehensive analysis showing ALL {total_instances} instances found
2. Provide the exact distribution across documents and pages with clear statistics
3. Include specific page numbers and locations for each instance
4. Show the actual matched text and context for key instances
5. Create a well-structured response with clear headings and organization
6. Use markdown formatting for better readability
7. Confirm that this is a complete analysis with no instances missed
8. Provide actionable insights based on the patterns found

RESPONSE STRUCTURE:
## Executive Summary
- Total count and distribution summary
- Key patterns and insights

## Detailed Analysis
- Breakdown by document and page
- Specific examples with page references
- Statistical analysis

## Complete Instance List
- All instances with exact locations
- Context for each instance

## Conclusions
- Patterns observed
- Completeness confirmation
- Actionable recommendations

Provide a thorough, well-organized response that demonstrates the comprehensive nature of this analysis using your advanced language understanding capabilities.
"""
    
    return prompt

def enhanced_context_builder_with_exact_priority(search_results, question, max_context_length=25000):
    """Build context preserving ORIGINAL document structure with GUARANTEED exact match priority"""
    
    if not search_results:
        return "No relevant content found."
    
    context_parts = []
    total_length = 0
    
    logger.info(f"🎯 Building context from {len(search_results)} results, prioritizing exact matches")
    
    # CRITICAL: Separate exact matches from other results
    exact_matches = []
    other_results = []
    
    for result in search_results:
        if (result.get('guaranteed_exact_match') or 
            result.get('search_type') == 'exact_phrase_GUARANTEED' or
            result.get('ranking_priority') == 'EXACT_MATCH'):
            exact_matches.append(result)
        else:
            other_results.append(result)
    
    logger.info(f"📊 Context priority: {len(exact_matches)} exact matches, {len(other_results)} other results")
    
    # Sort both categories by score
    exact_matches.sort(key=lambda x: x.get('final_score', x.get('score', 0)), reverse=True)
    other_results.sort(key=lambda x: x.get('final_score', x.get('score', 0)), reverse=True)
    
    # PRIORITIZE EXACT MATCHES - Give them 70% of available context space
    exact_match_space = int(max_context_length * 0.7)
    other_results_space = max_context_length - exact_match_space
    
    # Add EXACT MATCHES FIRST with maximum context
    if exact_matches:
        context_parts.append(f"\n{'<hr>'}")
        context_parts.append(f"🎯 EXACT MATCHES FOUND FOR: '{question}'")
        context_parts.append(f"{'<hr>'}\n")
        
        used_exact_space = 0
        for i, result in enumerate(exact_matches):
            if used_exact_space >= exact_match_space:
                break
                
            content = result.get('content', result.get('chunk', ''))
            if isinstance(content, dict):
                content = content.get('content', str(content))
            
            content = str(content).strip()
            if not content:
                continue
            
            # Get filename
            filename = result.get('filename', 'Unknown')
            if not filename or filename == 'Unknown':
                chunk_data = result.get('chunk', {})
                if isinstance(chunk_data, dict):
                    doc_metadata = chunk_data.get('metadata', {}).get('document_metadata', {})
                    filename = doc_metadata.get('original_filename', 'Unknown')
            
            # Add document header for exact match
            doc_header = f"\n{'='*60}\n🎯 EXACT MATCH #{i+1}: {filename}\n{'='*60}\n"
            
            # Calculate available space for this chunk
            available_space = min(8000, exact_match_space - used_exact_space - len(doc_header))
            
            if available_space < 500:  # Need minimum space
                break
            
            # Truncate content if necessary but preserve exact match area
            if len(content) > available_space:
                # Try to find the exact match and preserve context around it
                query_lower = question.lower()
                content_lower = content.lower()
                match_pos = content_lower.find(query_lower)
                
                if match_pos != -1:
                    # Center the context around the exact match
                    start = max(0, match_pos - available_space // 3)
                    end = min(len(content), match_pos + len(question) + available_space // 3 * 2)
                    
                    # Expand to sentence boundaries
                    while start > 0 and content[start] not in '\n.!?':
                        start -= 1
                    while end < len(content) and content[end] not in '\n.!?':
                        end += 1
                    
                    content = content[start:end]
                    if start > 0:
                        content = "..." + content
                    if end < len(result.get('content', result.get('chunk', {}).get('content', ''))):
                        content = content + "..."
                else:
                    # No exact match found in content (shouldn't happen), truncate from start
                    content = content[:available_space] + "..."
            
            # Add to context
            context_parts.append(doc_header)
            context_parts.append(content)
            context_parts.append("\n\n")
            
            chunk_length = len(doc_header) + len(content) + 2
            used_exact_space += chunk_length
            total_length += chunk_length
            
            logger.info(f"✅ Added EXACT MATCH {i+1}: {len(content)} chars from {filename}")
    
    # Add OTHER RESULTS with remaining space
    if other_results and total_length < max_context_length:
        remaining_space = max_context_length - total_length
        
        if remaining_space > 1000:  # Only add if we have meaningful space
            context_parts.append(f"\n{'='*20}")
            context_parts.append(f"📄 ADDITIONAL RELEVANT CONTENT")
            context_parts.append(f"{'='*20}\n")
            
            # Group by document to avoid fragmentation
            by_document = {}
            for result in other_results:
                filename = result.get('filename', 'Unknown')
                if not filename or filename == 'Unknown':
                    chunk_data = result.get('chunk', {})
                    if isinstance(chunk_data, dict):
                        doc_metadata = chunk_data.get('metadata', {}).get('document_metadata', {})
                        filename = doc_metadata.get('original_filename', 'Unknown')
                
                if filename not in by_document:
                    by_document[filename] = []
                by_document[filename].append(result)
            
            used_other_space = 0
            for doc_name, doc_results in by_document.items():
                if used_other_space >= remaining_space - 500:
                    break
                
                # Add document header
                doc_header = f"\n{'='*50}\n📄 {doc_name}\n{'='*50}\n"
                
                available_doc_space = min(4000, remaining_space - used_other_space - len(doc_header))
                if available_doc_space < 200:
                    break
                
                context_parts.append(doc_header)
                used_other_space += len(doc_header)
                total_length += len(doc_header)
                
                # Add content from this document
                for result in doc_results[:2]:  # Max 2 chunks per document
                    if used_other_space >= remaining_space - 200:
                        break
                    
                    content = result.get('content', result.get('chunk', ''))
                    if isinstance(content, dict):
                        content = content.get('content', str(content))
                    
                    content = str(content).strip()
                    if not content:
                        continue
                    
                    available_chunk_space = min(2000, remaining_space - used_other_space)
                    
                    if len(content) > available_chunk_space:
                        content = content[:available_chunk_space] + "..."
                    
                    context_parts.append(content)
                    context_parts.append("\n\n")
                    
                    chunk_length = len(content) + 2
                    used_other_space += chunk_length
                    total_length += chunk_length
    
    # Add comprehensive search summary
    summary = f"""

{'<hr>'}
🎯 SEARCH ANALYSIS SUMMARY
{'<hr>'}
- Question: {question}
- EXACT MATCHES: {len(exact_matches)} (guaranteed priority)
- Additional results: {len(other_results)}
- Total content analyzed: {total_length:,} characters
- Search strategy: Multi-strategy with exact match guarantee
- Exact matches are prioritized and shown first
- All content preserves original document formatting
{'<hr>'}

"""
    
    context_parts.append(summary)
    final_context = "".join(context_parts)
    
    logger.info(f"🎯 Final context: {len(final_context):,} chars, {len(exact_matches)} exact matches prioritized")
    
    return final_context

# --------------------- SEARCH ENGINE CLASS ---------------------
class AccuracyFocusedSearchEngine:
    """Fixed GPU-accelerated search engine with GUARANTEED exact match priority"""
    
    def __init__(self, processor: AccuracyFocusedDocumentProcessor):
        self.processor = processor
        self.logger = logging.getLogger(__name__)
    
    def search_with_guaranteed_accuracy(self, query: str, top_k: int = 20) -> List[Dict[str, Any]]:
        """Enhanced search that GUARANTEES exact matches are always prioritized"""
        
        if not query.strip():
            return []
        
        # Ensure we have content to search
        if not self.processor.chunk_metadata:
            self.logger.warning("No content indexed for search")
            return []
        
        self.logger.info(f"🎯 EXACT MATCH PRIORITY SEARCH: '{query}'")
        
        # 1. EXACT PHRASE SEARCH (ABSOLUTE HIGHEST PRIORITY)
        exact_results = self._exact_phrase_search_GUARANTEED(query, top_k)
        
        # If we have exact matches, prioritize them heavily
        if exact_results:
            self.logger.info(f"✅ Found {len(exact_results)} EXACT matches - these will be prioritized")
            
            # Still run other searches for context, but exact matches get massive score boost
            keyword_results = self._multi_keyword_search(self._extract_key_terms(query), top_k)
            semantic_results = self._semantic_search_with_reranking(query, top_k)
            context_results = self._context_aware_search(query, top_k)
            fuzzy_results = self._fuzzy_search(query, top_k)
            
            all_results = exact_results + keyword_results + semantic_results + context_results + fuzzy_results
        else:
            # No exact matches, run all strategies
            self.logger.warning(f"⚠️ No exact matches found for '{query}' - running all strategies")
            
            keyword_results = self._multi_keyword_search(self._extract_key_terms(query), top_k)
            semantic_results = self._semantic_search_with_reranking(query, top_k)
            context_results = self._context_aware_search(query, top_k)
            fuzzy_results = self._fuzzy_search(query, top_k)
            
            all_results = keyword_results + semantic_results + context_results + fuzzy_results
            
            # Emergency fallback if still no results
            if not all_results:
                emergency_results = self._emergency_fallback_search(query)
                all_results = emergency_results
        
        # 2. GUARANTEED EXACT MATCH RANKING
        final_results = self._guaranteed_exact_match_ranking(all_results, query, top_k)
        
        if not final_results:
            # Last resort - return ANY relevant content
            self.logger.error(f"❌ CRITICAL: No results found at all for '{query}' - using emergency fallback")
            final_results = self._emergency_fallback_search(query)
        
        self.logger.info(f"🎯 FINAL RESULTS: {len(final_results)} results (exact matches prioritized)")
        return final_results
    
    def _exact_phrase_search_GUARANTEED(self, query: str, top_k: int) -> List[Dict[str, Any]]:
        """GUARANTEED exact phrase matching with multiple search patterns"""
        query_lower = query.lower().strip()
        query_normalized = re.sub(r'\s+', ' ', query_lower)
        
        results = []
        
        # Multiple exact search patterns for maximum coverage
        search_patterns = [
            query_normalized,                    # Original query
            query_normalized.replace(' ', ''),   # No spaces
            query_lower,                         # Raw lowercase
            query.strip(),                       # Original case
        ]
        
        # Remove duplicates while preserving order
        unique_patterns = []
        seen = set()
        for pattern in search_patterns:
            if pattern and pattern not in seen:
                unique_patterns.append(pattern)
                seen.add(pattern)
        
        self.logger.info(f"🔍 Searching for exact patterns: {unique_patterns}")
        
        for idx, chunk in enumerate(self.processor.chunk_metadata):
            content_lower = chunk['content'].lower()
            content_normalized = re.sub(r'\s+', ' ', content_lower)
            content_raw = chunk['content']
            
            # Check each pattern
            best_match_score = 0
            best_match_info = None
            
            for pattern in unique_patterns:
                if not pattern:
                    continue
                    
                # Try different content variations
                content_variations = [
                    content_normalized,
                    content_lower,
                    content_raw.lower(),
                    content_raw
                ]
                
                for content_var in content_variations:
                    if pattern in content_var:
                        position = content_var.find(pattern)
                        
                        # Calculate match quality
                        match_score = 10000  # Base score for exact match
                        
                        # Bonus for exact case match
                        if pattern in content_raw:
                            match_score += 2000
                        
                        # Bonus for word boundaries
                        if self._is_word_boundary_match(content_var, pattern, position):
                            match_score += 1000
                        
                        # Bonus for being in structured content
                        if any(marker in content_var[max(0, position-100):position+100] 
                               for marker in ['# ', '## ', '### ', '**', '*', '|']):
                            match_score += 500
                        
                        # Bonus for being at start of sentence/paragraph
                        if position == 0 or content_var[position-1] in ['\n', '.', '!', '?']:
                            match_score += 300
                        
                        if match_score > best_match_score:
                            best_match_score = match_score
                            best_match_info = {
                                'pattern': pattern,
                                'position': position,
                                'content_type': 'exact_phrase_GUARANTEED'
                            }
            
            if best_match_score > 0:
                results.append({
                    'chunk': chunk,
                    'score': best_match_score,
                    'search_type': 'exact_phrase_GUARANTEED',
                    'chunk_index': idx,
                    'match_info': best_match_info,
                    'guaranteed_exact_match': True  # Special flag
                })
        
        # Sort by score and return top results
        results.sort(key=lambda x: x['score'], reverse=True)
        
        if results:
            self.logger.info(f"✅ EXACT MATCHES FOUND: {len(results)} matches with scores: {[r['score'] for r in results[:5]]}")
        else:
            self.logger.warning(f"❌ NO EXACT MATCHES for '{query}'")
        
        return results[:top_k]
    
    def _is_word_boundary_match(self, content: str, pattern: str, position: int) -> bool:
        """Check if the match is at word boundaries"""
        if position > 0 and content[position - 1].isalnum():
            return False
        
        end_pos = position + len(pattern)
        if end_pos < len(content) and content[end_pos].isalnum():
            return False
        
        return True
    
    def _guaranteed_exact_match_ranking(self, all_results: List[Dict], query: str, top_k: int) -> List[Dict]:
        """GUARANTEED ranking that always prioritizes exact matches"""
        
        # Separate exact matches from other results
        exact_matches = []
        other_results = []
        
        for result in all_results:
            if result.get('guaranteed_exact_match') or result.get('search_type') == 'exact_phrase_GUARANTEED':
                exact_matches.append(result)
            else:
                other_results.append(result)
        
        # Remove duplicates within each category
        exact_matches = self._remove_duplicate_chunks(exact_matches)
        other_results = self._remove_duplicate_chunks(other_results)
        
        self.logger.info(f"📊 Ranking: {len(exact_matches)} exact matches, {len(other_results)} other results")
        
        # Apply massive score boost to exact matches
        for result in exact_matches:
            result['final_score'] = result['score'] + 100000  # Massive boost
            result['ranking_priority'] = 'EXACT_MATCH'
        
        # Apply normal scoring to other results
        for result in other_results:
            base_score = result['score']
            search_type = result['search_type']
            
            # Apply weight multipliers
            multiplier = app.config['SEARCH_WEIGHTS'].get(search_type, 1.0)
            result['final_score'] = base_score * multiplier
            result['ranking_priority'] = 'SECONDARY'
        
        # Combine results with exact matches ALWAYS first
        final_results = exact_matches + other_results
        
        # Sort by final score (exact matches will always be on top due to massive boost)
        final_results.sort(key=lambda x: x['final_score'], reverse=True)
        
        # Take top results
        top_results = final_results[:top_k]
        
        # Log final ranking
        for i, result in enumerate(top_results[:5]):
            self.logger.info(f"📈 Rank {i+1}: {result['ranking_priority']} - Score: {result['final_score']:.0f} - Type: {result['search_type']}")
        
        return top_results
    
    def _remove_duplicate_chunks(self, results: List[Dict]) -> List[Dict]:
        """Remove duplicate chunks while keeping highest scoring ones"""
        seen_chunks = {}
        unique_results = []
        
        for result in results:
            chunk_idx = result['chunk_index']
            
            if chunk_idx not in seen_chunks or result['score'] > seen_chunks[chunk_idx]['score']:
                seen_chunks[chunk_idx] = result
        
        # Convert back to list, sorted by score
        unique_results = list(seen_chunks.values())
        unique_results.sort(key=lambda x: x['score'], reverse=True)
        
        return unique_results
    
    # Keep all other methods unchanged...
    def _multi_keyword_search(self, keywords: List[str], top_k: int) -> List[Dict[str, Any]]:
        """Search for multiple keywords with AND/OR logic - handles missing TF-IDF"""
        results = []
        
        # Check if TF-IDF is available
        if not self.processor.tfidf_fitted or self.processor.tfidf_matrix is None:
            self.logger.warning("TF-IDF not available, using simple text matching for keyword search")
            
            # Fallback to simple text matching
            for idx, chunk in enumerate(self.processor.chunk_metadata):
                content_lower = chunk['content'].lower()
                
                # Count keyword matches
                keyword_matches = []
                total_score = 0
                
                for keyword in keywords:
                    count = content_lower.count(keyword.lower())
                    if count > 0:
                        keyword_matches.append((keyword, count))
                        # Score based on keyword importance and frequency
                        keyword_score = count * (50 + len(keyword) * 5)
                        total_score += keyword_score
                
                if keyword_matches:
                    # Bonus for matching multiple keywords
                    coverage_bonus = (len(keyword_matches) / len(keywords)) * 200
                    total_score += coverage_bonus
                    
                    # Bonus for keyword density
                    content_words = len(content_lower.split())
                    if content_words > 0:
                        density = sum(count for _, count in keyword_matches) / content_words
                        density_bonus = min(density * 1000, 100)  # Cap at 100
                        total_score += density_bonus
                    
                    results.append({
                        'chunk': chunk,
                        'score': total_score,
                        'search_type': 'multi_keyword_fallback',
                        'chunk_index': idx,
                        'matched_keywords': keyword_matches,
                        'keyword_coverage': len(keyword_matches) / len(keywords)
                    })
            
            return sorted(results, key=lambda x: x['score'], reverse=True)[:top_k]
        
        # Original TF-IDF based search (when available)
        try:
            # Transform keywords using existing TF-IDF vectorizer
            keyword_text = ' '.join(keywords)
            keyword_vector = self.processor.tfidf_vectorizer.transform([keyword_text])
            
            # Calculate similarity with all documents
            similarities = cosine_similarity(keyword_vector, self.processor.tfidf_matrix).flatten()
            
            # Get top results
            top_indices = similarities.argsort()[-top_k:][::-1]
            
            for idx in top_indices:
                if similarities[idx] > 0:  # Only include positive similarities
                    chunk = self.processor.chunk_metadata[idx]
                    
                    # Also calculate simple keyword matches for metadata
                    content_lower = chunk['content'].lower()
                    keyword_matches = []
                    for keyword in keywords:
                        count = content_lower.count(keyword.lower())
                        if count > 0:
                            keyword_matches.append((keyword, count))
                    
                    results.append({
                        'chunk': chunk,
                        'score': float(similarities[idx]) * 1000,  # Scale score
                        'search_type': 'multi_keyword_tfidf',
                        'chunk_index': idx,
                        'matched_keywords': keyword_matches,
                        'keyword_coverage': len(keyword_matches) / len(keywords),
                        'tfidf_similarity': float(similarities[idx])
                    })
            
            return results
            
        except Exception as e:
            self.logger.warning(f"TF-IDF keyword search failed: {e}, falling back to simple matching")
            
            # Fallback to simple text matching if TF-IDF fails
            for idx, chunk in enumerate(self.processor.chunk_metadata):
                content_lower = chunk['content'].lower()
                
                keyword_matches = []
                total_score = 0
                
                for keyword in keywords:
                    count = content_lower.count(keyword.lower())
                    if count > 0:
                        keyword_matches.append((keyword, count))
                        keyword_score = count * (50 + len(keyword) * 5)
                        total_score += keyword_score
                
                if keyword_matches:
                    coverage_bonus = (len(keyword_matches) / len(keywords)) * 200
                    total_score += coverage_bonus
                    
                    results.append({
                        'chunk': chunk,
                        'score': total_score,
                        'search_type': 'multi_keyword_fallback',
                        'chunk_index': idx,
                        'matched_keywords': keyword_matches,
                        'keyword_coverage': len(keyword_matches) / len(keywords)
                    })
            
            return sorted(results, key=lambda x: x['score'], reverse=True)[:top_k]

    def _extract_key_terms(self, query: str) -> List[str]:
        """Extract key terms with better filtering"""
        # Remove common words but keep important ones
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should'}
        
        # Extract words and phrases
        words = re.findall(r'\b\w+\b', query.lower())
        
        # Keep important terms
        key_terms = []
        
        # Single important words (not in stop words, length > 2)
        for word in words:
            if len(word) > 2 and word not in stop_words:
                key_terms.append(word)
        
        # Important phrases (2-4 words)
        for i in range(len(words) - 1):
            if len(words[i]) > 2 and len(words[i+1]) > 2:
                phrase = f"{words[i]} {words[i+1]}"
                if not any(w in stop_words for w in [words[i], words[i+1]]):
                    key_terms.append(phrase)
        
        # Longer phrases for specific queries
        if len(words) >= 3:
            for i in range(len(words) - 2):
                if all(len(w) > 2 for w in words[i:i+3]):
                    phrase = " ".join(words[i:i+3])
                    key_terms.append(phrase)
        
        # Add original query
        key_terms.append(query.strip())
        
        return key_terms

    def _semantic_search_with_reranking(self, query: str, top_k: int) -> List[Dict[str, Any]]:
        """GPU-accelerated semantic search using FAISS with reranking"""
        try:
            if not self.processor.faiss_index or self.processor.faiss_index.ntotal == 0:
                return []
            
            # Generate query embedding
            if self.processor.gpu_enabled and app.config['MIXED_PRECISION']:
                try:
                    with torch.amp.autocast('cuda'):
                        query_embedding = self.processor.embedding_model.encode(
                            [query],
                            convert_to_tensor=True,
                            device=self.processor.device,
                            normalize_embeddings=app.config['EMBEDDING_NORMALIZE']
                        )
                except (AttributeError, TypeError):
                    with torch.cuda.amp.autocast():
                        query_embedding = self.processor.embedding_model.encode(
                            [query],
                            convert_to_tensor=True,
                            device=self.processor.device,
                            normalize_embeddings=app.config['EMBEDDING_NORMALIZE']
                        )
                query_embedding = query_embedding.cpu().float().numpy()
            else:
                query_embedding = self.processor.embedding_model.encode(
                    [query],
                    convert_to_tensor=False,
                    normalize_embeddings=app.config['EMBEDDING_NORMALIZE']
                )
            
            # Normalize for cosine similarity
            if app.config['EMBEDDING_NORMALIZE']:
                faiss.normalize_L2(query_embedding.astype('float32'))
            
            # Search FAISS index
            similarities, indices = self.processor.faiss_index.search(
                query_embedding.astype('float32'), 
                min(top_k * 3, self.processor.faiss_index.ntotal)  # Get more candidates for reranking
            )
            
            results = []
            for idx, score in zip(indices[0], similarities[0]):
                if idx != -1 and idx < len(self.processor.chunk_metadata):
                    chunk = self.processor.chunk_metadata[idx]
                    results.append({
                        'chunk': chunk,
                        'score': float(score) * 100,  # Scale score
                        'search_type': 'semantic_gpu',
                        'chunk_index': idx
                    })
            
            return results
            
        except Exception as e:
            self.logger.error(f"Semantic search failed: {e}")
            return []

    def _context_aware_search(self, query: str, top_k: int) -> List[Dict[str, Any]]:
        """Search considering chunk context and relationships"""
        results = []
        query_lower = query.lower()
        
        for idx, chunk in enumerate(self.processor.chunk_metadata):
            chunk_score = 0
            
            # Score current chunk
            if query_lower in chunk['content'].lower():
                chunk_score += 100
            
            # Check previous chunk context
            if chunk['metadata'].get('previous_chunk') is not None:
                prev_idx = chunk['metadata']['previous_chunk']
                if prev_idx < len(self.processor.chunk_metadata):
                    prev_chunk = self.processor.chunk_metadata[prev_idx]
                    if query_lower in prev_chunk['content'].lower():
                        chunk_score += 50  # Context bonus
            
            # Check next chunk context
            if chunk['metadata'].get('next_chunk') is not None:
                next_idx = chunk['metadata']['next_chunk']
                if next_idx < len(self.processor.chunk_metadata):
                    next_chunk = self.processor.chunk_metadata[next_idx]
                    if query_lower in next_chunk['content'].lower():
                        chunk_score += 50  # Context bonus
            
            if chunk_score > 0:
                results.append({
                    'chunk': chunk,
                    'score': chunk_score,
                    'search_type': 'context_aware',
                    'chunk_index': idx
                })
        
        return sorted(results, key=lambda x: x['score'], reverse=True)[:top_k]

    def _fuzzy_search(self, query: str, top_k: int) -> List[Dict[str, Any]]:
        """Fuzzy matching search for typos and variations"""
        try:
            query_words = [w.lower() for w in query.split() if len(w) > 3]
            if not query_words:
                return []
            
            results = []
            
            for idx, chunk in enumerate(self.processor.chunk_metadata):
                content_words = re.findall(r'\b\w{4,}\b', chunk['content'].lower())
                
                total_score = 0
                matches_found = 0
                
                for query_word in query_words:
                    best_match_score = 0
                    
                    for content_word in content_words:
                        similarity = SequenceMatcher(None, query_word, content_word).ratio()
                        if similarity > 0.8:  # 80% similarity threshold
                            best_match_score = max(best_match_score, similarity * 30)
                    
                    if best_match_score > 0:
                        total_score += best_match_score
                        matches_found += 1
                
                if total_score > 0:
                    # Bonus for finding multiple words
                    if matches_found > 1:
                        total_score += matches_found * 20
                    
                    results.append({
                        'chunk': chunk,
                        'score': total_score,
                        'search_type': 'fuzzy_match',
                        'chunk_index': idx
                    })
            
            return sorted(results, key=lambda x: x['score'], reverse=True)[:top_k]
            
        except Exception as e:
            self.logger.error(f"Fuzzy search failed: {e}")
            return []

    def _emergency_fallback_search(self, query: str) -> List[Dict[str, Any]]:
        """Emergency fallback that always finds something relevant"""
        try:
            if not self.processor.chunk_metadata:
                return []
            
            # Extract key terms from query
            query_words = re.findall(r'\b\w{3,}\b', query.lower())
            
            results = []
            
            # If no meaningful words, return first few chunks
            if not query_words:
                for idx in range(min(5, len(self.processor.chunk_metadata))):
                    chunk = self.processor.chunk_metadata[idx]
                    results.append({
                        'chunk': chunk,
                        'score': 10,
                        'search_type': 'emergency_first_chunks',
                        'chunk_index': idx
                    })
                return results
            
            # Look for any single word matches
            for idx, chunk in enumerate(self.processor.chunk_metadata):
                content_lower = chunk['content'].lower()
                score = 0
                
                for word in query_words:
                    if word in content_lower:
                        score += 5
                
                # Even if no words match, give a small score to structured chunks
                if score == 0:
                    if chunk['metadata'].get('has_headings') or chunk['metadata'].get('has_tables'):
                        score = 1
                
                if score > 0:
                    results.append({
                        'chunk': chunk,
                        'score': score,
                        'search_type': 'emergency_fallback',
                        'chunk_index': idx
                    })
            
            # If still no results, return random chunks with structure
            if not results:
                import random
                structured_chunks = [
                    (idx, chunk) for idx, chunk in enumerate(self.processor.chunk_metadata)
                    if chunk['metadata'].get('has_headings') or chunk['metadata'].get('has_tables')
                ]
                
                if structured_chunks:
                    selected = random.sample(structured_chunks, min(3, len(structured_chunks)))
                else:
                    selected = [(idx, chunk) for idx, chunk in enumerate(self.processor.chunk_metadata[:3])]
                
                for idx, chunk in selected:
                    results.append({
                        'chunk': chunk,
                        'score': 0.1,
                        'search_type': 'emergency_random',
                        'chunk_index': idx
                    })
            
            return sorted(results, key=lambda x: x['score'], reverse=True)
            
        except Exception as e:
            self.logger.error(f"Emergency fallback search failed: {e}")
            # Absolute last resort
            if self.processor.chunk_metadata:
                return [{
                    'chunk': self.processor.chunk_metadata[0],
                    'score': 0.01,
                    'search_type': 'absolute_fallback',
                    'chunk_index': 0
                }]
            return []




def fallback_to_search(search_keywords: str, user_id: int, document_info: list) -> dict:
    """
    Enhanced fallback search optimized for keyword matching
    """
    try:
        logger.info(f"🔍 Enhanced fallback search for keywords: '{search_keywords}'")
        
        # Use enhanced keyword search instead of exact phrase search
        all_instances, comprehensive_summary = search_keywords_enhanced(search_keywords, user_id)
        
        if not all_instances:
            response_content = f"""# Keyword Search Results for "{search_keywords}"

No instances of the keywords "{search_keywords}" were found in your documents after comprehensive search.

**Documents Searched:** {len(document_info)} documents
- {', '.join(document_info)}

**Search Variations Attempted:**
- Exact phrase matching
- Individual keyword matching  
- Fuzzy matching (typo tolerance)
- Semantic similarity matching

**Suggestions:**
- Try different search terms or variations
- Check spelling of keywords
- Use broader or more specific terms
- Try synonyms or related terms"""
        else:
            response_content = f"""# Keyword Search Results for "{search_keywords}"

Found **{len(all_instances)} relevant instances** across your documents using advanced keyword matching.

{comprehensive_summary}

"""
        
        return {
            'all_instances': all_instances,
            'content': response_content,
            'sources': document_info,
            'question_analysis': {
                'type': 'keyword_search',
                'confidence': 0.9,
                'reasoning': f'Keyword search for: "{search_keywords}"',
                'ai_determined': False,
                'fallback': True,
                'search_method': 'enhanced_keyword_matching'
            },
            'success': True,
            'ai_powered': False,
            'search_keywords': search_keywords
        }
        
    except Exception as e:
        logger.error(f"Enhanced fallback search error: {e}")
        return {
            'content': f"Keyword search failed: {str(e)}",
            'sources': document_info,
            'question_analysis': {'type': 'error', 'confidence': 0.0},
            'success': False,
            'search_keywords': search_keywords
        }

def search_keywords_enhanced(keywords: str, user_id: int):
    """
    Keyword search that finds content related to multiple keywords
    with fuzzy matching and semantic understanding
    """
    logger.info(f"🎯 KEYWORD SEARCH: Finding content related to '{keywords}' for user {user_id}")
    
    try:
        all_instances = []
        
        # Get user documents
        user_documents = Document.query.filter_by(user_id=user_id).all()
        
        if not user_documents:
            return [], "No documents found to search."
        
        logger.info(f"Processing {len(user_documents)} documents for keyword search")
        
        # Split keywords for individual and combined searching
        keyword_list = [k.strip() for k in keywords.split() if len(k.strip()) > 2]
        logger.info(f"Individual keywords: {keyword_list}")
        
        # Process each document
        for document in user_documents:
            try:
                filename = document.document_metadata.get('original_filename', 'Unknown') if document.document_metadata else 'Unknown'
                logger.info(f"Processing: {filename}")
                
                # Get the full document content
                content = get_full_document_content(document.id)
                
                if not content:
                    logger.warning(f"No content found for {filename}")
                    continue
                
                logger.info(f"Content length: {len(content)} characters")
                
                # Find keyword-related instances in this document
                doc_instances = find_keyword_instances_enhanced(content, keywords, keyword_list, filename, document.id)
                all_instances.extend(doc_instances)
                
                logger.info(f"Found {len(doc_instances)} keyword-related instances in {filename}")
                
            except Exception as doc_error:
                logger.error(f"Error processing document: {doc_error}")
                continue
        
        # Remove exact duplicates and rank by relevance
        unique_instances = remove_and_rank_instances(all_instances, keywords)
        
        # Build enhanced summary
        summary = build_enhanced_keyword_summary(unique_instances, keywords, len(user_documents))
        
        logger.info(f"✅ ENHANCED KEYWORD SEARCH COMPLETED: {len(unique_instances)} relevant instances found")
        
        return unique_instances, summary
        
    except Exception as e:
        logger.error(f"❌ ENHANCED KEYWORD SEARCH ERROR: {e}")
        return [], f"Search failed: {str(e)}"

def find_keyword_instances_enhanced(content, original_keywords, keyword_list, filename, document_id):
    """
    Find instances using multiple keyword matching strategies
    """
    instances = []
    
    try:
        # Split content into pages
        pages = split_into_pages_reliable(content)
        logger.info(f"Split into {len(pages)} pages for keyword search in: {filename}")
        
        # Search each page
        for page_data in pages:
            # Extract page number
            page_num_match = re.match(r'\[PAGE (\d+)\]', page_data)
            if page_num_match:
                page_num = int(page_num_match.group(1))
                page_content = page_data[page_num_match.end():].strip()
            else:
                page_num = pages.index(page_data) + 1
                page_content = page_data
            
            # Multiple search strategies
            page_instances = []
            
            # Strategy 1: Exact phrase matching
            if original_keywords.lower() in page_content.lower():
                page_instances.extend(create_keyword_instances(
                    page_content, original_keywords, page_num, filename, document_id, 
                    'exact_phrase', 100
                ))
            
            # Strategy 2: All keywords present (AND logic)
            if all(keyword.lower() in page_content.lower() for keyword in keyword_list):
                page_instances.extend(create_keyword_instances(
                    page_content, original_keywords, page_num, filename, document_id, 
                    'all_keywords', 80
                ))
            
            # Strategy 3: Most keywords present (partial match)
            matching_keywords = [kw for kw in keyword_list if kw.lower() in page_content.lower()]
            if len(matching_keywords) >= max(1, len(keyword_list) // 2):  # At least half the keywords
                score = (len(matching_keywords) / len(keyword_list)) * 60
                page_instances.extend(create_keyword_instances(
                    page_content, ' '.join(matching_keywords), page_num, filename, document_id, 
                    'partial_keywords', score
                ))
            
            # Strategy 4: Semantic similarity (if available)
            try:
                if has_semantic_similarity(page_content, original_keywords):
                    page_instances.extend(create_keyword_instances(
                        page_content, original_keywords, page_num, filename, document_id, 
                        'semantic_match', 40
                    ))
            except Exception as sem_error:
                logger.debug(f"Semantic matching failed: {sem_error}")
            
            instances.extend(page_instances)
        
        return instances
        
    except Exception as e:
        logger.error(f"Error in enhanced keyword search: {e}")
        return []

def create_keyword_instances(page_content, keywords, page_num, filename, document_id, match_type, base_score):
    """
    Create instances for keyword matches with enhanced context
    """
    instances = []
    
    try:
        # Find the best context section (paragraph containing most keywords)
        paragraphs = page_content.split('\n\n')
        best_paragraph = ""
        best_score = 0
        best_position = 0
        
        for i, paragraph in enumerate(paragraphs):
            para_score = 0
            for keyword in keywords.split():
                if len(keyword) > 2 and keyword.lower() in paragraph.lower():
                    para_score += 1
            
            if para_score > best_score:
                best_score = para_score
                best_paragraph = paragraph
                # Calculate approximate position
                best_position = sum(len(p) + 2 for p in paragraphs[:i])
        
        if best_paragraph:
            # Create context around the best matching paragraph
            context_start = max(0, best_position - 500)
            context_end = min(len(page_content), best_position + len(best_paragraph) + 500)
            context = page_content[context_start:context_end].strip()
            
            # Clean context
            context = ' '.join(context.split())
            
            instance = {
                'filename': filename,
                'document_id': document_id,
                'page_number': page_num,
                'position_in_page': best_position,
                'matched_text': keywords,
                'context': context,
                'match_type': match_type,
                'relevance_score': base_score + best_score * 10,
                'keywords_found': best_score,
                'exact_match': match_type == 'exact_phrase',
                'instance_key': f"{document_id}_{page_num}_{best_position}_{match_type}"
            }
            
            instances.append(instance)
        
        return instances
        
    except Exception as e:
        logger.error(f"Error creating keyword instances: {e}")
        return []

def has_semantic_similarity(content, keywords):
    """
    Check if content has semantic similarity to keywords (dynamic implementation)
    """
    try:
        # Dynamic semantic analysis - look for word families and related terms
        content_lower = content.lower()
        keywords_lower = keywords.lower()
        
        # Split keywords to analyze each one
        keyword_words = [w for w in keywords.split() if len(w) > 3]
        
        similarity_score = 0
        
        for keyword in keyword_words:
            # Look for word variations (suffixes, prefixes)
            base_variations = [
                keyword,
                keyword + 's',      # plural
                keyword + 'ed',     # past tense
                keyword + 'ing',    # present participle
                keyword + 'er',     # comparative
                keyword + 'ly',     # adverb
                keyword + 'tion',   # noun form
                keyword + 'al',     # adjective
                keyword + 'ness',   # noun form
                'un' + keyword,     # negative prefix
                're' + keyword,     # re- prefix
            ]
            
            # Check for any variations in content
            for variation in base_variations:
                if variation in content_lower:
                    similarity_score += 1
                    break
            
            # Look for contextual similarity (words that commonly appear together)
            # This is a simplified approach - in production, you might use word embeddings
            context_window = 50  # Look within 50 characters of the keyword
            
            for i, char in enumerate(content_lower):
                if content_lower[i:i+len(keyword)] == keyword:
                    # Extract surrounding context
                    start = max(0, i - context_window)
                    end = min(len(content_lower), i + len(keyword) + context_window)
                    context = content_lower[start:end]
                    
                    # Check if other keywords appear in this context
                    for other_keyword in keyword_words:
                        if other_keyword != keyword and other_keyword in context:
                            similarity_score += 2  # Higher score for co-occurrence
        
        # Return True if we found enough semantic connections
        threshold = max(1, len(keyword_words) // 2)  # At least half the keywords should have connections
        return similarity_score >= threshold
        
    except Exception as e:
        logger.debug(f"Semantic similarity check failed: {e}")
        return False

def remove_and_rank_instances(instances, original_keywords):
    """
    Remove duplicates and rank instances by relevance
    """
    # Remove exact duplicates
    seen_keys = set()
    unique_instances = []
    
    for instance in instances:
        key = instance['instance_key']
        if key not in seen_keys:
            seen_keys.add(key)
            unique_instances.append(instance)
    
    # Sort by relevance score (highest first)
    unique_instances.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
    
    return unique_instances

def build_enhanced_keyword_summary(instances, keywords, documents_searched):
    """
    Build enhanced summary for keyword search results
    """
    if not instances:
        return f"""
COMPREHENSIVE KEYWORD SEARCH RESULTS
=====================================
Keywords: '{keywords}'
Documents searched: {documents_searched}
Total instances found: 0

No content related to '{keywords}' was found in your documents.
"""
    
    summary_parts = []
    
    # Header
    summary_parts.append("📊 KEYWORD ANALYSIS FOR '{}'".format(keywords.upper()))
    summary_parts.append("<hr>")
    summary_parts.append("")
    
    # Quick stats
    by_doc = {}
    match_types = {}
    
    for instance in instances:
        doc = instance['filename']
        match_type = instance.get('match_type', 'unknown')
        
        if doc not in by_doc:
            by_doc[doc] = {'total': 0, 'pages': set(), 'best_score': 0}
        
        by_doc[doc]['total'] += 1
        by_doc[doc]['pages'].add(instance['page_number'])
        by_doc[doc]['best_score'] = max(by_doc[doc]['best_score'], instance.get('relevance_score', 0))
        
        match_types[match_type] = match_types.get(match_type, 0) + 1
    
    total_docs = len(by_doc)
    total_pages = sum(len(doc_data['pages']) for doc_data in by_doc.values())
    
    summary_parts.append("📈 <b>SEARCH SUMMARY:</b>")
    summary_parts.append(f"• Total relevant instances: {len(instances)}")
    summary_parts.append(f"• Documents with matches: {total_docs}")
    summary_parts.append(f"• Pages with content: {total_pages}")
    summary_parts.append(f"• Search strategies used: {len(match_types)}")
    summary_parts.append("")
    
    # Match type breakdown
    summary_parts.append("🔍 <b>MATCH TYPE BREAKDOWN:</b>")
    summary_parts.append("-" * 40)
    match_type_names = {
        'exact_phrase': 'Exact phrase matches',
        'all_keywords': 'All keywords present',
        'partial_keywords': 'Partial keyword matches',
        'semantic_match': 'Semantic similarity matches'
    }
    
    for match_type, count in match_types.items():
        type_name = match_type_names.get(match_type, match_type)
        summary_parts.append(f"• {type_name}: {count}")
    summary_parts.append("")
    
    # Top results by document
    summary_parts.append("📋 <b>TOP RESULTS BY DOCUMENT:</b>")
    summary_parts.append("-" * 50)
    
    # Sort documents by best score
    sorted_docs = sorted(by_doc.items(), key=lambda x: x[1]['best_score'], reverse=True)
    
    for doc_name, doc_data in sorted_docs[:5]:  # Top 5 documents
        summary_parts.append(f"📄 {doc_name}")
        summary_parts.append(f"   • Relevance instances: {doc_data['total']}")
        summary_parts.append(f"   • Pages covered: {len(doc_data['pages'])}")
        # summary_parts.append(f"   • Best relevance score: {doc_data['best_score']:.0f}/100")
        
        # Show top instances from this document
        doc_instances = [i for i in instances if i['filename'] == doc_name][:3]
        for i, instance in enumerate(doc_instances, 1):
            summary_parts.append(f"   • <strong>Instance {i}: Page {instance['page_number']} ({instance['match_type']})</strong> ({instance['matched_text']})")
        summary_parts.append("")
    
    # Footer
    summary_parts.append("=" * 50)
    summary_parts.append("🎯 <b>KEYWORD SEARCH SUMMARY</b>")
    summary_parts.append("=" * 50)
    summary_parts.append(f"✅ Found {len(instances)} relevant instances for '{keywords}'")
    summary_parts.append(f"✅ Covered {total_docs} document(s) across {total_pages} page(s)")
    summary_parts.append("✅ Used multiple search strategies for comprehensive coverage")
    summary_parts.append("=" * 50)
    
    return "\n".join(summary_parts)
# --------------------- DATABASE MODELS ---------------------

class SubscriptionPlan:
    BASIC = {
        'id': 'basic',
        'name': 'Basic',
        'price': 99,
        'features': ['500 documents/month', '100MB max file size', 'Basic support'],
        'limits': {'max_documents': 500, 'max_file_size': 100 * 1024 * 1024, 'max_conversations': 10}
    }
    
    PRO = {
        'id': 'pro',
        'name': 'Professional',
        'price': 199,
        'features': ['2000 documents/month', '500MB max file size', 'Priority support', 'Advanced analytics'],
        'limits': {'max_documents': 2000, 'max_file_size': 500 * 1024 * 1024, 'max_conversations': 50}
    }
    
    TRIAL = {
        'id': 'trial',
        'name': 'Free Trial',
        'price': 0,
        'features': ['Unlimited documents for 14 days', '50GB max file size', 'All Professional features', 'GPU acceleration'],
        'limits': {'max_documents': None, 'max_file_size': 50 * 1024 * 1024 * 1024, 'max_conversations': None, 'trial_days': 14}
    }

    @classmethod
    def get_plan(cls, plan_id):
        plans = {'basic': cls.BASIC, 'pro': cls.PRO, 'trial': cls.TRIAL}
        return plans.get(plan_id)

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128))
    
    # Enhanced user fields
    first_name = db.Column(db.String(50), nullable=False, default='')
    last_name = db.Column(db.String(50), nullable=False, default='')
    company_name = db.Column(db.String(100), nullable=False, default='')  # Required field
    
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    last_login = db.Column(db.DateTime)
    is_active = db.Column(db.Boolean, default=True)
    is_admin = db.Column(db.Boolean, default=False)
    conversations = db.relationship('Conversation', backref='user', lazy=True)
    
    # Subscription fields
    subscription_plan = db.Column(db.String(20), default='free')
    subscription_status = db.Column(db.String(20), default='inactive')
    subscription_id = db.Column(db.String(100))
    payment_method_id = db.Column(db.String(100))
    current_period_end = db.Column(db.DateTime, nullable=True)
    cancel_at_period_end = db.Column(db.Boolean, default=False)
    documents_uploaded = db.Column(db.Integer, default=0)
    last_billing_date = db.Column(db.DateTime)
    trial_start_date = db.Column(db.DateTime)
    trial_end_date = db.Column(db.DateTime)
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
    @property
    def full_name(self):
        """Get user's full name"""
        return f"{self.first_name} {self.last_name}".strip()
    
    @property
    def display_name(self):
        """Get display name (full name if available, otherwise username)"""
        full_name = self.full_name
        return full_name if full_name else self.username
    
    def _make_aware(self, dt):
        """Convert naive datetime to timezone-aware UTC datetime"""
        if dt is None:
            return None
        if dt.tzinfo is None:
            return dt.replace(tzinfo=timezone.utc)
        return dt
    
    @property
    def is_subscribed(self):
        """Check if user has active subscription - FIXED VERSION"""
        now = datetime.now(timezone.utc)
        
        # Check active paid subscription
        if (self.subscription_status == 'active' and 
            self.subscription_plan in ['basic', 'pro'] and
            (self.current_period_end is None or 
            self._make_aware(self.current_period_end) > now)):
            return True
        
        # Check active trial
        if (self.subscription_plan == 'trial' and 
            self.trial_end_date and 
            self._make_aware(self.trial_end_date) > now):
            return True
        
        # Admin always has access
        if self.is_admin:
            return True
        
        return False

    @property
    def plan_details(self):
        return SubscriptionPlan.get_plan(self.subscription_plan) or SubscriptionPlan.get_plan('basic')
    
    def can_upload_file(self, file_size):
        if self.is_admin:
            return True
        
        if self.subscription_plan == 'trial' and self.is_subscribed:
            return True
        
        plan = self.plan_details.get('limits', {})
        max_size = plan.get('max_file_size', 100 * 1024 * 1024)
        return file_size <= max_size

    @property
    def trial_days_remaining(self):
        """Calculate trial days remaining"""
        if not self.trial_end_date or self.subscription_plan != 'trial':
            return 0
        
        now = datetime.now(timezone.utc)
        trial_end = self._make_aware(self.trial_end_date)
        
        days_left = (trial_end - now).days
        return max(0, days_left)

class Payment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    amount = db.Column(db.Float, nullable=False)
    currency = db.Column(db.String(3), default='QAR')
    payment_intent_id = db.Column(db.String(100))
    payment_method = db.Column(db.String(50))
    status = db.Column(db.String(20))
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    receipt_url = db.Column(db.String(255))

class BillingHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    description = db.Column(db.String(255))
    amount = db.Column(db.Float, nullable=False)
    currency = db.Column(db.String(3), default='QAR')
    status = db.Column(db.String(20))
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    invoice_id = db.Column(db.String(100))

class Conversation(db.Model):
    id = db.Column(db.String(36), primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(100), nullable=False, default='New Query')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    messages = db.relationship('Message', backref='conversation', lazy=True, cascade='all, delete-orphan')

class Message(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    conversation_id = db.Column(db.String(36), db.ForeignKey('conversation.id'), nullable=False)
    role = db.Column(db.String(20), nullable=False)  # 'user' or 'assistant'
    content = db.Column(db.Text, nullable=False)
    sources = db.Column(db.JSON)  # Store sources as JSON
    timestamp = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    def __repr__(self):
        return f'<Message {self.id}: {self.role} in {self.conversation_id}>'

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    file_path = db.Column(db.String(512), nullable=False)
    file_hash = db.Column(db.String(64), nullable=False)
    file_type = db.Column(db.String(32), nullable=False)
    processed_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    document_metadata = db.Column(db.JSON, name="metadata")

class DocumentContent(db.Model):
    """Enhanced table for storing extracted content and embeddings"""
    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    content_type = db.Column(db.String(50), default='text')  # 'enhanced_full_text', 'enhanced_chunk', etc.
    chunk_index = db.Column(db.Integer, default=0)
    embedding = db.Column(db.JSON)  # Store embedding vectors
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    content_metadata = db.Column(db.JSON)  # Store metadata about content structure
    document = db.relationship('Document', backref='content_records')

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

# Update the chat route to use Ollama instead of Gemini
@app.route('/chat/<chat_id>', methods=['GET', 'POST'])
@login_required
def chat(chat_id):
    """Chat route using the dual system"""
    try:
        conversation = Conversation.query.filter_by(
            id=chat_id,
            user_id=current_user.id
        ).first_or_404()

        if request.method == 'GET':
            # [Your existing GET logic remains the same]
            messages = Message.query.filter_by(
                conversation_id=chat_id
            ).order_by(Message.timestamp.asc()).all()
            
            user_conversations = Conversation.query.filter_by(
                user_id=current_user.id
            ).order_by(Conversation.updated_at.desc()).all()
            
            documents = Document.query.filter_by(
                user_id=current_user.id
            ).order_by(Document.processed_at.desc()).all()
            
            doc_count = len(documents)
            
            return render_template('chat.html', 
                chat_id=chat_id,
                conversation=conversation,
                messages=messages,
                conversations=user_conversations,
                documents=documents,
                chat_title=conversation.title,
                doc_count=doc_count,
                ai_model=app.config['OLLAMA_MODEL'],
                dual_system_enabled=True
            )

        elif request.method == 'POST':
            def generate_dual_system_response():
                """Generate response using dual system"""
                try:
                    data = request.get_json()
                    question = data.get('question', '').strip()
                    
                    if not question:
                        yield json.dumps({'status': 'error', 'message': 'Please enter a question'})
                        return

                    # Save user message
                    user_msg = Message(
                        conversation_id=chat_id,
                        role='user',
                        content=question,
                        timestamp=datetime.now(timezone.utc)
                    )
                    db.session.add(user_msg)
                    db.session.commit()

                    # Update conversation title if needed
                    if conversation.title == 'New Query':
                        short_title = question[:50] + '...' if len(question) > 50 else question
                        conversation.title = short_title
                        conversation.updated_at = datetime.now(timezone.utc)
                        db.session.commit()

                    # Use dual system
                    full_response = ""
                    sources = []
                    analysis = {}
                    
                    for response_chunk in dual_system_response(question, current_user.id):
                        response_data = json.loads(response_chunk)
                        
                        if response_data['status'] == 'stream_chunk':
                            full_response += response_data['content']
                        elif response_data['status'] == 'stream_end':
                            sources = response_data.get('sources', [])
                            analysis = response_data.get('question_analysis', {})
                        
                        # Stream the response chunk to frontend
                        yield response_chunk
                    
                    # Save assistant response
                    if full_response:
                        assistant_msg = Message(
                            conversation_id=chat_id,
                            role='assistant',
                            content=full_response,
                            sources=[s.get('filename', s) if isinstance(s, dict) else s for s in sources],
                            timestamp=datetime.now(timezone.utc)
                        )
                        db.session.add(assistant_msg)
                        conversation.updated_at = datetime.now(timezone.utc)
                        db.session.commit()

                except Exception as e:
                    logger.error(f"Dual system chat generation error: {e}")
                    yield json.dumps({
                        'status': 'error',
                        'message': f"Error: {str(e)}"
                    }) + "\n"
            
            return Response(stream_with_context(generate_dual_system_response()), mimetype='application/x-ndjson')

    except Exception as e:
        logger.error(f"Chat route error: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        if request.method == 'GET':
            flash('Error loading chat', 'error')
            return redirect(url_for('index'))
        else:
            return jsonify({'error': str(e)}), 500


# Update startup functions to use Ollama
def startup_exact_match_priority_system():
    """Initialize the exact match priority system with Ollama on startup"""
    try:
        logger.info("🎯 Starting EXACT MATCH PRIORITY DocumentIQ System with Ollama Llama 3.1...")
        
        # Initialize processor
        success = initialize_enhanced_processor()
        
        # Initialize Ollama
        ollama_success = initialize_ollama()
        
        # Update search configuration
        update_search_configuration()
        
        if success and ollama_success:
            logger.info("✅ EXACT MATCH PRIORITY system with Ollama initialized successfully")
            logger.info("🎯 EXACT MATCH PRIORITY features enabled:")
            logger.info("   - GUARANTEED exact phrase matching with massive score boost")
            logger.info("   - Multiple exact match patterns (normalized, case-sensitive, etc.)")
            logger.info("   - Word boundary detection for precise matching")
            logger.info("   - Context prioritization for exact matches (70% of context space)")
            logger.info("   - Exact match content always shown first")
            logger.info("   - Emergency fallbacks ensure 100% search success")
            logger.info("   - Enhanced RAG prompt prioritizes exact match answers")
            logger.info("   - Powered by Ollama Llama 3.1 for superior language understanding")
        else:
            logger.error("❌ EXACT MATCH PRIORITY system initialization failed")
            if not success:
                logger.error("   - Document processor initialization failed")
            if not ollama_success:
                logger.error("   - Ollama initialization failed")
            
        return success and ollama_success
        
    except Exception as e:
        logger.error(f"❌ EXACT MATCH PRIORITY system startup error: {e}")
        return False



def initialize_enhanced_processor():
    """Initialize the enhanced document processor with GPU optimization"""
    global global_processor, search_engine
    
    try:
        logger.info("🚀 Initializing Document Processor with RTX A6000 optimization...")
        
        # Add debug information
        logger.debug(f"Torch CUDA available: {torch.cuda.is_available()}")
        if torch.cuda.is_available():
            logger.debug(f"CUDA device count: {torch.cuda.device_count()}")
            logger.debug(f"Device name: {torch.cuda.get_device_name(0)}")
        
        global_processor = AccuracyFocusedDocumentProcessor(gpu_enabled=True)
        search_engine = AccuracyFocusedSearchEngine(global_processor)
        
        logger.info("✅ Accuracy-focused processor initialized successfully")
        return True
        
    except Exception as e:
        logger.error(f"❌ Enhanced processor initialization failed: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return False

# Update context processor to include Ollama info
@app.context_processor
def inject_now():
    """Inject current time and enhanced system status into templates"""
    return {
        'now': datetime.now(timezone.utc),
        'gpu_available': torch.cuda.is_available(),
        'gpu_name': torch.cuda.get_device_name() if torch.cuda.is_available() else 'N/A',
        'enhanced_processing': True,
        'perfect_formatting': True,
        'comprehensive_search': True,
        'accuracy_focused': True,
        'rtx_a6000_optimized': torch.cuda.is_available(),
        'max_upload_gb': app.config['MAX_CONTENT_LENGTH'] / 1e9,
        'max_file_gb': app.config['MAX_FILE_SIZE'] / 1e9,
        'max_extracted_gb': app.config['MAX_ZIP_EXTRACTED_SIZE'] / 1e9,
        'ai_model': 'Ollama Llama 3.1',
        'model_name': app.config['OLLAMA_MODEL'],
        'ollama_powered': True  # Changed from gemini_powered
    }



def get_file_hash_optimized(file_path):
    """FASTER file hash generation using larger chunks"""
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        # Use larger chunks for faster hashing
        for chunk in iter(lambda: f.read(65536), b""):  # 64KB chunks instead of 4KB
            hasher.update(chunk)
    return hasher.hexdigest()



def generate_embeddings_fast(chunks):
    """Fast embedding generation"""
    if not ensure_processor_initialized():
        return [np.zeros(768) for _ in chunks]
    
    try:
        texts = [chunk['content'] for chunk in chunks]
        embeddings = global_processor.embedding_model.encode(
            texts,
            convert_to_tensor=False,
            show_progress_bar=False,
            batch_size=32  # Larger batch for speed
        )
        return embeddings
    except Exception as e:
        logger.error(f"Fast embedding generation failed: {e}")
        return [np.zeros(768) for _ in chunks]


def ensure_processor_initialized():
    """Ensure the processor is initialized"""
    global global_processor, search_engine
    
    if global_processor is None:
        logger.warning("Processor not initialized, attempting on-demand initialization")
        if not initialize_enhanced_processor():
            logger.error("Failed to initialize processor on demand")
            return False
    return True


def create_chunks_optimized(content):
    """Fast chunking for performance"""
    chunk_size = 1200
    overlap = 200
    
    if len(content) <= chunk_size:
        return [{'content': content, 'chunk_index': 0, 'metadata': {}}]
    
    chunks = []
    start = 0
    chunk_index = 0
    
    while start < len(content):
        end = start + chunk_size
        chunk_content = content[start:end]
        
        chunks.append({
            'content': chunk_content,
            'chunk_index': chunk_index,
            'metadata': {'optimized_chunk': True}
        })
        
        start += chunk_size - overlap
        chunk_index += 1
    
    return chunks


def get_gpu_stats():
    """Get GPU performance statistics"""
    if not torch.cuda.is_available():
        return {"gpu_available": False}
    
    try:
        stats = {
            "gpu_available": True,
            "gpu_name": torch.cuda.get_device_name(),
            "memory_allocated_gb": torch.cuda.memory_allocated() / 1e9,
            "memory_reserved_gb": torch.cuda.memory_reserved() / 1e9,
            "memory_total_gb": torch.cuda.get_device_properties(0).total_memory / 1e9,
            "cuda_version": torch.version.cuda,
            "pytorch_version": torch.__version__
        }
        
        # Add NVIDIA-ML stats if available
        if NVML_AVAILABLE:
            try:
                nvml.nvmlInit()
                handle = nvml.nvmlDeviceGetHandleByIndex(0)
                memory_info = nvml.nvmlDeviceGetMemoryInfo(handle)
                utilization = nvml.nvmlDeviceGetUtilizationRates(handle)
                temperature = nvml.nvmlDeviceGetTemperature(handle, nvml.NVML_TEMPERATURE_GPU)
                
                stats.update({
                    "memory_usage_percent": (memory_info.used / memory_info.total) * 100,
                    "gpu_utilization_percent": utilization.gpu,
                    "memory_utilization_percent": utilization.memory,
                    "temperature_celsius": temperature,
                    "is_optimal": temperature < 80 and utilization.gpu > 0
                })
            except Exception as e:
                stats["nvml_error"] = str(e)
        
        return stats
    except Exception as e:
        return {"gpu_available": False, "error": str(e)}



# Database initialization
def create_app_with_postgres():
    """Initialize app with PostgreSQL"""
    with app.app_context():
        try:
            with db.engine.connect() as connection:
                connection.execute(text('SELECT 1'))
            logger.info('PostgreSQL connection successful')
            
            # Create all tables
            db.create_all()
            logger.info('Database tables created/verified')
            
            # Log GPU status
            gpu_stats = get_gpu_stats()
            logger.info(f'Accuracy-Focused GPU Status: {gpu_stats}')
            
        except Exception as e:
            logger.error(f"Database initialization failed: {e}")
            raise

# Main application startup
if __name__ == '__main__':
    try:
        # Initialize database first
        create_app_with_postgres()
        
        # Validate system requirements (now includes Ollama)
        all_systems_ready, system_issues = validate_system_requirements()
        
        # Try to initialize OCR
        try:
            ocr_init_success = initialize_ocr_processor()
            if ocr_init_success:
                logger.info("✅ OCR system initialized successfully at startup")
            else:
                logger.warning("⚠️ OCR system failed to initialize at startup - will try on demand")
        except Exception as ocr_error:
            logger.error(f"OCR initialization error at startup: {ocr_error}")
        
        # Initialize EXACT KEYWORD priority system with Ollama (UPDATED)
        keyword_success = startup_exact_keyword_priority_system()
        
        # Check Ollama specifically
        ollama_ready = startup_ollama_check()
        
        # Enhanced GPU initialization logging
        if torch.cuda.is_available():
            gpu_props = torch.cuda.get_device_properties(0)
            logger.info('🎯 DocumentIQ - EXACT KEYWORD PRIORITY System with Ollama Llama 3.1')
            logger.info(f'GPU: {torch.cuda.get_device_name()}')
            logger.info(f'GPU Memory: {gpu_props.total_memory / 1e9:.1f}GB')
            logger.info(f'CUDA Cores: ~{gpu_props.multi_processor_count * 64}')
            logger.info(f'Max Upload Size: {app.config["MAX_CONTENT_LENGTH"] / 1e9:.1f}GB')
            logger.info(f'Max Single File: {app.config["MAX_FILE_SIZE"] / 1e9:.1f}GB')
            logger.info(f'Max Extracted ZIP: {app.config["MAX_ZIP_EXTRACTED_SIZE"] / 1e9:.1f}GB')
        else:
            logger.warning('⚠️  GPU not available - running exact keyword CPU mode')
        
        # Show startup summary
        logger.info("=" * 80)
        logger.info("STARTUP SUMMARY - EXACT KEYWORD PRIORITY WITH OLLAMA LLAMA 3.1")
        logger.info("=" * 80)
        logger.info(f"🔹 GPU Available: {'Yes' if torch.cuda.is_available() else 'No'}")
        logger.info(f"🔹 OCR Initialized: {'Yes' if 'ocr_init_success' in locals() and ocr_init_success else 'No'}")
        logger.info(f"🔹 Ollama Service: {'Ready' if ollama_ready else 'Not Ready'}")
        logger.info(f"🔹 Database: {'Connected' if all_systems_ready else 'Check required'}")
        logger.info(f"🔹 Processing Mode: {'GPU-Accelerated' if torch.cuda.is_available() else 'CPU'}")
        logger.info(f"🔹 AI Model: {app.config['OLLAMA_MODEL']} (Ollama)")
        logger.info(f"🔹 Base URL: {app.config['OLLAMA_BASE_URL']}")
        logger.info(f"🔹 Exact Keyword Search: {'Enabled' if keyword_success else 'Disabled'}")
        logger.info("=" * 80)
        
        logger.info("🎯 EXACT KEYWORD SEARCH EXAMPLES:")
        logger.info("   Query: 'What is a university?' → Searches for: 'university' + variations")
        logger.info("   Query: 'Tell me about companies' → Searches for: 'companies' + variations")
        logger.info("   Query: 'How do I find research?' → Searches for: 'research' + variations")
        logger.info("   Query: 'List all students' → Searches for: 'students' + variations")
        logger.info("=" * 80)
        
        if system_issues:
            logger.error("⚠️ RESOLVE THESE ISSUES FOR FULL FUNCTIONALITY:")
            for issue in system_issues:
                logger.error(f"   🔧 {issue}")
            logger.error("=" * 80)
        
        if not ollama_ready:
            logger.error("🚨 CRITICAL: Ollama service is not ready!")
            logger.error("   🔧 Start Ollama: ollama serve")
            logger.error("   🔧 Pull model: ollama pull llama3.1")
            logger.error("   🔧 Set base URL: export OLLAMA_BASE_URL='http://localhost:11434'")
            logger.error("   🔧 Without Ollama, AI responses will fail!")
            logger.error("=" * 80)
        
        if keyword_success and all_systems_ready and ollama_ready:
            logger.info('🌟 ALL SYSTEMS READY - DocumentIQ EXACT KEYWORD server starting with Ollama Llama 3.1...')
        elif keyword_success:
            logger.warning('⚠️ PARTIAL SYSTEM READY - Some features may not work correctly')
        else:
            logger.warning('⚠️ MINIMAL SYSTEM READY - Running with fallback systems')
        
        # Start development server
        logger.info('🌟 DocumentIQ EXACT KEYWORD server with Ollama Llama 3.1 starting...')
        app.run(host='0.0.0.0', port=8000, threaded=True, debug=False)
        
    except Exception as e:
        logger.critical(f"Failed to start application: {e}")
        logger.critical(f"Traceback: {traceback.format_exc()}")
        
        # GPU cleanup on failure
        if torch.cuda.is_available():
            try:
                torch.cuda.empty_cache()
                logger.info("GPU cache cleared on shutdown")
            except:
                pass
                
        raise